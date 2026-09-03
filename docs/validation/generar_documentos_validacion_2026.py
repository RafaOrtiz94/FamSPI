"""Genera el paquete documental de validacion retrospectiva FamSPI 2026.

Fuente documental principal:
- docs/validation/01_primeros_pasos_y_dq.md
- docs/validation/02_implementacion_iq.md
- docs/validation/03_operacion_oq.md
- docs/validation/04_desempeno_pq.md

La estructura se alinea con WHO TRS 1019 Annex 3, Appendix 5
para sistemas computarizados: DQ, IQ, OQ, SOP/training, PQ/UAT,
operacion/mantenimiento y estrategia retrospectiva para sistemas legacy.
"""

from __future__ import annotations

import json
import re
import subprocess
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


VALIDATION_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = VALIDATION_DIR.parents[1]
TRACEABILITY_FILE = VALIDATION_DIR / "traceability_map.json"
JEST_RESULTS_FILE = VALIDATION_DIR / "evidence" / "jest-results.json"
LOGO_PATH = VALIDATION_DIR / "assets" / "logo_famproject.png"

VERSION = "2026.1"
SYSTEM_NAME = "FamSPI"
COMPANY = "Famproject Cia. Ltda."
DEPARTMENT = "Departamento de Tecnologias de la Informacion y Comunicacion (TICS)"
TODAY = date.today().strftime("%d/%m/%Y")

WHO_REFERENCE = "WHO TRS 1019 Annex 3, Appendix 5 - Validation of computerized systems"
WHO_DATA_INTEGRITY_REFERENCE = "WHO TRS 1033 Annex 4 - Guideline on data integrity (ALCOA+)"

# Ventana historica evaluada (WHO §6.6, §11.4). Se define antes de revisar
# resultados y cubre el historial de cambios y operacion evidenciable por git
# y por la pista de auditoria del sistema.
HISTORICAL_PERIOD_START = "01/08/2025"
HISTORICAL_PERIOD_END = TODAY
HISTORICAL_PERIOD = (
    f"{HISTORICAL_PERIOD_START} a {HISTORICAL_PERIOD_END} "
    "(12 meses de historial de cambios y operacion cubiertos por el control de "
    "versiones git y la pista de auditoria del sistema)"
)

# Paleta monocromatica alineada al formato corporativo (docs/Formato
# Validacion.docx): un unico color (azul corporativo 0F4761) para
# titulos/subtitulos y para encabezados de tabla; sin variacion de tonos
# entre secciones. Se conservan los nombres historicos de las constantes
# para no tocar cada punto de uso, pero todas apuntan al mismo valor por
# categoria.
CORP_BLUE = "0F4761"
NAVY = CORP_BLUE
BLUE = CORP_BLUE
GREEN = CORP_BLUE
MINT = "F2F2F2"
SOFT_BLUE = "F2F2F2"
SOFT_GREEN = "F2F2F2"
SOFT_AMBER = "F2F2F2"
SOFT_RED = CORP_BLUE
LIGHT = "F2F2F2"
TEXT = "000000"
MUTED = "595959"
WHITE = "FFFFFF"


@dataclass(frozen=True)
class PhaseDocument:
    source: str
    output: str
    code: str
    title: str
    phase: str
    purpose: str
    who_focus: tuple[str, ...]
    include_test_evidence: bool = False


@dataclass(frozen=True)
class EvidenceRow:
    module: str
    req_id: str
    description: str
    test_file: str
    status: str
    duration_ms: int | None
    test_line: int | None = None


@dataclass(frozen=True)
class GapRow:
    module: str
    risk: str
    last_commit: str | None
    commits_last_12_months: int
    justification: str


@dataclass(frozen=True)
class TestEvidence:
    generated_at: str
    run_at: str
    success: bool
    total: int
    passed: int
    failed: int
    rows: list[EvidenceRow] = field(default_factory=list)
    gaps: list[GapRow] = field(default_factory=list)


STATUS_LABEL = {
    "passed": "Conforme",
    "failed": "No conforme",
    "pending": "Pendiente",
    "sin_evidencia": "Sin evidencia",
}


def load_test_evidence() -> TestEvidence:
    """Carga evidencia REAL: cruza traceability_map.json (requisito -> test)
    con evidence/jest-results.json (resultado real de la ultima corrida de
    `npm run test:validation`). No genera datos: falla si falta cualquiera
    de las dos fuentes, para no permitir un documento con evidencia inventada.
    """
    if not TRACEABILITY_FILE.exists():
        raise FileNotFoundError(
            "Falta docs/validation/traceability_map.json. "
            "Correr primero: node docs/validation/build_traceability_map.js"
        )
    if not JEST_RESULTS_FILE.exists():
        raise FileNotFoundError(
            "Falta docs/validation/evidence/jest-results.json. "
            "Correr primero: cd backend && npm run test:validation"
        )

    trace = json.loads(TRACEABILITY_FILE.read_text(encoding="utf-8"))
    jest = json.loads(JEST_RESULTS_FILE.read_text(encoding="utf-8"))

    # Cola de resultados reales por archivo de test, en el orden en que
    # Jest los reporta (mismo orden secuencial que build_traceability_map.js
    # usa al extraer describe/it, por lo que el emparejamiento por indice
    # es valido dentro de un mismo archivo).
    queues: dict[str, list[dict]] = {}
    for suite in jest.get("testResults", []):
        abs_path = Path(suite["name"])
        try:
            rel = abs_path.relative_to(PROJECT_ROOT).as_posix()
        except ValueError:
            rel = abs_path.as_posix()
        queues[rel] = list(suite.get("assertionResults", []))

    rows: list[EvidenceRow] = []
    for module_name, module_data in trace["modules"].items():
        if not module_data.get("has_automated_tests"):
            continue
        for req in module_data.get("requirements", []):
            test_file = req["test_file"]
            queue = queues.get(test_file, [])
            assertion = queue.pop(0) if queue else None
            status = assertion["status"] if assertion else "sin_evidencia"
            duration = assertion.get("duration") if assertion else None
            rows.append(
                EvidenceRow(
                    module=module_name,
                    req_id=req["req_id"],
                    description=req["description"],
                    test_file=test_file,
                    status=status,
                    duration_ms=duration,
                    test_line=req.get("test_line"),
                )
            )

    gaps: list[GapRow] = []
    for module_name, module_data in trace["modules"].items():
        if module_data.get("has_automated_tests"):
            continue
        git_evidence = module_data.get("git_evidence", {})
        gaps.append(
            GapRow(
                module=module_name,
                risk=module_data.get("risk_level", "sin clasificar"),
                last_commit=git_evidence.get("last_commit_at"),
                commits_last_12_months=git_evidence.get("commits_last_12_months", 0),
                justification=module_data.get("gap_justification", ""),
            )
        )

    run_at_ms = jest.get("startTime")
    run_at = (
        datetime.fromtimestamp(run_at_ms / 1000).strftime("%d/%m/%Y %H:%M:%S")
        if run_at_ms
        else "desconocido"
    )

    return TestEvidence(
        generated_at=trace.get("generated_at", "desconocido"),
        run_at=run_at,
        success=bool(jest.get("success")),
        total=jest.get("numTotalTests", 0),
        passed=jest.get("numPassedTests", 0),
        failed=jest.get("numFailedTests", 0),
        rows=rows,
        gaps=gaps,
    )


PHASE_DOCUMENTS = (
    PhaseDocument(
        source="01_primeros_pasos_y_dq.md",
        output="FAMSPI_01_PLAN_MAESTRO_DQ_2026.docx",
        code="FAM-SPI-VAL-DQ-2026",
        title="Plan Maestro y Calificacion de Diseno (DQ)",
        phase="Plan maestro de validacion, gestion de proveedores, desarrollo del sistema y Design Qualification",
        purpose=(
            "Establecer el plan maestro de validacion retrospectiva, la gestion de "
            "proveedores, el control de desarrollo del sistema y la calificacion de "
            "diseno (DQ) que sustentan las siguientes fases de calificacion."
        ),
        who_focus=(
            "Introduccion y alcance (§1)",
            "Protocolos y reportes de validacion (§3)",
            "Gestion de proveedores (§4)",
            "Especificaciones de requerimientos (§5)",
            "Especificacion de diseno/configuracion (§6)",
            "Design qualification (§7)",
            "Desarrollo del sistema e implementacion del proyecto (§8)",
            "Estrategia retrospectiva para sistemas existentes (§12.6-12.10)",
        ),
    ),
    PhaseDocument(
        source="02_implementacion_iq.md",
        output="FAMSPI_02_IMPLEMENTACION_IQ_2026.docx",
        code="FAM-SPI-VAL-IQ-2026",
        title="Implementacion IQ",
        phase="Installation Qualification",
        purpose=(
            "Demostrar que infraestructura, componentes, dependencias, rutas, "
            "variables y despliegues estan instalados bajo control."
        ),
        who_focus=(
            "Desarrollo del proyecto e implementacion",
            "Installation qualification",
            "Control documental",
            "Evidencia de instalacion y configuracion",
        ),
    ),
    PhaseDocument(
        source="03_operacion_oq.md",
        output="FAMSPI_03_OPERACION_OQ_2026.docx",
        code="FAM-SPI-VAL-OQ-2026",
        title="Operacion OQ",
        phase="Operational Qualification y SOPs/entrenamiento",
        purpose=(
            "Verificar que los flujos configurados operan conforme a permisos, "
            "estados, contratos API, excepciones y controles definidos, y confirmar "
            "que existen SOPs y material de entrenamiento vigentes antes de PQ."
        ),
        who_focus=(
            "Operational qualification (§10)",
            "Procedimientos operativos y entrenamiento (§11)",
            "Controles funcionales",
            "Operacion en condiciones esperadas y excepcionales",
        ),
        include_test_evidence=True,
    ),
    PhaseDocument(
        source="04_desempeno_pq.md",
        output="FAMSPI_04_DESEMPENO_PQ_2026.docx",
        code="FAM-SPI-VAL-PQ-2026",
        title="Desempeno PQ",
        phase="Performance Qualification / User Acceptance Testing",
        purpose=(
            "Confirmar que FamSPI sostiene el proceso real de negocio en uso "
            "productivo, con evidencia de aceptacion y desempeno operacional."
        ),
        who_focus=(
            "Performance qualification",
            "User acceptance testing",
            "Operacion y mantenimiento",
            "Revision periodica y control de cambios",
        ),
        include_test_evidence=True,
    ),
)


RETROSPECTIVE_CONTROLS = (
    (
        "Evaluacion de riesgo",
        "Determinar criticidad GxP/negocio por modulo, datos, aprobaciones y trazabilidad.",
        "Matriz de riesgo, impacto por modulo, justificacion de alcance.",
    ),
    (
        "Analisis de brechas",
        "Comparar evidencia disponible contra URS, DQ, IQ, OQ, PQ, SOP y entrenamiento.",
        "Gap log con responsable, fecha objetivo, severidad y decision CAPA.",
    ),
    (
        "Historia de uso",
        "Revisar uso real, incidentes, errores, mantenimiento y cambios ejecutados.",
        "Historial de tickets TI, codigos support_tickets.code, commits, despliegues, bitacoras y errores relevantes.",
    ),
    (
        "Control de cambios con ticket TI",
        "Confirmar que cada cambio validado tiene ticket TI asociado y ciclo verificable.",
        "Ticket en support_tickets, eventos en support_ticket_events, comentarios, estado resuelto/cerrado o cancelacion justificada.",
    ),
    (
        "Estado de control",
        "Confirmar que parametros actuales, roles, rutas, procesos y practicas siguen vigentes.",
        "Revision de configuracion actual y comparacion contra operacion validada.",
    ),
    (
        "Requerimientos y descripcion",
        "Mantener URS, descripcion del sistema, arquitectura y limites del sistema actualizados.",
        "URS aprobada, mapa de modulos, flujos, APIs, tablas y responsabilidades.",
    ),
    (
        "Aceptacion tecnica",
        "Evidenciar aceptacion de software, hardware, red, perifericos, procesos e integraciones.",
        "IQ/OQ/PQ, pruebas de regresion, evidencia de despliegue y aceptacion de usuarios.",
    ),
)


def _add_field_run(paragraph, field_code: str, *, size: int = 8, color: str = "667085"):
    run = paragraph.add_run()
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(color)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    return run


def _apply_heading_divider(style, color: str) -> None:
    p_pr = style.element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _cell(
    cell,
    text: object,
    *,
    bold: bool = False,
    size: int = 8,
    color: str = TEXT,
    fill: str | None = None,
    center: bool = False,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(color)
    if fill:
        _shade(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_widths(table, widths_cm: Iterable[float]) -> None:
    widths = list(widths_cm)
    for row in table.rows:
        for index, width in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = Cm(width)


def _table(
    doc: Document,
    headers: list[str],
    rows: Iterable[Iterable[object]],
    *,
    widths: Iterable[float] | None = None,
    header_fill: str = NAVY,
    font_size: int = 8,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for index, header in enumerate(headers):
        _cell(
            table.rows[0].cells[index],
            header,
            bold=True,
            size=font_size,
            color=WHITE,
            fill=header_fill,
            center=True,
        )

    for row_index, row_values in enumerate(rows):
        fill = LIGHT if row_index % 2 == 0 else WHITE
        row = table.add_row()
        for col_index, value in enumerate(row_values):
            _cell(row.cells[col_index], value, size=font_size, fill=fill)

    if widths:
        _set_widths(table, widths)
    doc.add_paragraph()


def _build_official_header(section, title: str, code: str) -> None:
    """Encabezado corporativo (docs/Formato Validacion.docx): logo, titulo,
    compania/departamento + fecha, y codigo del documento. Se repite en
    todas las paginas, igual que en la plantilla oficial."""
    header = section.header
    header.is_linked_to_previous = False
    # Limpia el parrafo por defecto que crea python-docx.
    default_p = header.paragraphs[0]._p
    default_p.getparent().remove(default_p)

    table = header.add_table(rows=4, cols=2, width=Cm(15.0))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_widths(table, [7.5, 7.5])

    # Fila 0: logo, ocupa las dos columnas, alineado a la derecha.
    logo_cell = table.cell(0, 0).merge(table.cell(0, 1))
    logo_p = logo_cell.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO_PATH.exists():
        logo_p.add_run().add_picture(str(LOGO_PATH), width=Cm(2.6))

    # Fila 1: titulo del documento, ocupa las dos columnas, centrado y negrita.
    title_cell = table.cell(1, 0).merge(table.cell(1, 1))
    title_p = title_cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(10)
    title_run.font.color.rgb = RGBColor.from_string(TEXT)

    # Fila 2: compania/departamento (izquierda) y fecha (derecha).
    left_cell = table.cell(2, 0)
    for i, line in enumerate((COMPANY, DEPARTMENT)):
        p = left_cell.paragraphs[0] if i == 0 else left_cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(TEXT)

    right_cell = table.cell(2, 1)
    right_cell.add_paragraph()
    date_p = right_cell.paragraphs[-1]
    date_run = date_p.add_run(f"FECHA: {TODAY}")
    date_run.font.name = "Arial"
    date_run.font.size = Pt(8)
    date_run.font.color.rgb = RGBColor.from_string(TEXT)

    # Fila 3: codigo del documento, ocupa las dos columnas.
    code_cell = table.cell(3, 0).merge(table.cell(3, 1))
    code_p = code_cell.paragraphs[0]
    code_run = code_p.add_run(code)
    code_run.font.name = "Arial"
    code_run.font.size = Pt(8)
    code_run.font.color.rgb = RGBColor.from_string(TEXT)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _build_official_footer(section) -> None:
    """Pie de pagina corporativo: numero de pagina + aviso de
    confidencialidad, igual que docs/Formato Validacion.docx."""
    footer = section.footer
    footer.is_linked_to_previous = False

    page_p = footer.paragraphs[0]
    page_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_field_run(page_p, "PAGE", size=9, color=CORP_BLUE)
    for run in page_p.runs:
        run.bold = True
        run.font.name = "Arial"

    notice_p = footer.add_paragraph()
    notice_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice_run = notice_p.add_run(
        f"Este documento es confidencial, la informacion contenida en el es propiedad de {COMPANY.rstrip('.')}."
    )
    notice_run.italic = True
    notice_run.font.name = "Arial"
    notice_run.font.size = Pt(8)
    notice_run.font.color.rgb = RGBColor.from_string(TEXT)


def _style_document(doc: Document, title: str, code: str = "") -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (
        ("Title", 24, NAVY),
        ("Heading 1", 15, NAVY),
        ("Heading 2", 12, BLUE),
        ("Heading 3", 10, GREEN),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(6)

    _apply_heading_divider(styles["Heading 1"], BLUE)

    _build_official_header(section, title, code)
    _build_official_footer(section)


def _add_cover_card(doc: Document, rows: Iterable[tuple[str, str]]) -> None:
    """Ficha de identificacion sin encabezado: etiqueta sombreada + valor."""
    data = list(rows)
    table = doc.add_table(rows=len(data), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(data):
        _cell(table.rows[i].cells[0], label, bold=True, size=9, color=NAVY, fill=SOFT_BLUE)
        _cell(table.rows[i].cells[1], value, size=9, fill=WHITE)
    _set_widths(table, [4.6, 12.2])
    doc.add_paragraph()


def _add_accent_box(doc: Document, text: str, *, fill: str = SOFT_GREEN, color: str = TEXT) -> None:
    """Caja de acento a todo el ancho para destacar el proposito/nota."""
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box.rows[0].cells[0]
    _shade(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(color)
    _set_widths(box, [16.8])
    doc.add_paragraph()


def _add_cover(doc: Document, phase: PhaseDocument) -> None:
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Expediente de validacion retrospectiva de sistema computarizado")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph()

    _add_cover_card(
        doc,
        [
            ("Codigo del documento", phase.code),
            ("Compania", COMPANY),
            ("Sistema", SYSTEM_NAME),
            ("Elaborado por", DEPARTMENT),
            ("Version documental", VERSION),
            ("Fecha de emision", TODAY),
            ("Norma base", WHO_REFERENCE),
            ("Norma de integridad de datos", WHO_DATA_INTEGRITY_REFERENCE),
            ("Fase de calificacion", phase.phase),
            ("Clasificacion", "Confidencial - uso interno controlado"),
            ("Estado", "Borrador controlado para revision/aprobacion"),
        ],
    )

    _add_accent_box(doc, phase.purpose, fill=SOFT_BLUE, color=NAVY)
    doc.add_page_break()


def _add_revision_history(doc: Document) -> None:
    """Historial de cambios del documento (WHO §8: control documental)."""
    doc.add_heading("Historial de cambios del documento", level=1)
    _table(
        doc,
        ["Version", "Fecha", "Autor", "Descripcion del cambio", "Estado"],
        [
            (
                VERSION,
                TODAY,
                DEPARTMENT,
                "Emision inicial del expediente retrospectivo generado con evidencia real "
                "(suite de pruebas, historial de git y configuracion vigente).",
                "Borrador controlado",
            ),
        ],
        widths=(1.8, 2.2, 4.6, 6.4, 1.8),
        header_fill=NAVY,
        font_size=8,
    )


def _add_table_of_contents(doc: Document) -> None:
    """Indice automatico (campo TOC de Word; se actualiza con F9 al abrir)."""
    doc.add_heading("Contenido", level=1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Actualice el indice en Word: clic derecho sobre esta tabla > Actualizar campos (o F9)."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)
    doc.add_page_break()


def _add_who_alignment(doc: Document, focus_items: Iterable[str]) -> None:
    doc.add_heading("Alineacion con WHO Appendix 5", level=1)
    doc.add_paragraph(
        "Este documento usa WHO TRS 1019 Annex 3 Appendix 5 como marco para "
        "validacion de sistemas computarizados. Para FamSPI se aplica con enfoque "
        "retrospectivo: se parte del sistema productivo existente, se documentan "
        "brechas y se confirma el estado de control actual."
    )
    _table(
        doc,
        ["Elemento WHO Appendix 5", "Aplicacion en FamSPI"],
        ((item, "Incluido como criterio de evidencia, revision o aceptacion.") for item in focus_items),
        widths=(6.5, 10.0),
        header_fill=GREEN,
    )


def _add_retrospective_strategy(doc: Document) -> None:
    doc.add_heading("Estrategia retrospectiva", level=1)
    doc.add_paragraph(
        "La validacion retrospectiva no se limita a historicos. La evidencia "
        "historica se acepta solo cuando demuestra que el estado actual del sistema "
        "sigue bajo control y que las practicas, parametros y procesos no cambiaron "
        "de forma que invaliden la conclusion."
    )
    _table(
        doc,
        ["Control", "Criterio de revision", "Evidencia esperada"],
        RETROSPECTIVE_CONTROLS,
        widths=(4.0, 6.4, 6.1),
        header_fill=NAVY,
    )


def _add_retrospective_document_controls(doc: Document) -> None:
    """Controles documentales para documento reconstruido (WHO §8):
    leyenda, fecha real, periodo historico, fuentes, supuestos, limitaciones
    y distincion entre lo observado directamente y lo inferido."""
    doc.add_heading("Control documental retrospectivo (WHO §8)", level=1)
    legend = doc.add_paragraph()
    legend_run = legend.add_run("DOCUMENTO RETROSPECTIVO")
    legend_run.bold = True
    legend_run.font.size = Pt(12)
    legend_run.font.color.rgb = RGBColor.from_string(GREEN)
    doc.add_paragraph(
        "Este documento fue reconstruido en la fecha real indicada, sobre un sistema "
        "ya operativo. No representa una aprobacion emitida antes de la instalacion "
        "original y no contiene firmas ni fechas retrofechadas. Su contenido describe "
        "el uso y la configuracion vigentes (as-is), verificados contra el codigo, la "
        "suite de pruebas y el historial real del sistema."
    )
    _table(
        doc,
        ["Campo (WHO §8)", "Contenido"],
        [
            ("Tipo de documento", "Reconstruccion retrospectiva del estado actual (as-is)."),
            ("Fecha real de elaboracion", TODAY),
            ("Periodo historico evaluado", HISTORICAL_PERIOD),
            (
                "Version/configuracion evaluada",
                "Linea base productiva declarada por TICS al " + TODAY
                + " (working tree vigente del repositorio FamSPI).",
            ),
            (
                "Fuentes de reconstruccion",
                "Codigo real (registerRoutes.js, AppRoutes.jsx), suite Jest ejecutada, "
                "historial de git, CONTEXT.md/AGENTS.md por modulo y tickets TI.",
            ),
            (
                "Supuestos",
                "El working tree evaluado corresponde a lo desplegado; los modulos con "
                "mantenimiento activo reciente permanecen bajo control; la pista de "
                "auditoria del sistema no fue deshabilitada durante el periodo.",
            ),
            (
                "Limitaciones",
                f"{MODULES_WITHOUT_TESTS} de {MODULE_COUNT} modulos en alcance no tienen prueba "
                "de verificacion directa y se sustentan con "
                "evidencia de mantenimiento (brecha declarada, WHO §12.8); la linea base "
                "es el working tree vigente y no un commit congelado; algunos requisitos "
                "se infieren del historial (tipo H).",
            ),
            (
                "Observado directamente",
                "Rutas montadas, resultados de pruebas, configuracion de seguridad, "
                "estructura modular y evidencia de cambios en git.",
            ),
            (
                "Inferido y sustentado por historial",
                "Requisitos operativos de modulos sin prueba de verificacion directa, cubiertos por "
                "evidencia de mantenimiento y riesgo declarado.",
            ),
            ("Responsable de verificar", DEPARTMENT),
        ],
        widths=(4.5, 12.0),
        header_fill=GREEN,
        font_size=8,
    )


ALCOA_ROWS = (
    (
        "Atribuible",
        "¿Se conoce quien realizo cada accion?",
        "Autenticacion JWT obligatoria por usuario individual; auditoria automatica de "
        "operaciones POST/PUT/PATCH/DELETE con usuario y accion.",
        "`middlewares/auth.js` (verifyToken), `auditMiddleware`, tabla `auditoria.logs`",
    ),
    (
        "Legible",
        "¿El dato y su contexto se leen durante toda la retencion?",
        "Datos en PostgreSQL con exportaciones legibles y documentos conservados en Drive.",
        "Base Neon (PostgreSQL), `driveClientManager.js`",
    ),
    (
        "Contemporaneo",
        "¿El registro se genera al momento de la actividad?",
        "Sellos de tiempo generados por el servidor (NOW()) en cada transaccion y evento de auditoria.",
        "Consultas SQL con NOW(); zona horaria America/Guayaquil",
    ),
    (
        "Original",
        "¿Se conserva el dato fuente o copia exacta verificada?",
        "Registro electronico nativo en la base de datos; respaldo periodico verificable.",
        "PostgreSQL nativo, `databaseBackupToDrive.js`",
    ),
    (
        "Exacto",
        "¿El dato es correcto y esta protegido?",
        "Validaciones de entrada por modulo, control de acceso por rol y revision/aprobacion en flujos de firma.",
        "`middlewares/roles.js` (requireRole), `signature-workflows`",
    ),
    (
        "Completo",
        "¿Incluye repeticiones, errores y cambios?",
        "La pista de auditoria registra el historial integro de operaciones sobre datos criticos.",
        "`auditMiddleware`, tabla `auditoria.logs`",
    ),
    (
        "Consistente",
        "¿Secuencia, fecha, hora y relaciones son coherentes?",
        "Reloj de servidor unico, claves referenciales en el esquema y orden cronologico de eventos.",
        "Zona horaria unica, restricciones referenciales PostgreSQL",
    ),
    (
        "Duradero",
        "¿Permanece protegido durante el periodo requerido?",
        "Respaldo comprimido periodico enviado a almacenamiento externo con endpoint protegido.",
        "`jobs/databaseBackupToDrive.js`, `middlewares/jobsAuth.js`",
    ),
    (
        "Disponible",
        "¿Puede recuperarse oportunamente para revision?",
        "Datos y documentos recuperables; prueba de restauracion como control operativo requerido.",
        "Respaldo en Drive; procedimiento de restauracion (WHO §13.11)",
    ),
)

DIRA_ROWS = (
    (
        "Firmas y aprobaciones documentales",
        "Modulo signature / signature-workflows",
        "Captura y verificacion electronica multifirmante",
        "Pista de auditoria + verificador publico",
        "Alteracion de firma sin trazabilidad",
        "Bajo",
    ),
    (
        "Aprobaciones de business case",
        "Modulo business-case",
        "Flujo con estados y gate de determinaciones",
        "Auditoria + pruebas de verificacion (37 casos)",
        "Aprobacion no atribuible",
        "Bajo",
    ),
    (
        "Registros de viaticos y finanzas",
        "Modulos finanzas / viaticos",
        "Captura manual con revision y aprobacion",
        "Auditoria de operaciones; control de acceso por rol",
        "Modificacion no controlada (modulo sin prueba de verificacion directa)",
        "Medio",
    ),
    (
        "Marcaciones de asistencia",
        "Modulo attendance",
        "Captura con geolocalizacion y sello horario del servidor",
        "Pruebas de verificacion (82 casos) + auditoria",
        "Marca fuera de flujo esperado",
        "Bajo",
    ),
)


def _add_data_integrity_alcoa(doc: Document) -> None:
    """Integridad de datos ALCOA+ y DIRA (WHO TRS 1033 Annex 4; guia §4, §16.4)."""
    doc.add_heading("Integridad de datos ALCOA+ (WHO TRS 1033 Annex 4)", level=1)
    doc.add_paragraph(
        "FamSPI genera y conserva datos GxP/de negocio criticos. Conforme a WHO TRS 1033 "
        "Annex 4, TICS mapea los principios ALCOA+ contra los controles reales del "
        "sistema. Una impresion en PDF no sustituye el registro electronico nativo: "
        "cuando el significado depende de metadatos, historial o pista de auditoria, "
        "esos elementos se conservan y pueden revisarse."
    )
    _table(
        doc,
        ["Principio", "Pregunta de control", "Implementacion en FamSPI", "Evidencia"],
        ALCOA_ROWS,
        widths=(2.2, 4.0, 6.0, 4.3),
        header_fill=GREEN,
        font_size=7,
    )
    doc.add_heading("Evaluacion de integridad de datos (DIRA)", level=2)
    doc.add_paragraph(
        "Evaluacion de riesgo de integridad para datos criticos representativos "
        "(WHO guia §16.4). El riesgo residual se reduce con auditoria, control de "
        "acceso por rol y respaldo; los modulos sin prueba de verificacion directa mantienen riesgo "
        "medio hasta priorizar su cobertura."
    )
    _table(
        doc,
        ["Dato critico", "Origen", "Captura/proceso", "Control", "Riesgo", "Residual"],
        DIRA_ROWS,
        widths=(3.2, 2.8, 3.4, 3.4, 2.4, 1.3),
        header_fill=NAVY,
        font_size=7,
    )


def _add_config_baseline(doc: Document) -> None:
    """Linea base de configuracion as-built/as-configured (WHO §15)."""
    doc.add_heading("Linea base de configuracion as-built (WHO §15)", level=1)
    doc.add_paragraph(
        "Parametros de configuracion vigentes obtenidos del sistema actual. Las "
        "diferencias frente a documentos historicos se registran como brecha. Los "
        "elementos configurables quedan sujetos a control de cambios con ticket TI."
    )
    _table(
        doc,
        ["Elemento", "Valor aprobado", "Fuente", "Ambiente", "Evidencia"],
        [
            ("Zona horaria", "America/Guayaquil", "Estandar corporativo", "Produccion", "Config de servidor / consultas NOW()"),
            ("Auditoria GxP", "Habilitada (no desactivable por usuario rutinario)", "Requisito de trazabilidad", "Produccion", "`auditMiddleware`, `auditoria.logs`"),
            ("Autenticacion", "JWT firmado, iss `spi-fam-backend` / aud `spi-fam-frontend`", "Diseno de seguridad", "Produccion", "`config/oauth.js`, `middlewares/auth.js`"),
            ("Sesion", "Bearer + x-refresh-token en headers; sin cookies", "Diseno de seguridad", "Produccion", "`middlewares/auth.js`, `session.repository.js`"),
            ("Control de acceso", "RBAC por rol y grupos (ROLE_GROUPS); guard por modulo", "Matriz de roles", "Produccion", "`middlewares/roles.js`, `moduleAccess.js`"),
            ("Custodia de secretos", "gcloud Secret Manager (DB_PASSWORD, proyecto famspi-sbox)", "Politica de secretos", "Produccion", "`config/db.js`; nunca .env en produccion"),
            ("Respaldo", "Job programado a almacenamiento externo (Drive), endpoint protegido", "Politica de respaldo", "Produccion", "`jobs/databaseBackupToDrive.js`, `jobsAuth.js`"),
            ("Migraciones", "Archivos SQL numerados, versionados y aplicados de forma controlada", "Control de cambios BD", "Produccion", "`backend/migrations/`"),
        ],
        widths=(2.8, 5.2, 3.2, 2.0, 3.3),
        header_fill=BLUE,
        font_size=7,
    )


GLOSSARY_ROWS = (
    ("DQ / IQ / OQ / PQ", "Calificacion de Diseno / Instalacion / Operacion / Desempeno"),
    ("URS / FRS / DS", "Especificacion de Requisitos de Usuario / Funcional / de Diseno"),
    ("RTM", "Matriz de Trazabilidad de Requisitos"),
    ("UAT", "Prueba de Aceptacion de Usuario"),
    ("RBAC", "Control de Acceso Basado en Roles"),
    ("JWT", "JSON Web Token (autenticacion sin sesion persistente)"),
    ("ALCOA+", "Atribuible, Legible, Contemporaneo, Original, Exacto (+ Completo, Consistente, Duradero, Disponible)"),
    ("DIRA", "Evaluacion de Riesgo de Integridad de Datos"),
    ("GxP / GMP", "Buenas Practicas reguladas / Buenas Practicas de Manufactura"),
    ("CAPA", "Accion Correctiva y Preventiva"),
    ("SOP", "Procedimiento Operativo Estandar"),
    ("TICS", "Departamento de Tecnologias de la Informacion y Comunicacion"),
    ("WHO / OMS", "Organizacion Mundial de la Salud (TRS 1019 Annex 3; TRS 1033 Annex 4)"),
)


def _add_glossary(doc: Document) -> None:
    """Definiciones y abreviaturas del expediente (WHO §8)."""
    doc.add_heading("Definiciones", level=1)
    _table(
        doc,
        ["Termino / abreviatura", "Definicion"],
        GLOSSARY_ROWS,
        widths=(3.5, 13.0),
        header_fill=BLUE,
        font_size=8,
    )


def _clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = text.replace("`", "")
    return text.strip()


def _parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines: list[str] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        table_lines.append(lines[index].strip())
        index += 1

    parsed: list[list[str]] = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if all(set(cell) <= {"-", ":", " "} for cell in cells):
            continue
        parsed.append([_clean_inline(cell) for cell in cells])
    return parsed, index


def _add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    headers = rows[0]
    body = rows[1:] or [["" for _ in headers]]
    width = max(1.0, 16.5 / max(1, len(headers)))
    _table(doc, headers, body, widths=[width] * len(headers), header_fill=BLUE, font_size=7)


def _add_markdown(doc: Document, source_path: Path) -> None:
    lines = source_path.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()

        if not line:
            index += 1
            continue

        if line.startswith("```"):
            in_code = not in_code
            index += 1
            continue

        if line.startswith("<!--") and line.endswith("-->"):
            # Marcadores internos de regeneracion (control de versiones del
            # equipo de TICS): nunca deben aparecer como texto visible.
            index += 1
            continue

        if in_code:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(raw)
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED)
            index += 1
            continue

        if line.startswith("|"):
            parsed, index = _parse_table(lines, index)
            _add_markdown_table(doc, parsed)
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            text = _clean_inline(heading.group(2))
            if level == 1:
                doc.add_heading(text, level=1)
            elif level == 2:
                doc.add_heading(text, level=2)
            else:
                doc.add_heading(text, level=3)
            index += 1
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", line)
        if bullet:
            doc.add_paragraph(_clean_inline(bullet.group(1)), style="List Bullet")
            index += 1
            continue

        number = re.match(r"^\d+\.\s+(.*)$", line)
        if number:
            doc.add_paragraph(_clean_inline(number.group(1)), style="List Number")
            index += 1
            continue

        paragraph = doc.add_paragraph(_clean_inline(line))
        if line.startswith(">"):
            paragraph.style = doc.styles["Intense Quote"]
        index += 1


def _add_test_case_detail(doc: Document, row: EvidenceRow, run_at: str) -> None:
    """Bloque por caso de verificacion: explicacion + codigo real del test +
    prueba (resultado real y recuadro para la captura de la ejecucion)."""
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(1)
    tr = title.add_run(f"{row.req_id} — {row.description}")
    tr.bold = True
    tr.font.size = Pt(9)
    tr.font.name = "Arial"
    tr.font.color.rgb = RGBColor.from_string(NAVY)

    exp = doc.add_paragraph()
    el = exp.add_run("Explicacion: ")
    el.bold = True
    el.font.color.rgb = RGBColor.from_string(NAVY)
    exp.add_run(
        f"Prueba de verificacion del modulo {row.module} que verifica el comportamiento "
        "descrito. Un resultado Conforme evidencia que la funcion se mantiene bajo "
        "control en la version evaluada."
    )

    if row.test_line:
        frag, start_line, _ = _extract_code(row.test_file, at_line=row.test_line, before=1, after=8)
    else:
        frag, start_line, _ = _extract_code(row.test_file, after=8)
    _add_code_fragment(doc, row.test_file, frag, start_line)

    res = doc.add_paragraph()
    rl = res.add_run("Prueba (resultado real): ")
    rl.bold = True
    rl.font.color.rgb = RGBColor.from_string(TEXT)
    dur = f" — duracion {row.duration_ms} ms" if row.duration_ms is not None else ""
    status_run = res.add_run(f"{STATUS_LABEL.get(row.status, row.status)}")
    status_run.bold = True
    res.add_run(f" (corrida {run_at}){dur}")

    _add_screenshot_placeholder(doc)


def _add_test_evidence_section(doc: Document, evidence: TestEvidence, detailed: bool = False) -> None:
    doc.add_heading("Evidencia de ejecucion de pruebas", level=1)
    doc.add_paragraph(
        "El Departamento de Tecnologias de la Informacion y Comunicacion (TICS) ejecuto "
        "la suite de pruebas de verificacion de regresion del sistema y consolido a "
        "continuacion los resultados obtenidos, como evidencia objetiva de esta "
        "calificacion conforme a WHO TRS 1019 Annex 3 Appendix 5, §3.7."
    )
    doc.add_paragraph(
        f"Fecha de ejecucion: {evidence.run_at}. Resultado global: "
        f"{'CONFORME' if evidence.success else 'CON DESVIACIONES'} — "
        f"{evidence.passed}/{evidence.total} pruebas conformes, {evidence.failed} con desviacion."
    )

    if evidence.rows:
        _table(
            doc,
            ["Requisito", "Modulo", "Descripcion", "Archivo de test", "Resultado", "Duracion (ms)"],
            [
                (
                    r.req_id,
                    r.module,
                    r.description,
                    Path(r.test_file).name,
                    STATUS_LABEL.get(r.status, r.status),
                    r.duration_ms if r.duration_ms is not None else "",
                )
                for r in evidence.rows
            ],
            widths=(2.2, 1.8, 5.3, 3.0, 1.9, 1.8),
            header_fill=GREEN,
            font_size=7,
        )
    else:
        doc.add_paragraph("No se registran requisitos con prueba de verificacion asociada.")

    if detailed and evidence.rows:
        doc.add_heading("Detalle por caso (explicacion, codigo y prueba)", level=2)
        doc.add_paragraph(
            "Cada caso de verificacion se documenta con su explicacion, el fragmento real "
            "del codigo de prueba y su resultado. El recuadro permite adjuntar la captura "
            "de la ejecucion (por ejemplo, la salida de la corrida que respalda el caso)."
        )
        for row in evidence.rows:
            _add_test_case_detail(doc, row, evidence.run_at)

    doc.add_heading("Discrepancias detectadas", level=2)
    failed_rows = [r for r in evidence.rows if r.status == "failed"]
    if not failed_rows:
        doc.add_paragraph("Ninguna. Todas las pruebas ejecutadas resultaron conformes.")
    else:
        _table(
            doc,
            ["Requisito", "Modulo", "Archivo de test", "Investigacion / causa raiz", "Decision"],
            [(r.req_id, r.module, Path(r.test_file).name, "", "") for r in failed_rows],
            widths=(2.2, 1.8, 3.3, 5.5, 3.4),
            header_fill=SOFT_RED,
            font_size=7,
        )
        doc.add_paragraph(
            "Por WHO §3.7, toda discrepancia critica/mayor debe investigarse y justificarse "
            "antes de aceptar el resultado. Completar causa raiz y decision antes de aprobar."
        )

    doc.add_heading("Modulos con evidencia por historial de control y mantenimiento", level=2)
    if evidence.gaps:
        doc.add_paragraph(
            "Para estos modulos, TICS sustenta el estado de control mediante evidencia de "
            "mantenimiento activo (historial de cambios) como evidencia de control, "
            "conforme al enfoque de validacion retrospectiva de WHO TRS 1019 Annex 3 "
            "Appendix 5, §12.8-12.9."
        )
        _table(
            doc,
            ["Modulo", "Riesgo", "Ultimo cambio registrado", "Cambios (12 meses)", "Justificacion de TICS"],
            [
                (g.module, g.risk, g.last_commit or "sin datos", g.commits_last_12_months, g.justification)
                for g in evidence.gaps
            ],
            widths=(2.2, 1.5, 2.6, 1.5, 8.7),
            header_fill=NAVY,
            font_size=7,
        )
    else:
        doc.add_paragraph("Ninguno: todos los modulos cuentan con evidencia de verificacion.")


def _add_conclusion_block(doc: Document, evidence: TestEvidence) -> None:
    doc.add_heading("Conclusion y decision de liberacion", level=1)
    estado = "CONFORME" if evidence.success else "CON DESVIACIONES ABIERTAS"
    doc.add_paragraph(
        f"Estado de la calificacion: {estado}. {evidence.passed}/{evidence.total} pruebas "
        f"de verificacion resultaron conformes; {evidence.failed} presentaron desviacion; "
        f"{len(evidence.gaps)} modulos permanecen bajo brecha de cobertura documentada."
    )
    doc.add_paragraph(
        "Este documento se libera para uso GMP/operativo solo si las discrepancias listadas "
        "arriba fueron investigadas y su aceptacion quedo justificada por escrito, y si las "
        "brechas declaradas fueron revisadas por el responsable de calidad/validacion."
    )


def _add_approval_block(doc: Document) -> None:
    doc.add_heading("Aprobaciones", level=1)
    _table(
        doc,
        ["Rol", "Nombre", "Firma", "Fecha", "Decision"],
        [
            ("Responsable del proceso", "", "", "", ""),
            ("Responsable Tecnico", "", "", "", ""),
            ("Gerencia General", "", "", "", ""),
        ],
        widths=(3.8, 3.8, 3.2, 2.8, 2.9),
        header_fill=NAVY,
    )


# ── Evidencia tangible por item de checklist (explicacion + codigo real +
#    recuadro para la captura de ejecucion). El codigo se extrae del archivo
#    fuente real; nunca se fabrica. La captura la pega el ejecutor. ─────────

def _read_json_field(rel_path: str, *keys: str) -> str | None:
    try:
        data = json.loads((PROJECT_ROOT / rel_path).read_text(encoding="utf-8"))
        for key in keys:
            data = data[key]
        return str(data)
    except Exception:
        return None


def _git_head_short() -> str | None:
    try:
        out = subprocess.run(
            ["git", "rev-parse", "--short", "HEAD"],
            cwd=PROJECT_ROOT, capture_output=True, text=True, timeout=10,
        )
        return out.stdout.strip() or None
    except Exception:
        return None


def _count_files(rel_dir: str, pattern: str) -> str | None:
    try:
        return str(len(list((PROJECT_ROOT / rel_dir).glob(pattern))))
    except Exception:
        return None


BACKEND_VERSION = _read_json_field("backend/package.json", "version")
FRONTEND_VERSION = _read_json_field("spi_front/package.json", "version")
MIGRATIONS_COUNT = _count_files("backend/migrations", "*.sql")
GIT_HEAD = _git_head_short()
MODULE_COUNT = _read_json_field("docs/validation/traceability_map.json", "module_count")
MODULES_WITHOUT_TESTS = _read_json_field("docs/validation/traceability_map.json", "modules_without_tests")


def _extract_code(rel_path: str, anchor: str | None = None, before: int = 0,
                  after: int = 9, max_chars: int = 1100,
                  at_line: int | None = None) -> tuple[str, int, bool]:
    """Devuelve (fragmento, linea_inicial, anclado). Extrae codigo REAL del
    archivo. Si el archivo o el anclaje no existen, lo indica sin inventar."""
    path = PROJECT_ROOT / rel_path
    if not path.exists():
        return (f"[Archivo no encontrado al generar: {rel_path}. Verificar manualmente.]", 0, False)
    lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    idx, anchored = 0, True
    if at_line:
        idx = max(0, min(at_line - 1, len(lines) - 1))
    elif anchor:
        found = next((i for i, ln in enumerate(lines) if anchor in ln), None)
        if found is None:
            idx, anchored = 0, False
        else:
            idx = found
    start = max(0, idx - before)
    end = min(len(lines), idx + after + 1)
    frag = "\n".join(lines[start:end]).rstrip()
    if len(frag) > max_chars:
        frag = frag[:max_chars].rstrip() + "\n..."
    return (frag, start + 1, anchored)


def _add_code_fragment(doc: Document, rel_path: str, frag: str, start_line: int) -> None:
    caption = doc.add_paragraph()
    cap_run = caption.add_run(f"Fragmento de codigo — {rel_path} (desde linea {start_line})")
    cap_run.bold = True
    cap_run.font.size = Pt(7.5)
    cap_run.font.name = "Arial"
    cap_run.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_after = Pt(2)

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    _shade(cell, LIGHT)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    for i, line in enumerate(frag.splitlines() or [""]):
        run = para.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor.from_string(TEXT)
        if i < len(frag.splitlines()) - 1:
            run.add_break()
    _set_widths(table, [16.8])
    doc.add_paragraph()


def _add_screenshot_placeholder(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    _shade(cell, SOFT_AMBER)
    head = cell.paragraphs[0]
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = head.add_run("Evidencia visual de la ejecucion — insertar captura de pantalla aqui")
    hr.bold = True
    hr.font.size = Pt(8.5)
    hr.font.name = "Arial"
    hr.font.color.rgb = RGBColor.from_string(MUTED)
    for _ in range(3):
        cell.add_paragraph()
    legend = cell.add_paragraph()
    legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = legend.add_run(
        "Ambiente: ____________    Version/commit: ____________    "
        "Usuario: ____________    Fecha/hora: ____________"
    )
    lr.font.size = Pt(7.5)
    lr.font.name = "Arial"
    lr.font.color.rgb = RGBColor.from_string(MUTED)
    _set_widths(table, [16.8])
    doc.add_paragraph()


def _add_item_evidence(doc: Document, item: dict) -> None:
    doc.add_heading(f"{item['id']} — {item['title']}", level=2)

    explanation = doc.add_paragraph()
    lbl = explanation.add_run("Explicacion: ")
    lbl.bold = True
    lbl.font.color.rgb = RGBColor.from_string(NAVY)
    explanation.add_run(item["explanation"])

    verified = item.get("verified")
    if verified:
        vp = doc.add_paragraph()
        vlbl = vp.add_run("Dato verificado al generar: ")
        vlbl.bold = True
        vlbl.font.color.rgb = RGBColor.from_string(GREEN)
        vp.add_run(verified)

    if item.get("file"):
        frag, start_line, anchored = _extract_code(
            item["file"], item.get("anchor"),
            before=item.get("before", 0), after=item.get("after", 9),
        )
        if not anchored and "no encontrado" not in frag:
            note = doc.add_paragraph()
            nr = note.add_run(
                "Nota: el anclaje exacto cambio; se muestra el inicio del archivo. "
                "Verificar la seccion vigente."
            )
            nr.italic = True
            nr.font.size = Pt(7.5)
            nr.font.color.rgb = RGBColor.from_string(MUTED)
        _add_code_fragment(doc, item["file"], frag, start_line)

    _add_screenshot_placeholder(doc)


def _add_item_evidence_section(doc: Document, header: str, intro: str, items: list[dict]) -> None:
    doc.add_heading(header, level=1)
    doc.add_paragraph(intro)
    for item in items:
        _add_item_evidence(doc, item)


ITEM_EVIDENCE: dict[str, tuple[str, str, list[dict]]] = {
    "FAM-SPI-VAL-DQ-2026": (
        "Evidencia por item — Calificacion de diseno (DQ)",
        "Cada verificacion de diseno se sustenta con el archivo fuente real que la "
        "respalda y con espacio para la captura de la revision ejecutada.",
        [
            {"id": "DQ-01", "title": "Inventario de rutas backend actualizado",
             "explanation": "El montaje real de rutas privadas del backend confirma que el inventario de endpoints del diseno corresponde al codigo desplegado.",
             "verified": (f"{MODULE_COUNT} modulos backend inventariados." if MODULE_COUNT else None),
             "file": "backend/src/routes/registerRoutes.js", "anchor": "function mountPrivateRoutes", "after": 12},
            {"id": "DQ-02", "title": "Inventario de pantallas privadas actualizado",
             "explanation": "Las rutas de interfaz protegidas confirman que el catalogo de pantallas del diseno esta alineado con la navegacion real.",
             "file": "spi_front/src/routes/AppRoutes.jsx", "anchor": "ProtectedRoute", "after": 10},
            {"id": "DQ-03", "title": "Matriz de roles actualizada",
             "explanation": "Los grupos de rol (ROLE_GROUPS) que expande el control de acceso demuestran que la matriz de roles del diseno es la que aplica el sistema.",
             "file": "backend/src/middlewares/roles.js", "anchor": "const ROLE_GROUPS", "after": 12},
            {"id": "DQ-04", "title": "Catalogo de modulos y workspaces vigente",
             "explanation": "El registro central de rutas monta cada modulo del catalogo; es la fuente de verdad del alcance funcional.",
             "verified": (f"{MODULE_COUNT} modulos; commit {GIT_HEAD}." if MODULE_COUNT and GIT_HEAD else None),
             "file": "backend/src/routes/registerRoutes.js", "anchor": "function mountPrivateRoutes", "after": 14},
            {"id": "DQ-05", "title": "URS por dominio sin contradicciones criticas",
             "explanation": "La URS reconstruida por dominio describe los requisitos de usuario vigentes contra los que se evalua el sistema.",
             "file": "docs/validation/URS/URS_modulo_autenticacion_sesiones.md", "after": 12},
            {"id": "DQ-06", "title": "FRS por dominio sin contradicciones criticas",
             "explanation": "La especificacion funcional por dominio detalla como el sistema satisface cada requisito de usuario.",
             "file": "docs/validation/FRS/FRS_modulo_autenticacion_sesiones.md", "after": 12},
            {"id": "DQ-07", "title": "DS por dominio alineado con implementacion",
             "explanation": "La especificacion de diseno por dominio documenta la solucion tecnica y su alineacion con la implementacion.",
             "file": "docs/validation/DS/DS_modulo_autenticacion_sesiones.md", "after": 12},
            {"id": "DQ-08", "title": "RTM refleja modulos y flujos vigentes",
             "explanation": "La matriz de trazabilidad enlaza requisito con evidencia de prueba o brecha declarada; se regenera desde el codigo real.",
             "file": "docs/validation/RTM/RTM_sistema_spi.md", "after": 12},
            {"id": "DQ-09", "title": "Navegacion, layout y patrones alineados al diseno",
             "explanation": "El documento de diseno visual define los patrones de navegacion y layout que la interfaz debe respetar.",
             "file": "DESIGN.md", "after": 12},
            {"id": "DQ-10", "title": "Areas validadas y pendientes separadas",
             "explanation": "El indice de validacion distingue el alcance vigente de la documentacion historica, evitando gobernar flujos ya cambiados.",
             "file": "docs/validation/README.md", "anchor": "Alcance funcional vigente", "after": 12},
        ],
    ),
    "FAM-SPI-VAL-IQ-2026": (
        "Evidencia por item — Calificacion de instalacion (IQ)",
        "Cada verificacion de instalacion se sustenta con el archivo fuente/configuracion "
        "real y con espacio para la captura de la ejecucion (terminal, panel o consulta).",
        [
            {"id": "IQ-01", "title": "Version del backend identificada",
             "explanation": "La version del artefacto backend bajo validacion se declara en su package.json; identifica de forma inequivoca lo desplegado.",
             "verified": (f"backend version {BACKEND_VERSION}" + (f", commit {GIT_HEAD}" if GIT_HEAD else "") if BACKEND_VERSION else None),
             "file": "backend/package.json", "anchor": "\"version\"", "after": 4},
            {"id": "IQ-02", "title": "Version del frontend identificada",
             "explanation": "La version del artefacto frontend se declara en su package.json, identificando la build de interfaz validada.",
             "verified": (f"frontend version {FRONTEND_VERSION}" if FRONTEND_VERSION else None),
             "file": "spi_front/package.json", "anchor": "\"version\"", "after": 4},
            {"id": "IQ-03", "title": "Rutas backend montadas segun alcance",
             "explanation": "El registro central monta las rutas privadas tras el pipeline de seguridad (verifyToken, moduleAccessGuard, auditoria).",
             "file": "backend/src/routes/registerRoutes.js", "anchor": "function mountPrivateRoutes", "after": 14},
            {"id": "IQ-04", "title": "Rutas frontend privadas y publicas segun alcance",
             "explanation": "Las rutas de la SPA se protegen con ProtectedRoute segun rol; confirma que las pantallas del alcance responden desde la version correcta.",
             "file": "spi_front/src/routes/AppRoutes.jsx", "anchor": "ProtectedRoute", "after": 10},
            {"id": "IQ-05", "title": "Base de datos accesible y coherente",
             "explanation": "El pool de conexiones a PostgreSQL (Neon) se configura con validacion y logs; es el punto de acceso unico y controlado a los datos.",
             "file": "backend/src/config/db.js", "anchor": "new Pool(", "after": 12},
            {"id": "IQ-06", "title": "Variables logicas y dependencias externas documentadas",
             "explanation": "La conexion se parametriza por variables de entorno y la credencial se obtiene de gcloud Secret Manager, nunca de un .env en produccion.",
             "file": "backend/src/config/db.js", "anchor": "process.env", "after": 8},
            {"id": "IQ-07", "title": "Integraciones activas identificadas",
             "explanation": "El flujo OAuth2 de Google autentica usuarios y emite el JWT con iss/aud propios; es una de las integraciones externas activas del sistema.",
             "file": "backend/src/config/oauth.js", "after": 12},
            {"id": "IQ-08", "title": "Usuarios y roles de prueba definidos",
             "explanation": "El control de acceso por rol y sus grupos define los perfiles con los que se ejecutan las pruebas y se opera cada modulo.",
             "file": "backend/src/middlewares/roles.js", "anchor": "const ROLE_GROUPS", "after": 12},
            {"id": "IQ-09", "title": "Trazabilidad tecnica disponible",
             "explanation": "El middleware de auditoria registra automaticamente toda operacion POST/PUT/PATCH/DELETE en auditoria.logs, garantizando trazabilidad tecnica.",
             "file": "backend/src/middlewares/auditMiddleware.js", "anchor": "async function auditMiddleware", "after": 12},
            {"id": "IQ-10", "title": "Mecanismo de recuperacion o rollback documentado",
             "explanation": "Los cambios de esquema se versionan como migraciones SQL numeradas aplicadas de forma controlada; el script de aplicacion documenta el procedimiento.",
             "verified": (f"{MIGRATIONS_COUNT} migraciones SQL versionadas." if MIGRATIONS_COUNT else None),
             "file": "backend/run_migrations.ps1", "after": 12},
        ],
    ),
    "FAM-SPI-VAL-OQ-2026": (
        "Evidencia por item — Casos criticos de operacion (OQ)",
        "Cada caso critico se sustenta con el control real que lo implementa en el codigo "
        "y con espacio para la captura de la ejecucion funcional. Complementa la evidencia "
        "de la suite de pruebas de verificacion incluida mas adelante.",
        [
            {"id": "OQ-001", "title": "Login correcto e incorrecto",
             "explanation": "La validacion del token de sesion determina el acceso; un login/credencial invalido no obtiene un JWT valido y es rechazado.",
             "file": "backend/src/middlewares/auth.js", "anchor": "const verifyToken", "after": 14},
            {"id": "OQ-002", "title": "Ruta privada sin token",
             "explanation": "Toda ruta privada pasa por verifyToken; sin Authorization Bearer valido la peticion se rechaza antes de llegar al handler.",
             "file": "backend/src/middlewares/auth.js", "anchor": "const verifyToken", "after": 16},
            {"id": "OQ-003", "title": "Ruta con rol no autorizado",
             "explanation": "requireRole expande grupos de rol y bloquea al usuario cuyo rol no pertenece al conjunto permitido para el endpoint.",
             "file": "backend/src/middlewares/roles.js", "anchor": "requireRole", "after": 14},
            {"id": "OQ-004", "title": "Cambio de estado con auditoria",
             "explanation": "Cualquier operacion de escritura queda registrada por el middleware de auditoria con usuario, accion y modulo, evidenciando el cambio de estado.",
             "file": "backend/src/middlewares/auditMiddleware.js", "anchor": "async function auditMiddleware", "after": 14},
            {"id": "OQ-005", "title": "Flujo de solicitud y aprobacion",
             "explanation": "El servicio de aprobaciones implementa la logica de revision/decision sobre solicitudes, base del flujo solicitud-aprobacion.",
             "file": "backend/src/modules/approvals/approvals.service.js", "anchor": "module.exports", "after": 16},
            {"id": "OQ-006", "title": "Asignacion o reordenamiento segun reglas del dominio",
             "explanation": "El motor del business case aplica las reglas de calculo/estado del dominio comercial sobre los registros criticos.",
             "file": "backend/src/modules/business-case/businessCase.service.js", "anchor": "module.exports", "after": 16},
            {"id": "OQ-007", "title": "Guardado de ficha/perfil y persistencia correcta",
             "explanation": "El servicio de perfil de usuario persiste los datos de la ficha y los recupera de forma consistente.",
             "file": "backend/src/modules/user-profile/userProfile.service.js", "anchor": "module.exports", "after": 16},
            {"id": "OQ-008", "title": "Flujo de documento o firma",
             "explanation": "El motor de firmas resuelve firmantes y estados del workflow documental multifirmante, con snapshot del firmante para trazabilidad.",
             "file": "backend/src/modules/signature-workflows/signatureWorkflows.service.js", "anchor": "async function resolveSignerSnapshot", "after": 14},
            {"id": "OQ-009", "title": "Carga o visualizacion de expediente",
             "explanation": "El servicio de documentos gestiona el almacenamiento y recuperacion de los archivos que componen el expediente.",
             "file": "backend/src/modules/documents/document.service.js", "anchor": "module.exports", "after": 16},
            {"id": "OQ-010", "title": "Validaciones operativas de formularios criticos",
             "explanation": "La normalizacion de payloads de API sanea y valida la entrada antes de que llegue a la logica de negocio.",
             "file": "backend/src/middlewares/apiNormalization.js", "anchor": "module.exports", "after": 16},
        ],
    ),
    "FAM-SPI-VAL-PQ-2026": (
        "Evidencia por item — Escenarios de desempeno por dominio (PQ)",
        "Cada escenario de negocio de punta a punta se sustenta con el servicio de dominio "
        "real que lo soporta y con espacio para la captura del flujo ejecutado con usuarios "
        "reales.",
        [
            {"id": "PQ-TH", "title": "Talento humano — solicitud, revision, aprobacion y consulta",
             "explanation": "El servicio de permisos soporta el ciclo completo de solicitud, revision y aprobacion del dominio de talento humano.",
             "file": "backend/src/modules/permisos/permisos.service.js", "anchor": "module.exports", "after": 16},
            {"id": "PQ-COM", "title": "Comercial — cliente/oportunidad de registro a seguimiento",
             "explanation": "El servicio de clientes soporta el registro y seguimiento operativo del dominio comercial.",
             "file": "backend/src/modules/clients/clients.service.js", "anchor": "module.exports", "after": 16},
            {"id": "PQ-BC", "title": "Business Case — creacion, analisis, aprobacion y trazabilidad",
             "explanation": "El motor del business case sostiene el ciclo de creacion, calculo, aprobacion y consulta con trazabilidad.",
             "file": "backend/src/modules/business-case/businessCase.service.js", "anchor": "module.exports", "after": 16},
            {"id": "PQ-ST", "title": "Servicio tecnico — solicitud, planificacion, ejecucion y cierre",
             "explanation": "El servicio de mantenimientos soporta la planificacion y cierre de la operacion de servicio tecnico.",
             "file": "backend/src/modules/mantenimientos/mantenimientos.service.js", "anchor": "module.exports", "after": 16},
            {"id": "PQ-CMP", "title": "Compras — solicitud, revision, aprobacion y expediente",
             "explanation": "El servicio de compras privadas soporta el flujo de solicitud a expediente de compra.",
             "file": "backend/src/modules/private-purchases/privatePurchases.service.js", "anchor": "module.exports", "after": 16},
            {"id": "PQ-TI", "title": "Inventario y TI — alta/asignacion/entrega o ticket completo",
             "explanation": "El servicio de activos TI soporta el alta, asignacion y entrega de equipos del dominio TI.",
             "file": "backend/src/modules/ti-assets/tiAssets.service.js", "anchor": "module.exports", "after": 16},
            {"id": "PQ-FIN", "title": "Finanzas — viatico con revision y resultado final",
             "explanation": "El servicio de viaticos soporta el flujo financiero de solicitud, revision y resultado del viatico.",
             "file": "backend/src/modules/viaticos/viaticos.service.js", "anchor": "module.exports", "after": 16},
            {"id": "PQ-FIR", "title": "Firma y documentos — firma completa y verificacion posterior",
             "explanation": "El motor de firmas sostiene la firma electronica del documento y su verificacion posterior.",
             "file": "backend/src/modules/signature-workflows/signatureWorkflows.service.js", "anchor": "async function resolveRecipientOrThrow", "after": 14},
        ],
    ),
}


# ── Correccion ortografica de tildes (post-proceso) ───────────────────────
# Los textos se redactan sin acentos por compatibilidad; este paso los
# restituye en la salida .docx. Los fragmentos de codigo (fuente Consolas)
# se excluyen para no alterar el codigo real.

ACCENT_MAP = {
    # esdrujulas y otras
    "modulo": "módulo", "modulos": "módulos", "codigo": "código", "codigos": "códigos",
    "numero": "número", "numeros": "números", "parametro": "parámetro", "parametros": "parámetros",
    "analisis": "análisis", "tecnico": "técnico", "tecnica": "técnica", "tecnicos": "técnicos",
    "tecnicas": "técnicas", "practica": "práctica", "practicas": "prácticas", "politica": "política",
    "politicas": "políticas", "historico": "histórico", "historica": "histórica",
    "historicos": "históricos", "historicas": "históricas", "automatico": "automático",
    "automatica": "automática", "estatico": "estático", "estatica": "estática",
    "periodico": "periódico", "periodica": "periódica", "publico": "público", "publica": "pública",
    "publicos": "públicos", "publicas": "públicas", "unico": "único", "unica": "única",
    "unicos": "únicos", "unicas": "únicas", "minimo": "mínimo", "minima": "mínima",
    "maximo": "máximo", "maxima": "máxima", "rapido": "rápido", "generico": "genérico",
    "especifico": "específico", "especifica": "específica", "logico": "lógico", "fisico": "físico",
    "electronico": "electrónico", "electronica": "electrónica", "cronologico": "cronológico",
    "termino": "término", "terminos": "términos", "ultimo": "último", "ultima": "última",
    "ultimos": "últimos", "ultimas": "últimas", "proposito": "propósito", "metrica": "métrica",
    "metricas": "métricas", "informatico": "informático", "trafico": "tráfico",
    # -ia acentuada
    "auditoria": "auditoría", "categoria": "categoría", "categorias": "categorías",
    "tecnologia": "tecnología", "tecnologias": "tecnologías", "garantia": "garantía",
    "dia": "día", "dias": "días", "guia": "guía", "guias": "guías", "metodologia": "metodología",
    "energia": "energía", "jerarquia": "jerarquía", "bateria": "batería", "compania": "compañía",
    "companias": "compañías",
    # con ñ
    "diseno": "diseño", "disenos": "diseños", "desempeno": "desempeño", "espanol": "español",
    "senal": "señal", "senales": "señales", "pequeno": "pequeño", "ano": "año", "anos": "años",
    # agudas (verbos/adverbios)
    "tambien": "también", "segun": "según", "despues": "después", "ademas": "además",
    "asi": "así", "aqui": "aquí", "alli": "allí", "mas": "más", "demas": "demás",
    "estan": "están", "estara": "estará", "estaran": "estarán", "sera": "será", "seran": "serán",
    "podra": "podrá", "podran": "podrán", "debera": "deberá", "deberan": "deberán",
    "atras": "atrás", "quiza": "quizá",
    # nombres/adjetivos varios
    "pais": "país", "paises": "países", "linea": "línea", "lineas": "líneas", "area": "área",
    "areas": "áreas", "facil": "fácil", "dificil": "difícil", "util": "útil", "comun": "común",
    "limite": "límite", "limites": "límites", "razon": "razón", "telefono": "teléfono",
    # -sion / -tion españolas (la regla automatica solo cubre -cion/-xion para
    # no dañar palabras en ingles como "registration" o "session")
    "version": "versión", "revision": "revisión", "decision": "decisión", "sesion": "sesión",
    "dimension": "dimensión", "extension": "extensión", "gestion": "gestión",
    "cuestion": "cuestión", "presion": "presión", "provision": "provisión",
    "division": "división", "precision": "precisión", "prevision": "previsión",
    "comprension": "comprensión", "impresion": "impresión", "expresion": "expresión",
    "tension": "tensión", "inclusion": "inclusión", "exclusion": "exclusión",
    "conclusion": "conclusión", "emision": "emisión", "omision": "omisión",
    "supervision": "supervisión", "conexion": "conexión",
}

_WORD_RE = re.compile(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+")
_ION_SUFFIXES = (("cion", "ción"), ("xion", "xión"))


def _accentuate(text: str) -> str:
    def _apply_case(base: str, sample: str) -> str:
        if sample.isupper():
            return base.upper()
        if sample[:1].isupper():
            return base[:1].upper() + base[1:]
        return base

    def _fix(match: re.Match) -> str:
        word = match.group(0)
        low = word.lower()
        rep = ACCENT_MAP.get(low)
        if rep is None:
            for suf, acc in _ION_SUFFIXES:
                if len(low) >= 5 and low.endswith(suf):
                    rep = low[:-4] + acc
                    break
        if rep is None:
            return word
        return _apply_case(rep, word)

    return _WORD_RE.sub(_fix, text)


def _iter_doc_paragraphs(doc: Document):
    def _from_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    yield from doc.paragraphs
    yield from _from_tables(doc.tables)
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from _from_tables(section.header.tables)
        yield from section.footer.paragraphs
        yield from _from_tables(section.footer.tables)


def _apply_accents(doc: Document) -> None:
    for paragraph in _iter_doc_paragraphs(doc):
        for run in paragraph.runs:
            if run.font.name == "Consolas":
                continue
            fixed = _accentuate(run.text)
            if fixed != run.text:
                run.text = fixed


def _enable_update_fields(doc: Document) -> None:
    """Marca los campos (indice/TOC) para que Word los actualice al abrir."""
    try:
        settings = doc.settings.element
        existing = settings.find(qn("w:updateFields"))
        if existing is None:
            existing = OxmlElement("w:updateFields")
            settings.append(existing)
        existing.set(qn("w:val"), "true")
    except Exception:
        pass


def build_phase_document(phase: PhaseDocument, evidence: TestEvidence | None) -> Path:
    source_path = VALIDATION_DIR / phase.source
    if not source_path.exists():
        raise FileNotFoundError(f"No existe la fuente documental: {source_path}")

    output_path = VALIDATION_DIR / phase.output
    doc = Document()
    _style_document(doc, phase.title, phase.code)
    _add_cover(doc, phase)
    _add_revision_history(doc)
    _add_table_of_contents(doc)
    _add_glossary(doc)
    _add_retrospective_document_controls(doc)
    _add_who_alignment(doc, phase.who_focus)
    _add_retrospective_strategy(doc)
    _add_markdown(doc, source_path)
    if phase.code == "FAM-SPI-VAL-IQ-2026":
        _add_config_baseline(doc)
    if phase.code in ITEM_EVIDENCE:
        header, intro, items = ITEM_EVIDENCE[phase.code]
        _add_item_evidence_section(doc, header, intro, items)
    if phase.include_test_evidence:
        if evidence is None:
            raise RuntimeError(
                f"{phase.output} requiere evidencia real de pruebas pero no se cargo. "
                "Correr load_test_evidence() antes de generar este documento."
            )
        _add_test_evidence_section(doc, evidence, detailed=(phase.code == "FAM-SPI-VAL-OQ-2026"))
        _add_conclusion_block(doc, evidence)
    _add_approval_block(doc)
    _apply_accents(doc)
    _enable_update_fields(doc)
    doc.save(output_path)
    return output_path


def _add_source_inventory(doc: Document) -> None:
    doc.add_heading("Inventario de fuentes revisadas", level=1)
    sources = [
        ("Documentacion", "docs/validation/README.md", "Indice documental actualizado."),
        ("Documentacion", "docs/validation/01_primeros_pasos_y_dq.md", "Primeros pasos, DQ y brechas."),
        ("Documentacion", "docs/validation/02_implementacion_iq.md", "IQ e instalacion."),
        ("Documentacion", "docs/validation/03_operacion_oq.md", "OQ y pruebas operativas."),
        ("Documentacion", "docs/validation/04_desempeno_pq.md", "PQ y aceptacion."),
        ("Backend", "backend/src/routes/registerRoutes.js", "Modulos y rutas registradas."),
        ("Frontend", "spi_front/src/routes/AppRoutes.jsx", "Rutas de interfaz y navegacion."),
        ("Diseno", "DESIGN.md", "Criterios visuales y coherencia UI."),
        ("Tickets TI", "backend/src/modules/support-tickets/CONTEXT.md", "Flujo operativo de mesa de ayuda y workspace TI."),
        ("Tickets TI", "backend/src/modules/support-tickets/supportTickets.service.js", "Tablas support_tickets, support_ticket_events y transiciones."),
        ("Tickets TI", "docs/validation/URS/URS_modulo_ti_soporte_tickets.md", "Requerimientos de trazabilidad, estados y eventos de tickets."),
        ("Control de cambios", "docs/validation/general/14A_control_cambios.md", "Procedimiento de control de cambios con ticket TI obligatorio."),
    ]
    _table(doc, ["Tipo", "Ruta", "Uso en validacion"], sources, widths=(3.0, 6.8, 6.7), header_fill=BLUE)


def _add_gap_log_template(doc: Document, evidence: TestEvidence) -> None:
    doc.add_heading("Registro de brechas retrospectivas", level=1)
    doc.add_paragraph(
        "Brechas de cobertura de pruebas identificadas y evaluadas por TICS, con su "
        "riesgo declarado y decision de aceptacion."
    )
    if not evidence.gaps:
        doc.add_paragraph("Ninguna brecha: todos los modulos cuentan con evidencia de verificacion.")
        return
    _table(
        doc,
        ["ID", "Modulo", "Riesgo", "Ultimo cambio registrado", "Cambios (12 meses)", "Justificacion", "Estado"],
        [
            (
                f"GAP-{index:03d}",
                g.module,
                g.risk,
                g.last_commit or "sin datos",
                g.commits_last_12_months,
                g.justification,
                "Abierta" if g.risk == "alto" else "Aceptada",
            )
            for index, g in enumerate(evidence.gaps, start=1)
        ],
        widths=(1.5, 2.2, 1.5, 2.4, 1.5, 5.9, 1.5),
        header_fill=NAVY,
        font_size=7,
    )


def _add_traceability_template(doc: Document, evidence: TestEvidence) -> None:
    doc.add_heading("Matriz de trazabilidad", level=1)
    doc.add_paragraph(
        "Relacion entre cada requisito del sistema y su evidencia de prueba, elaborada "
        "por TICS a partir de la ejecucion de la suite de pruebas de verificacion."
    )
    if not evidence.rows:
        doc.add_paragraph("No se registran requisitos con prueba de verificacion asociada.")
        return
    _table(
        doc,
        ["Req.", "Modulo", "Evidencia de prueba", "Calificacion asociada", "Resultado"],
        [
            (r.req_id, r.module, Path(r.test_file).name, "OQ / PQ", STATUS_LABEL.get(r.status, r.status))
            for r in evidence.rows
        ],
        widths=(2.2, 1.8, 5.0, 1.5, 2.0),
        header_fill=GREEN,
        font_size=7,
    )


def _add_change_control_ticket_verification(doc: Document) -> None:
    doc.add_heading("Verificacion de control de cambios contra tickets TI", level=1)
    doc.add_paragraph(
        "El control de cambios se considera valido solo cuando cada cambio material "
        "esta asociado a un ticket TI real del modulo support-tickets. El ticket "
        "funciona como evidencia operacional primaria y el registro documental de "
        "validacion conserva la referencia, decision, impacto y pruebas."
    )
    _table(
        doc,
        ["Punto de control", "Fuente real", "Criterio de aceptacion"],
        [
            ("Codigo de ticket", "support_tickets.code", "Debe figurar en el registro de control de cambios."),
            ("Estado", "support_tickets.status", "Resuelto/cerrado para implementados; justificado si cancelado."),
            ("Eventos", "support_ticket_events", "Debe existir ciclo suficiente: creado, asignado, cambio de estado, comentarios y cierre/resolucion."),
            ("Comentarios", "support_ticket_comments", "Deben incluir analisis, pruebas, despliegue, rollback o decision de no ejecutar."),
            ("Workspace TI", "/api/v1/support-tickets/workspace/list", "TI debe poder verificar el ticket por codigo, estado y responsable."),
            ("Revalidacion", "DQ/IQ/OQ/PQ afectados", "El ticket debe enlazarse con la evidencia de prueba o desviacion correspondiente."),
        ],
        widths=(4.0, 5.2, 7.3),
        header_fill=BLUE,
    )
    _table(
        doc,
        ["ID Cambio", "Ticket TI", "Estado ticket", "Eventos verificados", "Pruebas / evidencia", "Decision", "Responsable", "Fecha"],
        [
            ("CC-001", "", "", "", "", "", "", ""),
            ("CC-002", "", "", "", "", "", "", ""),
            ("CC-003", "", "", "", "", "", "", ""),
        ],
        widths=(1.5, 1.8, 1.9, 2.3, 3.0, 2.0, 2.0, 1.5),
        header_fill=NAVY,
        font_size=7,
    )


def _add_operation_maintenance_section(doc: Document) -> None:
    doc.add_heading("Operacion y mantenimiento del sistema (WHO §13)", level=1)
    doc.add_paragraph(
        "TICS mantiene controles de seguridad, respaldo, migracion de datos y revision "
        "periodica sobre FamSPI conforme a WHO §13.1-13.16."
    )
    _table(
        doc,
        ["Control WHO §13", "Implementacion en FamSPI", "Evidencia"],
        [
            (
                "Seguridad y control de acceso (§13.1-13.8)",
                "Autenticacion JWT obligatoria y control de acceso por rol en cada endpoint privado.",
                "`middlewares/auth.js` (verifyToken), `middlewares/roles.js` (requireRole, ROLE_GROUPS)",
            ),
            (
                "Trazabilidad de operaciones (§13.6, §13.8)",
                "Toda operacion POST/PUT/PATCH/DELETE queda registrada automaticamente con usuario y accion.",
                "`auditMiddleware`, tabla `auditoria.logs`",
            ),
            (
                "Respaldo y recuperacion (§13.11)",
                "Job programado que genera respaldo comprimido de la base de datos y lo sube a almacenamiento externo, con endpoint protegido.",
                "`backend/src/jobs/databaseBackupToDrive.js`, `middlewares/jobsAuth.js` (JOBS_KEY), ejecucion via Cloud Scheduler",
            ),
            (
                "Migracion de datos (§13.13-13.14)",
                "Cambios de esquema versionados y aplicados de forma controlada.",
                f"{MIGRATIONS_COUNT} archivos SQL numerados en `backend/migrations/` al {TODAY}",
            ),
            (
                "Revision periodica (§13.15-13.16)",
                "Historial de cambios por modulo se revisa como parte del analisis de riesgo y brechas.",
                "Seccion 11 del documento Plan Maestro y DQ (evidencia de mantenimiento activo por modulo)",
            ),
        ],
        widths=(3.3, 6.7, 6.5),
        header_fill=BLUE,
        font_size=7,
    )


def _add_retirement_review(doc: Document) -> None:
    doc.add_heading("Retiro del sistema (WHO §14)", level=1)
    doc.add_paragraph(
        "WHO Appendix 5 §14 contempla el retiro del sistema, de componentes o de "
        "integraciones como parte del ciclo de vida. Para FamSPI, cualquier modulo "
        "reemplazado, deshabilitado o migrado debe mantener evidencia de datos "
        "retenidos, accesibilidad, integridad y decision aprobada."
    )
    doc.add_paragraph(
        "Caso verificable: la integracion con Odoo fue retirada del sistema por "
        "decision de negocio. El modulo backend/src/modules/integrations/ documenta "
        "la remocion en su AGENTS.md, y la prueba de verificacion "
        "integrationOutboxWorker.service.test.js confirma que los eventos no-CRM se "
        "marcan como omitidos por ausencia de proveedor, evidenciando que el retiro "
        "quedo reflejado en comportamiento verificable del sistema, no solo en "
        "documentacion."
    )
    _table(
        doc,
        ["Elemento", "Criterio minimo", "Evidencia"],
        [
            (
                "Retencion de datos",
                "Datos historicos recuperables durante el periodo definido.",
                "Respaldo periodico via `databaseBackupToDrive.js`",
            ),
            (
                "Integridad",
                "No perdida ni alteracion no autorizada durante migracion o archivo.",
                "Migraciones versionadas en `backend/migrations/`",
            ),
            (
                "Comportamiento verificable tras el retiro",
                "El sistema debe reflejar el retiro en su comportamiento, no solo en texto.",
                "`integrationOutboxWorker.service.test.js` (evento no-CRM se omite, Odoo ya no existe)",
            ),
            (
                "Aprobacion",
                "Decision documentada por responsables del proceso y TI.",
                "",
            ),
        ],
        widths=(4.0, 5.5, 7.0),
        header_fill=BLUE,
        font_size=7,
    )


def build_annexes(evidence: TestEvidence) -> Path:
    output_path = VALIDATION_DIR / "FAMSPI_05_ANEXOS_RETROSPECTIVA_2026.docx"
    doc = Document()
    _style_document(doc, "Anexos de validacion retrospectiva", "FAM-SPI-VAL-ANX-2026")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Soporte transversal del expediente de validacion (DQ/IQ/OQ/PQ)")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph()

    _add_cover_card(
        doc,
        [
            ("Codigo del documento", "FAM-SPI-VAL-ANX-2026"),
            ("Compania", COMPANY),
            ("Sistema", SYSTEM_NAME),
            ("Elaborado por", DEPARTMENT),
            ("Version documental", VERSION),
            ("Fecha de emision", TODAY),
            ("Norma base", WHO_REFERENCE),
            ("Norma de integridad de datos", WHO_DATA_INTEGRITY_REFERENCE),
            ("Clasificacion", "Confidencial - uso interno controlado"),
            ("Uso", "Soporte transversal para DQ/IQ/OQ/PQ."),
        ],
    )
    doc.add_page_break()
    _add_revision_history(doc)
    _add_table_of_contents(doc)
    _add_glossary(doc)

    _add_who_alignment(
        doc,
        (
            "Analisis de riesgo",
            "Analisis de brechas",
            "Revision de historia de uso",
            "Control de cambios",
            "Seguridad, respaldo y migracion de datos (§13)",
            "Revision periodica (§13.15-13.16)",
            "Retiro del sistema (§14)",
        ),
    )
    _add_retrospective_document_controls(doc)
    _add_source_inventory(doc)
    _add_retrospective_strategy(doc)
    _add_data_integrity_alcoa(doc)
    _add_change_control_ticket_verification(doc)
    _add_operation_maintenance_section(doc)
    _add_gap_log_template(doc, evidence)
    _add_traceability_template(doc, evidence)
    _add_retirement_review(doc)
    _add_approval_block(doc)
    _apply_accents(doc)
    _enable_update_fields(doc)
    doc.save(output_path)
    return output_path


def build_all() -> list[Path]:
    evidence = load_test_evidence()
    outputs = [
        build_phase_document(phase, evidence if phase.include_test_evidence else None)
        for phase in PHASE_DOCUMENTS
    ]
    outputs.append(build_annexes(evidence))
    for output in outputs:
        print(f"Generado: {output}")
    return outputs


if __name__ == "__main__":
    build_all()
