"""Protocolo Maestro de Validacion FamSPI v1.0.0 + Anexos A-E.
Alineado con WHO TRS 1019 Annex 3, Appendix 5 y metodologia GAMP 5 (V-Model, HLRA, FRA).
Genera un unico documento auditable con narrativa tecnica profunda.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = (
    r"c:\Users\Departamento de TI\Desktop\PROYECTOS\FamSPI"
    r"\docs\validation\PROTOCOLO_MAESTRO_VALIDACION_FAMSPI_V1_0_0.docx"
)
FECHA   = "18/06/2026"
VERSION = "1.0"
SISTEMA = "FamSPI v1.0.0"
DEPTO   = "Departamento de Tecnologia de la Informacion (TI)"
NORMA   = "WHO TRS 1019 Annex 3, Appendix 5"
GAMP    = "GAMP 5 — Second Edition (ISPE, 2022)"

AZUL   = RGBColor(0x1F, 0x49, 0x7D)
VERDE  = RGBColor(0x37, 0x86, 0x50)
GRIS   = RGBColor(0x80, 0x80, 0x80)
ROJO   = RGBColor(0xC0, 0x00, 0x00)


# ═══════════════════════════════════════════════════════════════════════════════
#  PRIMITIVAS
# ═══════════════════════════════════════════════════════════════════════════════

def _shd(cell, hex_color):
    tcp = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), hex_color)
    tcp.append(s)


def _cell(cell, text, bold=False, size=8, color=None, fill=None,
          center=False, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if fill:
        _shd(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _W(t, widths):
    for row in t.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def _estado_fill(v):
    if v in ("Conforme", "CONFORME"):   return "E2EFDA"
    if v in ("Alto", "ALTO"):           return "FCE4D6"
    if v in ("Medio", "MEDIO"):         return "FFF2CC"
    if v in ("Pendiente", "PENDIENTE", "—", ""): return "FAFAFA"
    return None


def _table(doc, headers, rows, widths=None, hbg="1F497D", font=8,
           last_estado=True, alt=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        _cell(t.rows[0].cells[i], h, bold=True,
              color="FFFFFF", fill=hbg, center=True, size=font)
    for ri, rd in enumerate(rows):
        row = t.add_row()
        base = "F2F2F2" if (alt and ri % 2 == 0) else "FFFFFF"
        last = len(rd) - 1
        for ci, v in enumerate(rd):
            sf = _estado_fill(v) if (last_estado and ci == last) else None
            bold_v = bool(sf and v in ("Conforme", "CONFORME", "Alto", "ALTO"))
            _cell(row.cells[ci], v, size=font, fill=sf or base,
                  bold=bold_v, center=(sf is not None and ci == last))
    if widths:
        _W(t, widths)
    doc.add_paragraph()
    return t


def _kv(doc, rows, widths=(5, 11)):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        _cell(t.rows[i].cells[0], k, bold=True,
              fill="1F497D", color="FFFFFF", size=9)
        _cell(t.rows[i].cells[1], v, size=9, fill="F8F8F8")
    _W(t, widths)
    doc.add_paragraph()


def _h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.color.rgb = AZUL if level <= 1 else (VERDE if level == 2 else GRIS)
    return p


def _p(doc, text, size=10, after=5, italic=False, bold=False, color=None):
    para = doc.add_paragraph()
    para.style = doc.styles["Normal"]
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = para.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(size)
    r.italic = italic; r.bold = bold
    if color:
        r.font.color.rgb = color
    return para


def _bullet(doc, items, size=10):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(3)
        r = para.add_run(item)
        r.font.name = "Calibri"; r.font.size = Pt(size)


def _numbered(doc, items, size=10):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(3)
        r = para.add_run(item)
        r.font.name = "Calibri"; r.font.size = Pt(size)


def _label(doc, text, fill="1F497D"):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(4)
    pPr = para._p.get_or_add_pPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), fill)
    pPr.append(s)
    r = para.add_run(text)
    r.bold = True; r.font.size = Pt(9)
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _code(doc, lines):
    for line in lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.left_indent  = Cm(0.3)
        pPr = para._p.get_or_add_pPr()
        s = OxmlElement("w:shd")
        s.set(qn("w:val"), "clear")
        s.set(qn("w:color"), "auto")
        s.set(qn("w:fill"), "1E1E1E")
        pPr.append(s)
        r = para.add_run(line if line else " ")
        r.font.name = "Courier New"; r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
    doc.add_paragraph()


# Columnas de calificacion con Firma/Fecha
QH  = ["ID", "Descripcion de la Prueba / Verificacion",
        "Criterio de Aceptacion", "Ref. de Evidencia",
        "Resultado Obtenido", "Firma / Fecha"]
QW  = [1.2, 4.8, 3.8, 3.2, 2.5, 2.0]

PQH = ["ID PQ", "Escenario de Uso Real", "Participantes",
       "Resultado Esperado", "Ref. Evidencia",
       "Resultado Obtenido", "Firma / Fecha"]
PQW = [1.0, 4.2, 2.8, 3.2, 2.5, 2.0, 1.8]

CONF = "Conforme"
PEND = "Pendiente"
FF_CONF = "TI — 13/05/2026"
FF_PEND = "_______________"


# ═══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ═══════════════════════════════════════════════════════════════════════════════

def portada(doc):
    doc.add_paragraph(); doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PROTOCOLO MAESTRO DE VALIDACION\nFamSPI v1.0.0")
    r.bold = True; r.font.name = "Calibri"
    r.font.size = Pt(22); r.font.color.rgb = AZUL

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run(f"{DEPTO}  |  {FECHA}")
    rs.italic = True; rs.font.size = Pt(10)
    rs.font.name = "Calibri"; rs.font.color.rgb = GRIS

    _kv(doc, [
        ("Sistema",           SISTEMA),
        ("Marco normativo",   NORMA),
        ("Metodologia",       GAMP),
        ("Version",           VERSION),
        ("Ambito",            "12 modulos / 6 areas / 120+ endpoints"),
        ("Clasificacion",     "INTERNO — Uso restringido"),
        ("Estado",            "Para aprobacion — liberacion condicionada Areas 01 y 02"),
    ], widths=[4.5, 11])
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECCION 1 — CONTROL DOCUMENTAL
# ═══════════════════════════════════════════════════════════════════════════════

def s1_control(doc):
    _h(doc, "1. Control Documental y Aprobaciones")
    _kv(doc, [
        ("Titulo",                   "Protocolo Maestro de Validacion de FamSPI v1.0.0"),
        ("Codigo",                   "TI-VAL-PROTO-001"),
        ("Version",                  VERSION),
        ("Fecha de elaboracion",     FECHA),
        ("Fecha de revision",        "_______________"),
        ("Fecha de aprobacion",      "_______________"),
        ("Alcance",                  "12 modulos / 6 areas de negocio / todas las funciones criticas de FamSPI v1.0.0"),
        ("Departamento responsable", DEPTO),
        ("Marco normativo principal",NORMA),
        ("Metodologia de validacion",GAMP),
        ("Referencias adicionales",  "21 CFR Part 11 (referencia); LOPDP Ecuador; ISO/IEC 27001 (seguridad)"),
        ("Estado",                   "Para aprobacion — liberacion parcial Areas 01 y 02 condicionada a firma"),
        ("Localizacion del expediente","docs/validation/ — repositorio FamSPI (rama principal)"),
    ], widths=[4.5, 12])

    _h(doc, "Historial de revisiones", level=2)
    _table(doc,
        ["Version", "Fecha", "Descripcion del cambio", "Autor", "Aprobado por"],
        [["1.0", FECHA,
          "Version inicial — ciclo completo DQ/IQ/OQ/PQ con HLRA/FRA alineado WHO TRS 1019 y GAMP 5",
          "Departamento de TI", "_______________"]],
        widths=[1.5, 2.5, 8, 3.5, 2.5], last_estado=False)

    _h(doc, "Firmas de aprobacion", level=2)
    t = doc.add_table(rows=4, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(["Realizado por", "Revisado por", "Aprobado por QA"]):
        _cell(t.rows[0].cells[ci], h, bold=True,
              fill="1F497D", color="FFFFFF", center=True, size=9)
    labels = ["Nombre:", "Cargo:", "Firma / Fecha:"]
    for ri in range(1, 4):
        for ci in range(3):
            _cell(t.rows[ri].cells[ci],
                  labels[ri-1] + "  ________________________", size=9, fill="FAFAFA")
    _W(t, [5.5, 5.5, 5.5])
    doc.add_paragraph()
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECCION 2 — GLOSARIO
# ═══════════════════════════════════════════════════════════════════════════════

def s2_glosario(doc):
    _h(doc, "2. Glosario y Definiciones Tecnicas")
    _p(doc, (
        "El presente glosario unifica la terminologia regulatoria establecida por la "
        "Organizacion Mundial de la Salud (WHO TRS 1019) y la guia GAMP 5 con los "
        "componentes tecnicos especificos del sistema FamSPI v1.0.0. Su proposito es "
        "garantizar una interpretacion inequivoca de los criterios de aceptacion, "
        "las evidencias y los registros de prueba durante la revision por parte de "
        "auditores internos o externos, eliminando ambiguedades terminologicas que "
        "podrian comprometer la auditabilidad del expediente de validacion."
    ))

    _h(doc, "2.1 Terminos normativos — WHO TRS 1019 y GAMP 5", level=2)
    _table(doc,
        ["Termino", "Definicion"],
        [
            ["Audit Trail\n(Bitacora de auditoria)",
             "Registro electronico seguro, cronologico e inmutable de las acciones que crean, "
             "modifican o eliminan registros regulados. WHO TRS 1019 Sec. 3.2 exige que sea "
             "automatico, protegido contra alteracion y auditado regularmente. En FamSPI: "
             "tabla auditoria.logs con campos usuario_email, modulo, accion, ip, duracion_ms "
             "y timestamp de base de datos (inmutable)."],
            ["Data Integrity\n(Integridad de datos)",
             "Garantia de que los datos son atribuibles, legibles, contemporaneos, originales "
             "y exactos (principios ALCOA+) durante todo su ciclo de vida. En sistemas "
             "computarizados, la integridad se protege mediante controles de acceso (RBAC), "
             "trazabilidad (audit trail) y validacion de entradas."],
            ["V-Model (Modelo en V)",
             "Metodologia de desarrollo y verificacion de GAMP 5 que empareja cada nivel de "
             "especificacion con su protocolo de prueba correspondiente: URS ↔ PQ, "
             "FS/FRS ↔ OQ, DS/DDS ↔ IQ, Codigo ↔ DQ. Garantiza trazabilidad bidireccional "
             "entre requerimientos y evidencia de verificacion."],
            ["HLRA\n(High Level Risk Assessment)",
             "Evaluacion de riesgos a alto nivel que categoriza el sistema segun su impacto "
             "potencial en la calidad del producto, seguridad del paciente o integridad de "
             "los datos. Define la categoria GAMP del sistema y el nivel de rigor de "
             "validacion requerido (GAMP 5, Cap. 7.3)."],
            ["FRA\n(Functional Risk Assessment)",
             "Evaluacion de riesgos a nivel funcional aplicada mediante FMEA "
             "(Failure Mode and Effects Analysis) para identificar modos de falla en cada "
             "funcion critica. Informa directamente los casos de prueba OQ que deben "
             "retar cada control de mitigacion (GAMP 5, Cap. 7.4)."],
            ["FMEA\n(Failure Mode & Effects Analysis)",
             "Tecnica analitica que evalua cada modo de falla potencial segun: "
             "S = Severidad (impacto del fallo), P = Probabilidad (frecuencia esperada) y "
             "D = Detectabilidad (facilidad de deteccion). RPN = S x P x D. "
             "RPN >= 30 requiere verificacion obligatoria en OQ."],
            ["Qualification\n(Calificacion)",
             "Conjunto de actividades documentadas que demuestran que la infraestructura, "
             "el software y el entorno estan instalados, operan y rinden conforme a "
             "especificaciones. Fases: DQ (Diseno), IQ (Instalacion), OQ (Operacional), "
             "PQ (Desempeno). Cada fase es prerequisito de la siguiente."],
            ["Estado Validado",
             "Condicion de un sistema computarizado en la que se ha demostrado mediante "
             "evidencia documentada que opera consistentemente conforme a sus especificaciones "
             "aprobadas. Solo se mantiene si los cambios siguen el proceso de control "
             "de cambios y se ejecutan revisiones periodicas (WHO TRS 1019, Sec. 14)."],
        ],
        widths=[3.8, 13.7], last_estado=False)

    _h(doc, "2.2 Terminos del sistema FamSPI", level=2)
    _table(doc,
        ["Termino", "Definicion en el contexto de FamSPI"],
        [
            ["API REST",
             "Interfaz de programacion del backend de FamSPI que expone recursos mediante "
             "HTTP (GET, POST, PUT, PATCH, DELETE). Cada endpoint es un punto de control "
             "auditable con autenticacion JWT obligatoria (excepto rutas declaradas "
             "en publicPaths.js)."],
            ["JWT\n(JSON Web Token)",
             "Credencial firmada (RFC 7519) emitida por FamSPI tras autenticacion Google. "
             "Contiene claims obligatorios: iss=spi-fam-backend, aud=spi-fam-frontend, "
             "sub=usuario_id. La firma se verifica con SECRET_KEY en cada solicitud "
             "a rutas privadas. Duracion corta; renovable con refreshToken."],
            ["OAuth2 (Google)",
             "Protocolo de autorizacion delegada (RFC 6749) usado para autenticar usuarios "
             "de FamSPI mediante cuentas Google corporativas. El sistema valida que el "
             "dominio del correo pertenezca al dominio autorizado antes de emitir sesion. "
             "Riesgo principal: R-002 (RPN=30) de la matriz FMEA."],
            ["RBAC\n(Role-Based Access Control)",
             "Modelo de control de acceso implementado en roles.js con 15 grupos de roles "
             "(comercial, ti, finanzas, talento_humano, gerencia, etc.) y SUPER_ROLES "
             "{admin, administrador} con bypass total. Aplicado mediante "
             "requireRole() en cada endpoint privado."],
            ["Endpoint",
             "URL especifica de la API REST. Ejemplo: POST /api/v1/talento-humano/permisos. "
             "Cada endpoint tiene roles permitidos, validaciones de entrada y registro "
             "en auditoria.logs para metodos de mutacion."],
            ["PostgreSQL / Neon",
             "Sistema de gestion de base de datos relacional serverless en la nube. "
             "FamSPI configura un pool de 20 conexiones maximo, SSL habilitado, "
             "retry automatico en errores transitorios (ETIMEDOUT, ECONNRESET, 57P01) "
             "y transacciones ACID para garantizar integridad de datos."],
            ["PM2",
             "Gestor de procesos Node.js en produccion. Garantiza disponibilidad continua "
             "del sistema mediante reinicio automatico, separacion de instancias y "
             "logs de proceso. Equivalente operativo de un watchdog de proceso."],
        ],
        widths=[3.8, 13.7], last_estado=False)

    _h(doc, "2.3 Terminos de negocio", level=2)
    _table(doc,
        ["Termino", "Definicion regulatoria en FamSPI"],
        [
            ["Permiso / Vacaciones",
             "Solicitud formal de ausencia laboral gestionada en FamSPI con flujo de "
             "aprobacion jerarquica (creacion → pendiente → aprobado/rechazado), "
             "link de verificacion legal por QR y registro inmutable en auditoria.logs. "
             "Afecta directamente el saldo de dias del colaborador y la nomina."],
            ["Saldo de vacaciones",
             "Contador de dias disponibles por colaborador. Su integridad es critica "
             "(FMEA R-011, RPN=24): un calculo incorrecto impacta directamente "
             "en nomina y cumplimiento laboral. Verificado en OQ-007 y OQ-008."],
            ["Viatico",
             "Anticipo o reembolso de gastos de viaje gestionado en FamSPI mediante "
             "workspace con asistente de 4 pasos, carga de facturas SRI formato TXT "
             "14 columnas y flujo de 9 estados de aprobacion con trazabilidad completa."],
            ["Business Case",
             "Expediente de oportunidad comercial vinculado a cliente, con estados, "
             "flujo de revision y documentacion requerida por el area comercial. "
             "Gestionado en los modulos del Area 03."],
        ],
        widths=[3.8, 13.7], last_estado=False)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECCION 3 — INTRODUCCION, ALCANCE Y MODELO EN V
# ═══════════════════════════════════════════════════════════════════════════════

def s3_intro(doc):
    _h(doc, "3. Introduccion, Alcance y Estrategia de Validacion (Modelo en V)")
    _p(doc, (
        "La validacion de FamSPI v1.0.0 se rige por WHO TRS 1019 Annex 3, Appendix 5 "
        "y se ejecuta aplicando el Modelo en V (V-Model) de GAMP 5 como marco metodologico "
        "para minimizar los riesgos de calidad, garantizar la integridad de los datos y "
        "asegurar la trazabilidad bidireccional entre requerimientos y evidencia de prueba. "
        "El Modelo en V establece que cada nivel de especificacion tiene un nivel de "
        "verificacion correspondiente: las Especificaciones de Usuario (URS) se verifican "
        "mediante la Calificacion de Desempeno (PQ/UAT), las Especificaciones Funcionales "
        "(FRS/FS) mediante la Calificacion Operacional (OQ), las Especificaciones de Diseno "
        "(DDS) mediante la Calificacion de Instalacion (IQ) y el codigo fuente mediante la "
        "Calificacion de Diseno (DQ). Esta correspondencia garantiza que ningun requerimiento "
        "quede sin verificacion documentada."
    ))
    _p(doc, (
        "El ciclo de vida de la validacion comprende cinco subprocesos secuenciales: "
        "(1) Planificacion — definicion del alcance, estrategia de riesgo y este protocolo; "
        "(2) Especificacion — documentacion de URS, FRS y DDS en los Anexos A y B; "
        "(3) Construccion y Configuracion — desarrollo interno verificado mediante control "
        "de versiones Git y revision de codigo; "
        "(4) Verificacion — ejecucion de los protocolos DQ, IQ, OQ y PQ con evidencia "
        "objetiva; y (5) Reporte — consolidacion del informe final de validacion, "
        "declaracion del estado validado y apertura del periodo de mantenimiento."
    ))

    _h(doc, "3.1 Modulos y areas en alcance", level=2)
    _table(doc,
        ["Area", "Modulos incluidos", "Impacto en datos criticos", "Estado validacion"],
        [
            ["Area 01\nGobierno, Seguridad y Acceso",
             "Autenticacion Google OAuth2, usuarios y roles (RBAC), "
             "bitacora de auditoria, deteccion de accesos fuera de horario",
             "Alto — datos de acceso, sesiones, trazabilidad de todas las acciones",
             "DQ/IQ/OQ/PQ — Conforme 13/05/2026"],
            ["Area 02\nPersonas y Talento Humano",
             "Colaboradores y perfiles, permisos laborales, vacaciones, "
             "solicitudes de personal, asistencia, departamentos",
             "Alto — saldos laborales, cumplimiento LOPDP, datos de nomina",
             "DQ/IQ/OQ/PQ — Conforme 13/05/2026"],
            ["Area 03\nComercial y Business Case",
             "Clientes, oportunidades, business cases, "
             "entregas de colaboradores, integracion CRM",
             "Medio — datos comerciales y de clientes",
             "DQ/IQ conforme — OQ/PQ pendientes"],
            ["Area 04\nServicio Tecnico y Operaciones",
             "Mantenimientos preventivos/correctivos, capacitaciones, "
             "proyectos externos, certificaciones",
             "Medio — registros tecnicos y de servicio",
             "DQ/IQ conforme — OQ/PQ pendientes"],
            ["Area 05\nCompras, Inventario y Logistica",
             "Activos TI con actas, tickets de soporte, "
             "gestion de equipos, compras privadas",
             "Medio — trazabilidad de activos y solicitudes TI",
             "DQ/IQ conforme — OQ/PQ pendientes"],
            ["Area 06\nFinanzas y Transacciones",
             "Viaticos (wizard 4 pasos, 9 estados), firma digital SHA-256, "
             "notificaciones, dashboard, reportes y exportacion ATS/XML",
             "Alto — datos financieros, documentos con validez legal",
             "DQ/IQ conforme — OQ/PQ pendientes"],
        ],
        widths=[2.8, 6.5, 4.5, 3.7], last_estado=False)

    _h(doc, "3.2 Exclusiones explicitas del alcance", level=2)
    _bullet(doc, [
        "Infraestructura fisica del servidor y conectividad de red de terceros "
        "(calificacion de proveedores fuera de alcance).",
        "Servicios de terceros como Google OAuth, Neon PostgreSQL y Gmail API: "
        "se consideran proveedores calificados; se verifica la integracion, no el servicio.",
        "Modulos en desarrollo activo no desplegados en produccion al corte v1.0.0.",
        "Funcionalidades de Business Intelligence y reportes avanzados pendientes.",
        "Aplicacion movil (no existe en v1.0.0 — SPA web unicamente).",
        "Procesos de negocio ejecutados fuera del sistema (flujos en papel o por correo).",
    ])
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECCION 4 — QMS Y SOPs
# ═══════════════════════════════════════════════════════════════════════════════

def s4_qms(doc):
    _h(doc, "4. Sistema de Gestion de Calidad (QMS) y Procedimientos Operativos (SOPs)")
    _p(doc, (
        "De conformidad con la seccion 11 de WHO TRS 1019 Annex 3, el estado validado de "
        "FamSPI se mantiene no solo mediante la ejecucion del protocolo de calificacion, "
        "sino a traves de un ecosistema de Procedimientos Operativos Estandar (SOPs) que "
        "gobiernan el ciclo de vida post-validacion del sistema. Estos SOPs son documentos "
        "controlados, aprobados y de acceso restringido, que constituyen un prerequisito "
        "para iniciar la fase de Calificacion de Desempeno (PQ). Sin la existencia y "
        "aceptacion formal de estos procedimientos, el sistema no puede ser declarado "
        "en estado validado completo, independientemente del resultado de las pruebas tecnicas."
    ))
    _p(doc, "Los siguientes procedimientos son requisitos previos al cierre de la validacion:",
       bold=True, size=9)
    _bullet(doc, [
        "SOP-VAL-001 — Control de Accesos y Gestion de Roles: define el proceso de "
        "alta, modificacion y baja de usuarios en FamSPI, la asignacion y revision "
        "de roles RBAC, y la revocacion inmediata de accesos al terminar la relacion "
        "laboral. Referencia: docs/validation/general/14_operacion_mantenimiento.md.",
        "SOP-VAL-002 — Copias de Seguridad y Restauracion de Datos: documenta la "
        "frecuencia de respaldos de la base de datos PostgreSQL/Neon, el procedimiento "
        "de restauracion verificado y el registro del resultado en bitacora operacional. "
        "Prerequisito para cierre de IQ-012.",
        "SOP-VAL-003 — Gestion de Incidentes de Seguridad: define los pasos de contencion, "
        "investigacion, documentacion y notificacion ante eventos de acceso no autorizado, "
        "perdida de datos o deteccion de cambios no controlados, incluyendo apertura de "
        "desviaciones formales.",
        "SOP-VAL-004 — Control de Cambios de Software (Change Control): establece que "
        "todo cambio al codigo fuente, configuracion o esquema de base de datos de FamSPI "
        "debe ser evaluado, aprobado y documentado antes de su implementacion, con "
        "re-ejecucion de los casos de prueba OQ/PQ afectados. Referencia: 14A_control_cambios.md.",
        "SOP-VAL-005 — Revisiones Periodicas del Sistema Validado: establece la frecuencia "
        "(minimo anual), los criterios y el responsable de la revision del estado validado, "
        "incluyendo auditoria de usuarios/roles, revision de logs, dependencias con "
        "vulnerabilidades (npm audit) y verificacion de certificados TLS. "
        "Referencia: 14B_revision_periodica.md.",
        "SOP-VAL-006 — Gestion de Desviaciones y No Conformidades: define la "
        "clasificacion (critica / mayor / menor), el proceso de apertura, investigacion "
        "de causa raiz, plan de accion correctiva (CAPA) y cierre formal de toda "
        "desviacion detectada durante o despues de la validacion.",
    ])
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECCION 5 — ESPECIFICACIONES
# ═══════════════════════════════════════════════════════════════════════════════

def s5_specs(doc):
    _h(doc, "5. Especificaciones del Sistema (URS / FS / DDS)")
    _p(doc, (
        "En el marco del Modelo en V (GAMP 5), las especificaciones constituyen el brazo "
        "izquierdo del ciclo de vida: cada nivel de especificacion dicta que se construira "
        "y sera el referente contra el que se verificara durante las calificaciones. "
        "La ausencia de especificaciones formales y aprobadas impide la trazabilidad "
        "requerida por WHO TRS 1019 y hace inauditable cualquier resultado de prueba."
    ))
    _p(doc, (
        "Las Especificaciones de Requerimientos de Usuario (URS) cubren 12 modulos del "
        "sistema con 120+ requerimientos catalogados mediante identificador unico "
        "(REQ-{MODULO}-NNN), nivel de criticidad (Alta / Media / Baja) y actor responsable. "
        "Las Especificaciones Funcionales (FS/FRS) describen el endpoint exacto, las "
        "entradas requeridas, el proceso ejecutado y la salida HTTP esperada para cada "
        "requerimiento critico, estableciendo la base tecnica para los casos de prueba OQ. "
        "La Especificacion de Diseno del Sistema (DDS) documenta la arquitectura verificada "
        "en codigo fuente: capa de presentacion (React 19), capa de logica de negocio "
        "(Express 5.1.0 con cadena de middlewares: Helmet → Rate Limit → CORS → JWT → "
        "Audit → Routes) y capa de persistencia (PostgreSQL Neon, pool de 20 conexiones, "
        "SSL, retry en errores transitorios ETIMEDOUT/ECONNRESET/57P01). "
        "Las tres especificaciones estan consolidadas formalmente con extractos de codigo "
        "fuente como evidencia objetiva en los Anexos A (URS/FS) y B (DDS) de este "
        "expediente de validacion."
    ))
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECCION 6 — HLRA Y FRA
# ═══════════════════════════════════════════════════════════════════════════════

def s6_riesgos(doc):
    _h(doc, "6. Evaluacion de Riesgos a Doble Nivel (HLRA y FRA)")
    _p(doc, (
        "La gestion de riesgos en este protocolo se articula en dos fases metodologicas "
        "complementarias segun GAMP 5 Segunda Edicion, Capitulo 7: la Evaluacion de Riesgos "
        "a Alto Nivel (HLRA) que categoriza el sistema y define el rigor de validacion "
        "requerido, y la Evaluacion de Riesgos Funcional (FRA) que desciende a nivel de "
        "cada funcion critica para identificar, cuantificar y mitigar los modos de falla "
        "especificos, alimentando directamente los casos de prueba OQ obligatorios."
    ))

    _h(doc, "6.1 HLRA — Evaluacion a Alto Nivel", level=2)
    _p(doc, (
        "La HLRA evalua el sistema FamSPI en su conjunto para determinar su categoria GAMP "
        "y el nivel de impacto regulatorio de una falla del sistema. Los criterios evaluados "
        "son: impacto en la seguridad del personal, impacto en la integridad de registros "
        "regulados (LOPDP Ecuador), impacto en la calidad de los datos de negocio y "
        "criticidad de las funciones para la continuidad operacional."
    ))
    _table(doc,
        ["Criterio HLRA", "Evaluacion", "Justificacion"],
        [
            ["Categoria GAMP del sistema",
             "Categoria 4 — Sistema computarizado desarrollado a medida",
             "FamSPI es un sistema de informacion desarrollado internamente sobre un "
             "framework estandar (Express.js). No es Categoria 3 (configurable) ni "
             "Categoria 5 (infraestructura), ya que incluye logica de negocio "
             "especificamente codificada."],
            ["Impacto en registros electronicos regulados",
             "ALTO",
             "FamSPI gestiona registros laborales (permisos, vacaciones) y documentos "
             "con validez legal (actas con firma SHA-256). La perdida o alteracion no "
             "controlada constituye incumplimiento de la LOPDP Ecuador."],
            ["Impacto en la seguridad del personal",
             "MEDIO",
             "El sistema no gestiona datos de salud criticos ni procesos de seguridad "
             "fisica. El impacto en seguridad del personal se limita a la confidencialidad "
             "de datos laborales (nomina, asistencia, saldo de vacaciones)."],
            ["Continuidad operacional",
             "ALTO",
             "Una indisponibilidad prolongada del sistema paraliza la gestion de permisos, "
             "viaticos y asistencia de todos los colaboradores. RTO objetivo: < 4 horas."],
            ["Nivel de rigor de validacion resultante",
             "COMPLETO (DQ + IQ + OQ + PQ)",
             "Dado el impacto ALTO en registros regulados y continuidad operacional, "
             "el nivel de rigor requerido es la ejecucion del ciclo completo de "
             "calificacion con evidencia documentada para todas las funciones criticas."],
        ],
        widths=[3.8, 2.8, 11], last_estado=False)

    _h(doc, "6.2 FRA — Evaluacion Funcional (FMEA): riesgos que requieren verificacion obligatoria en OQ", level=2)
    _p(doc, (
        "La FRA aplica FMEA (RPN = Severidad x Probabilidad x Detectabilidad, escala 1-5) "
        "a las funciones criticas identificadas en la HLRA. Los modos de falla con "
        "RPN >= 30 o Severidad = 5 sin mitigacion documentada son condicion necesaria "
        "para bloquear la declaracion del estado validado hasta su verificacion en OQ. "
        "La matriz completa con los 15 modos de falla (R-001 a R-015) se encuentra "
        "en el Anexo C."
    ))
    _table(doc,
        ["ID", "Funcion afectada", "Modo de falla", "S", "P", "D", "RPN",
         "Nivel", "Control implementado", "Verificacion OQ"],
        [
            ["R-002", "Autenticacion OAuth2",
             "Compromiso de credenciales Google / replay de token JWT",
             "5", "2", "3", "30", "ALTO",
             "JWT de corta duracion; claims iss/aud/sub validados; "
             "dominio corporativo restringido; refresh token rotation",
             "OQ-014 — OBLIGATORIO"],
            ["R-009", "Control de acceso RBAC",
             "Asignacion incorrecta de rol — acceso a funciones no autorizadas",
             "4", "2", "3", "24", "Medio",
             "requireRole() en cada endpoint; SUPER_ROLES solo admin; "
             "doble capa con guards assertX() en servicios criticos",
             "OQ-005"],
            ["R-011", "Permisos / Vacaciones",
             "Calculo incorrecto de saldo — impacto en nomina",
             "4", "2", "3", "24", "Medio",
             "Validacion unitaria de saldo; pruebas OQ especificas; "
             "historial auditable de cada transicion de estado",
             "OQ-007 / OQ-008"],
            ["R-005", "Bitacora de auditoria",
             "Perdida o alteracion de registros de audit trail",
             "5", "1", "2", "10", "Bajo",
             "auditoria.logs en PostgreSQL persistente; "
             "usuarios no pueden eliminar registros; exportacion CSV protegida",
             "OQ-006"],
        ],
        widths=[1.0, 2.5, 3.8, 0.5, 0.5, 0.5, 0.8, 1.2, 4.5, 2.5],
        last_estado=False, font=8)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECCION 7 — PLAN DE VERIFICACION Y PROTOCOLOS
# ═══════════════════════════════════════════════════════════════════════════════

def s7_plan(doc):
    _h(doc, "7. Plan de Verificacion y Protocolos de Prueba")
    _p(doc, (
        "Esta seccion contiene los cuatro protocolos de calificacion del brazo derecho "
        "del Modelo en V. Cada subseccion presenta primero el contexto metodologico "
        "de la fase — que es, por que se realiza y que controles verifica — seguido "
        "de la tabla de ejecucion con columnas: ID, descripcion precisa, criterio de "
        "aceptacion verificable, referencia de evidencia, resultado obtenido y firma con "
        "fecha del ejecutor. Las Areas 01 y 02 tienen resultado conforme (13/05/2026). "
        "Las Areas 03 a 06 tienen campos de resultado listos para ser completados durante "
        "la ejecucion formal del protocolo."
    ), after=6)

    # ── 7.1 DQ ─────────────────────────────────────────────────────────────────
    _h(doc, "7.1 Calificacion de Diseno (DQ)", level=2)
    _p(doc, (
        "La Calificacion de Diseno (DQ) verifica documentalmente, antes de la instalacion, "
        "que el diseno tecnico propuesto de FamSPI es adecuado para su proposito previsto "
        "y que las URS aprobadas se reflejan en las especificaciones de diseno (DDS). "
        "En el Modelo en V, la DQ garantiza que el lado izquierdo (especificaciones) "
        "es coherente, completo y trazable antes de construir o instalar cualquier "
        "componente del sistema. Una DQ fallida indica deficiencias de diseno que, "
        "de no corregirse, propagarian errores a todas las fases siguientes. "
        "La DQ de FamSPI se ejecuta mediante revision documental del DDS, del repositorio "
        "de codigo y de la matriz FMEA."
    ), size=9)
    _table(doc, QH, [
        ["DQ-001",
         "Arquitectura de tres capas documentada con tecnologias, versiones y "
         "archivo de configuracion de cada capa verificado en codigo fuente",
         "DDS con descripcion completa: presentacion (React 19), logica (Express 5.1.0), "
         "persistencia (PostgreSQL Neon). Cada capa con archivo de config identificado",
         "Anexo B, Sec. B.1 — tabla de arquitectura con evidencia de codigo",
         CONF, FF_CONF],
        ["DQ-002",
         "Segmentacion explicita de rutas publicas/privadas declarada y verificada "
         "en publicPaths.js y registerRoutes.js",
         "15 prefijos publicos declarados en publicPaths.js; mountPublicRoutes() y "
         "mountPrivateRoutes() separados sin superposicion",
         "Anexo B, Sec. B.3; Anexo D, EV-D-005",
         CONF, FF_CONF],
        ["DQ-003",
         "URS documentada para 12 modulos con REQ-IDs unicos, criticidad "
         "y actor responsable en cada requerimiento",
         "12 archivos URS en docs/validation/URS/ con 120+ REQ catalogados. "
         "Todo REQ de criticidad Alta tiene FRS correspondiente",
         "Anexo A — tabla de 12 modulos con estado documental",
         CONF, FF_CONF],
        ["DQ-004",
         "Evaluacion FMEA completada: 15 modos de falla con S/P/D/RPN, "
         "mitigacion documentada y R-002 identificado como ALTO (RPN=30)",
         "Matriz FMEA con R-001 a R-015; columnas S/P/D/RPN/mitigacion/estado; "
         "R-002 marcado ACTIVO con verificacion obligatoria en OQ-014",
         "Anexo C — Matriz FMEA completa; general/07A_evaluacion_riesgos_fmea.md",
         CONF, FF_CONF],
        ["DQ-005",
         "Todos los REQ de criticidad Alta tienen endpoint FRS implementado "
         "y verificable en codigo fuente activo en produccion",
         "REQ-AUTH-001 a REQ-AUTH-012 y REQ-TH-001 a REQ-TH-006 con FRS "
         "verificado y endpoint activo en produccion",
         "Anexo A, Tablas FRS-AUTH y FRS-TH; Anexo D, EV-D-001 a EV-D-005",
         CONF, FF_CONF],
    ], widths=QW)

    # ── 7.2 IQ ─────────────────────────────────────────────────────────────────
    _h(doc, "7.2 Calificacion de Instalacion (IQ)", level=2)
    _p(doc, (
        "La Calificacion de Instalacion (IQ) provee evidencia documentada de que el "
        "software (Node.js, Express, dependencias npm), la infraestructura (PM2, HTTPS/TLS) "
        "y la base de datos (PostgreSQL/Neon) estan correctamente instalados y configurados "
        "en el entorno de produccion de acuerdo con el DDS aprobado. La IQ no verifica "
        "funcionalidad — eso corresponde a la OQ — sino la correcta instalacion y "
        "configuracion del entorno. Un fallo en la IQ indica que el entorno no esta "
        "preparado para ejecutar pruebas funcionales validas. La IQ es prerequisito "
        "formal para iniciar la OQ."
    ), size=9)
    _table(doc, QH, [
        ["IQ-001",
         "Node.js instalado en servidor de produccion con version compatible (v18.x+)",
         "node --version retorna v18.x o superior en el servidor productivo",
         "Captura de terminal: node --version en servidor de produccion",
         CONF, FF_CONF],
        ["IQ-002",
         "Express 5.1.0 instalado segun package.json; version activa verificada",
         "package.json declara 'express': '^5.1.0'; npm list retorna 5.1.x instalado",
         "backend/package.json; salida npm list express",
         CONF, FF_CONF],
        ["IQ-003",
         "Variables de entorno criticas presentes y no vacias en produccion",
         "SECRET_KEY, DATABASE_URL, GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET, "
         "NODE_ENV=production: todas presentes y no vacias",
         "Verificacion con script de arranque (valores no registrados por seguridad)",
         CONF, FF_CONF],
        ["IQ-004",
         "Conexion a PostgreSQL/Neon activa; pool inicializado; "
         "/health retorna db_status: connected",
         "Aplicacion inicia sin error de conexion; /health retorna HTTP 200 "
         "con db_status: connected",
         "Log de arranque PM2; respuesta de /health en produccion",
         CONF, FF_CONF],
        ["IQ-005",
         "Proceso PM2 activo en estado 'online'; restart_policy configurada",
         "pm2 list muestra proceso FamSPI en estado 'online'; "
         "pm2 show reporta restart correctamente configurado",
         "Captura de pm2 list y pm2 show en servidor de produccion",
         CONF, FF_CONF],
        ["IQ-006",
         "HTTPS/TLS activo en produccion; HTTP redirige a HTTPS (HTTP 301)",
         "Acceso https:// retorna HTTP 200; acceso http:// retorna HTTP 301 "
         "con Location: https://",
         "Verificacion via curl del dominio productivo en ambos protocolos",
         CONF, FF_CONF],
        ["IQ-007",
         "Rate limiting activo: 3000 req/15min en produccion; "
         "cabecera X-RateLimit-Limit presente",
         "Cabecera X-RateLimit-Limit: 3000 presente en respuestas de la API; "
         "configuracion windowMs=900000 verificada en security.js",
         "Inspeccion de cabeceras HTTP; backend/src/config/security.js",
         CONF, FF_CONF],
        ["IQ-008",
         "Cabeceras de seguridad Helmet activas: CSP, HSTS y X-Frame-Options",
         "Respuesta HTTP contiene Content-Security-Policy, "
         "Strict-Transport-Security (max-age>=31536000) y X-Frame-Options: DENY",
         "Inspeccion de cabeceras HTTP en produccion via curl -I",
         CONF, FF_CONF],
        ["IQ-009",
         "Esquema de base de datos desplegado: tablas principales creadas y accesibles",
         "Consulta a information_schema.tables retorna todas las tablas "
         "principales del sistema sin errores",
         "Script de verificacion de esquema ejecutado y registrado",
         CONF, FF_CONF],
        ["IQ-010",
         "Frontend React compilado y servido; SPA carga sin errores HTTP 4xx/5xx",
         "Acceso al dominio carga la SPA React; herramientas de dev muestran "
         "sin errores criticos en la carga inicial",
         "Captura de pantalla del dashboard; sin errores en consola del navegador",
         CONF, FF_CONF],
        ["IQ-011",
         "Usuario administrador inicial activo y funcional en produccion",
         "Al menos un usuario con rol 'administrador' puede autenticarse "
         "y acceder al modulo de administracion",
         "Captura de sesion activa de usuario admin en produccion",
         CONF, FF_CONF],
        ["IQ-012",
         "Procedimiento de rollback documentado; restauracion desde respaldo verificada",
         "Existe SOP-VAL-002 escrito; al menos una restauracion de prueba "
         "ejecutada y documentada con resultado exitoso",
         "SOP-VAL-002; registro de prueba de restauracion",
         PEND, FF_PEND],
    ], widths=QW)

    # ── 7.3 OQ ─────────────────────────────────────────────────────────────────
    _h(doc, "7.3 Calificacion Operacional (OQ)", level=2)
    _p(doc, (
        "La Calificacion Operacional (OQ) somete al sistema a retos funcionales para "
        "demostrar que opera conforme a las especificaciones funcionales (FS/FRS) en "
        "todo su rango operativo definido. La OQ es la fase en que los controles de "
        "mitigacion identificados en el FRA (Sec. 6.2) son retados deliberadamente: "
        "se prueban accesos invalidos, roles incorrectos, tokens expirados y flujos "
        "de error para demostrar que el sistema rechaza las condiciones no autorizadas "
        "con la respuesta esperada. Ademas, la OQ verifica el correcto funcionamiento "
        "del Audit Trail (bitacora de auditoria) como control critico de integridad de "
        "datos exigido por WHO TRS 1019. La OQ de las Areas 01 y 02 se ejecuto en "
        "entorno de produccion el 13/05/2026 con resultado conforme."
    ), size=9)
    _table(doc, QH, [
        ["OQ-001",
         "Login Google OAuth2 con cuenta del dominio corporativo autorizado; "
         "JWT emitido con claims iss/aud/sub correctos; sesion registrada",
         "HTTP 200 con accessToken y refreshToken; JWT decodificado con "
         "iss=spi-fam-backend, aud=spi-fam-frontend, sub=usuario_id; "
         "registro en user_sessions con timestamp",
         "Registro en user_sessions DB; token decodificado; log auditMiddleware",
         CONF, FF_CONF],
        ["OQ-002",
         "Rechazo de cuenta Google con dominio no autorizado; "
         "sin creacion de sesion ni registro en user_sessions",
         "HTTP 403 con mensaje de dominio no permitido; "
         "user_sessions sin nuevo registro post-intento",
         "Respuesta HTTP capturada; verificacion de user_sessions sin registro nuevo",
         CONF, FF_CONF],
        ["OQ-003",
         "Renovacion de sesion via refresh token; nuevo accessToken emitido "
         "sin re-autenticacion; user_sessions actualizado",
         "POST /api/v1/auth/refresh con refresh token valido: HTTP 200 "
         "con nuevo accessToken; user_sessions con timestamp actualizado",
         "Respuesta HTTP 200 con accessToken; registro en user_sessions",
         CONF, FF_CONF],
        ["OQ-004",
         "Logout invalida sesion; token anterior rechazado con HTTP 401 "
         "post-logout (R-003 FMEA)",
         "POST /api/v1/auth/logout: sesion cerrada en user_sessions; "
         "uso del token previo retorna HTTP 401 code: INVALID_TOKEN",
         "Estado en user_sessions; intento de reuso del token con respuesta 401",
         CONF, FF_CONF],
        ["OQ-005",
         "Endpoint protegido rechaza a usuario con rol sin permiso: "
         "HTTP 403 con lista de roles; sin datos filtrados (R-009 FMEA)",
         "Usuario rol 'comercial' en GET /api/v1/auditoria: HTTP 403 "
         "con body {ok:false, error:'Acceso denegado...'}; sin datos del sistema",
         "Respuesta HTTP 403 capturada; body sin datos sensibles del sistema",
         CONF, FF_CONF],
        ["OQ-006",
         "Toda accion POST/PUT/DELETE registrada en auditoria.logs con "
         "los campos ALCOA+ requeridos por WHO (R-005 FMEA)",
         "Despues de cualquier mutacion: auditoria.logs contiene usuario_email, "
         "modulo, accion, ip, duracion_ms, timestamp — todos presentes e inmediatos",
         "Consulta SELECT a auditoria.logs post-mutacion; verificacion de campos",
         CONF, FF_CONF],
        ["OQ-007",
         "Creacion de solicitud de permiso con validaciones; "
         "estado inicial correcto; HTTP 400 en campos obligatorios ausentes",
         "POST /api/v1/talento-humano/permisos: HTTP 201 con estado 'pendiente'; "
         "POST sin campos obligatorios: HTTP 400 con detalle de validacion",
         "Registro en DB; prueba negativa con campos vacios y resultado HTTP 400",
         CONF, FF_CONF],
        ["OQ-008",
         "Aprobacion por responsable autorizado; estado actualizado; "
         "historial con actor y timestamp; registro en auditoria (R-011 FMEA)",
         "PUT estado a 'aprobado': DB con nuevo estado; historial con actor "
         "y timestamp; auditoria.logs con evento de la transicion",
         "DB con estado; historial de estados; registro en auditoria.logs",
         CONF, FF_CONF],
        ["OQ-009",
         "Rechazo de vacaciones con motivo obligatorio; "
         "notificacion emitida al solicitante",
         "Estado 'rechazado' con motivo en DB; notificacion en cola de dispatchers; "
         "registro en auditoria.logs con motivo",
         "DB con estado y motivo; cola de notificaciones; registro de auditoria",
         CONF, FF_CONF],
        ["OQ-010",
         "Link de verificacion legal accesible sin JWT; "
         "datos correctos retornados en HTTP 200",
         "GET /api/v1/permisos/legal-verification/:token sin Bearer token: "
         "HTTP 200 con datos del documento; sin acceso a otros datos del sistema",
         "Acceso al link sin autenticacion; respuesta HTTP 200 correcta",
         CONF, FF_CONF],
        ["OQ-011",
         "Business Case creado y vinculado a cliente/oportunidad; "
         "estado inicial correcto en DB",
         "POST /api/v1/business-case: HTTP 201; BC con estado inicial; "
         "vinculacion a cliente y oportunidad verificada en DB",
         "Registro en DB; respuesta HTTP 201 con ID; vinculacion verificada",
         PEND, FF_PEND],
        ["OQ-012",
         "Flujo de viatico: borrador → revision jefe → aprobacion financiero → pago; "
         "auditoria registra cada transicion de estado",
         "Cada cambio de estado: actor validado, permisos verificados, "
         "estado previo correcto; auditoria.logs con registro de cada paso",
         "Registros de auditoria por transicion; estado final 'pagado' en DB",
         PEND, FF_PEND],
        ["OQ-013",
         "Ticket de soporte TI creado, asignado y cerrado; "
         "trazabilidad completa con auditoria en cada transicion",
         "Flujo completo sin errores; auditoria registra cada transicion; "
         "cierre con causa, tiempo y datos completos en DB",
         "DB con todos los estados; auditoria de cada transicion; cierre documentado",
         PEND, FF_PEND],
        ["OQ-014",
         "R-002 — VERIFICACION OBLIGATORIA FMEA: replay de access token "
         "expirado rechazado; sin acceso concedido bajo ninguna condicion",
         "Token expirado retorna HTTP 401 con code: INVALID_TOKEN; "
         "token de sesion invalida idem; sin datos en respuesta; log en auditoria",
         "Respuesta HTTP 401 capturada; body correcto; verificacion en auth.js",
         PEND, FF_PEND],
    ], widths=QW)

    # ── 7.4 PQ ─────────────────────────────────────────────────────────────────
    _h(doc, "7.4 Calificacion de Desempeno y Pruebas de Usuario (PQ / UAT)", level=2)
    _p(doc, (
        "La Calificacion de Desempeno (PQ) y las Pruebas de Aceptacion de Usuario (UAT) "
        "se ejecutan en el entorno en vivo con usuarios funcionales previamente capacitados "
        "segun el Anexo E. En el Modelo en V, la PQ es el nivel mas alto de verificacion "
        "y corresponde directamente con las URS: confirma que el sistema satisface los "
        "requerimientos del usuario en condiciones representativas de uso real. "
        "Los participantes son los propios usuarios de negocio — no el equipo de TI — "
        "quienes ejecutan flujos de trabajo reales de extremo a extremo. "
        "El entrenamiento formal documentado (Anexo E) es un prerequisito irrenunciable "
        "para la participacion en la PQ, segun WHO TRS 1019 Annex 3, Sec. 12. "
        "Las actas de aceptacion firmadas por el responsable funcional y gerencia "
        "constituyen la evidencia de cierre de cada area."
    ), size=9)
    _table(doc, PQH, [
        ["PQ-001",
         "Permiso punta a punta: creacion → aprobacion → verificacion QR legal",
         "Solicitante + Aprobador + TI",
         "Flujo completo sin errores; estado final correcto; QR funcional; "
         "trazabilidad completa en auditoria.logs",
         "Acta UAT Area 02; capturas del flujo",
         CONF, FF_CONF],
        ["PQ-002",
         "Vacaciones punta a punta: creacion → aprobacion → link legal accesible",
         "Solicitante + Aprobador",
         "Flujo completo; estado final correcto en DB; "
         "link de verificacion legal activo sin JWT",
         "Acta UAT Area 02; capturas del flujo completo",
         CONF, FF_CONF],
        ["PQ-003",
         "Aprobacion y rechazo con distintos motivos; "
         "notificaciones emitidas; historial auditable",
         "Responsable funcional + TI",
         "Ambas decisiones ejecutadas; notificaciones en cola; "
         "historial con actor y timestamp auditable",
         "Capturas de ambas decisiones; notificaciones en DB",
         CONF, FF_CONF],
        ["PQ-004",
         "Consulta de historial con filtros de estado, fecha y usuario; paginacion",
         "TI + Funcional",
         "Filtros retornan resultados correctos; paginacion operativa; "
         "consistencia con datos directos de DB",
         "Capturas de consultas; comparacion con datos en DB",
         CONF, FF_CONF],
        ["PQ-005",
         "Aceptacion formal Areas 01 y 02 por responsable funcional y gerencia",
         "Responsable funcional + Gerencia",
         "Acta de aceptacion firmada con declaracion de conformidad "
         "con el uso previsto del sistema",
         "Actas firmadas en area_01_*/ y area_02_*/",
         CONF, FF_CONF],
        ["PQ-006",
         "Business Case completo: creacion → seguimiento → cierre",
         "Comercial + Gerencia",
         "BC pasa todos los estados; documentacion completa; "
         "vinculacion activa a cliente",
         "Capturas del flujo; datos en DB verificados",
         PEND, FF_PEND],
        ["PQ-007",
         "Registro de visita comercial y asignacion de cliente",
         "Comercial + Administrador",
         "Visita registrada con fecha, actor y resultado; "
         "cliente asignado correctamente",
         "Capturas de registro; datos en DB; auditoria del evento",
         PEND, FF_PEND],
        ["PQ-008",
         "Aceptacion formal modulo comercial",
         "Jefe Comercial + Gerencia",
         "Acta de aceptacion firmada declarando conformidad",
         "Acta de aceptacion Area 03",
         PEND, FF_PEND],
        ["PQ-009",
         "Capacitacion punta a punta: planificacion → ejecucion → certificado PDF",
         "TH + Instructor + Tecnico",
         "Flujo completo sin errores; certificado generado en PDF con datos correctos",
         "Capturas del flujo; certificado PDF; datos en DB",
         PEND, FF_PEND],
        ["PQ-010",
         "Orden de mantenimiento preventivo completa: creacion → asignacion → cierre",
         "Tecnico + Supervisor",
         "Orden pasa todos los estados validos; cierre con evidencia de ejecucion",
         "Capturas del flujo; datos en DB; auditoria de transiciones",
         PEND, FF_PEND],
        ["PQ-011",
         "Aceptacion formal modulo servicio tecnico",
         "Jefe Tecnico + Gerencia",
         "Acta de aceptacion firmada declarando conformidad",
         "Acta de aceptacion Area 04",
         PEND, FF_PEND],
        ["PQ-012",
         "Alta, asignacion y baja de activo TI con acta generada en cada estado",
         "TI + Usuario receptor",
         "Activo: alta → asignado → dado_de_baja; acta PDF en cada transicion",
         "Capturas de cada estado; actas PDF generadas",
         PEND, FF_PEND],
        ["PQ-013",
         "Ticket de soporte TI de creacion a cierre con SLA calculado",
         "Usuario + Tecnico TI",
         "Ticket: creacion → asignado → en_progreso → resuelto → cerrado; "
         "SLA calculado y visible en DB",
         "Capturas del flujo; SLA verificado en DB",
         PEND, FF_PEND],
        ["PQ-014",
         "Aceptacion formal modulo inventario/TI",
         "Jefe TI + Gerencia",
         "Acta de aceptacion firmada declarando conformidad",
         "Acta de aceptacion Area 05",
         PEND, FF_PEND],
        ["PQ-015",
         "Viatico punta a punta: borrador → carga facturas TXT SRI → "
         "revision → aprobacion → pago registrado",
         "Solicitante + Jefe area + Finanzas",
         "Flujo completo; facturas TXT parseadas; estado final 'pagado' en DB; "
         "auditoria de cada transicion",
         "Capturas de cada paso del wizard; facturas en DB; estado final",
         PEND, FF_PEND],
        ["PQ-016",
         "Categorizacion de facturas y exportacion ATS/XML formato SRI Ecuador",
         "Finanzas",
         "Facturas categorizadas; archivo ATS/XML generado con estructura "
         "valida segun formato SRI",
         "Archivo ATS/XML exportado; validacion de estructura",
         PEND, FF_PEND],
        ["PQ-017",
         "Aceptacion formal modulo finanzas/viaticos",
         "Jefe Finanzas + Gerencia",
         "Acta de aceptacion firmada declarando conformidad",
         "Acta de aceptacion Area 06",
         PEND, FF_PEND],
    ], widths=PQW)
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECCION 8 — RTM
# ═══════════════════════════════════════════════════════════════════════════════

def s8_rtm(doc):
    _h(doc, "8. Matriz de Trazabilidad Continua (TM / RTM)")
    _p(doc, (
        "En la metodologia GAMP 5, la Matriz de Trazabilidad (RTM) no es un documento "
        "final que se elabora al terminar la validacion — es una herramienta viva que "
        "comienza en la fase DQ uniendo URS con Especificaciones (FS y DDS), se extiende "
        "durante la OQ al vincular cada especificacion con su caso de prueba, y se cierra "
        "durante el reporte final al documentar la evidencia y el estado de cada "
        "verificacion. Su funcion es garantizar que ningun requerimiento del usuario "
        "quede sin probar y que todo riesgo identificado en la FRA este mitigado "
        "y verificado con evidencia objetiva. La RTM completa del sistema se mantiene "
        "en docs/validation/RTM/RTM_sistema_spi.md y se actualiza con cada ciclo "
        "de control de cambios."
    ))
    _table(doc,
        ["ID URS", "Especificacion (FS / DDS)", "Riesgo FMEA",
         "Prueba (DQ/IQ/OQ/PQ)", "Evidencia", "Estado"],
        [
            ["REQ-AUTH-001", "FRS-AUTH-001\n(GET /api/v1/auth/google → callback)",
             "R-002 (RPN=30)", "OQ-001 / PQ-001",
             "user_sessions DB; token decodificado", CONF],
            ["REQ-AUTH-002", "FRS-AUTH-001\n(validacion de dominio en callback)",
             "R-002 (RPN=30)", "OQ-002",
             "Respuesta HTTP 403 capturada; user_sessions sin registro", CONF],
            ["REQ-AUTH-005", "FRS-AUTH-002\n(emision accessToken + refreshToken)",
             "R-003 (RPN=24)", "OQ-003",
             "Respuesta HTTP 200 con tokens; user_sessions actualizado", CONF],
            ["REQ-AUTH-007", "FRS-AUTH-004\n(POST /api/v1/auth/logout)",
             "R-003 (RPN=24)", "OQ-004",
             "user_sessions cerrado; token previo rechazado HTTP 401", CONF],
            ["REQ-USR-002",  "DDS — roles.js requireRole()",
             "R-009 (RPN=24)", "OQ-005",
             "HTTP 403 con roles permitidos; sin datos del sistema filtrados", CONF],
            ["REQ-AUD-001",  "DDS — auditMiddleware.js logAction()",
             "R-005 (RPN=10)", "OQ-006",
             "Campos ALCOA+ en auditoria.logs verificados post-mutacion", CONF],
            ["REQ-SEC-001",  "FRS-SEC-001\n(deteccion off-hours)",
             "R-004 (RPN=24)", "OQ-001",
             "Evento off-hours registrado; notificacion al canal de TI", CONF],
            ["REQ-TH-001",   "FRS-TH-001\n(POST permisos)",
             "R-011 (RPN=24)", "OQ-007 / PQ-001",
             "Registro en DB estado 'pendiente'; HTTP 201 verificado", CONF],
            ["REQ-TH-002",   "FRS-TH-003\n(PUT permisos/:id/estado)",
             "R-011 (RPN=24)", "OQ-008 / OQ-009 / PQ-002",
             "DB con nuevo estado; historial con actor; auditoria.logs", CONF],
            ["REQ-TH-006",   "FRS-TH-004\n(GET legal-verification/:token)",
             "R-007 (RPN=20)", "OQ-010 / PQ-003",
             "HTTP 200 sin JWT; datos del documento verificados", CONF],
            ["REQ-COM-001",  "FRS-COM-001\n(POST business-case)",
             "R-013 (RPN=27)", "OQ-011 / PQ-006",
             "Pendiente ejecucion", PEND],
            ["REQ-VT-001",   "FRS-VT-001\n(wizard 4 pasos)",
             "R-011 (RPN=24)", "OQ-012 / PQ-015",
             "Pendiente ejecucion", PEND],
            ["REQ-INV-001",  "FRS-INV-001\n(ciclo activo TI)",
             "R-008 (RPN=12)", "PQ-012",
             "Pendiente ejecucion", PEND],
            ["REQ-TI-001",   "FRS-TI-001\n(ticket soporte)",
             "R-008 (RPN=12)", "OQ-013 / PQ-013",
             "Pendiente ejecucion", PEND],
            ["REQ-SEC-002",  "DDS — auth.js verifyToken()",
             "R-002 ALTO\n(RPN=30 — OBLIGATORIO)", "OQ-014",
             "Pendiente — BLOQUEA declaracion del estado validado Area 01", PEND],
        ],
        widths=[2.2, 3.5, 2.8, 3.2, 4, 1.8], font=8)
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECCION 9 — INDICE DE ANEXOS
# ═══════════════════════════════════════════════════════════════════════════════

def s9_anexos(doc):
    _h(doc, "9. Indice de Anexos")
    _p(doc, (
        "Los siguientes anexos son parte integral e inseparable del expediente de "
        "validacion de FamSPI v1.0.0. Los documentos pesados — tablas de requerimientos, "
        "extractos de codigo fuente y registros de entrenamiento — se consolidan en "
        "esta seccion para mantener el cuerpo del protocolo enfocado en la narrativa "
        "metodologica y los protocolos de ejecucion. El expediente completo se almacena "
        "en docs/validation/ del repositorio de control de versiones de FamSPI."
    ))
    _table(doc,
        ["Anexo", "Titulo", "Contenido", "Estado"],
        [
            ["A", "URS y FS completas",
             "Tablas de REQ-IDs por los 12 modulos (120+ requerimientos); "
             "tablas FRS con endpoint, entradas, proceso y salida HTTP para "
             "todos los modulos criticos",
             "Incluido en este documento"],
            ["B", "DDS — Especificacion de Diseno",
             "Arquitectura de 3 capas con evidencia de codigo; cadena de "
             "middlewares en orden de ejecucion verificado en app.js; "
             "segmentacion de rutas; configuracion de pool BD; dependencias criticas",
             "Incluido en este documento"],
            ["C", "Matriz HLRA / FRA completa",
             "Tabla HLRA de 5 criterios; matriz FMEA con 15 modos de falla "
             "R-001 a R-015 con S/P/D/RPN, mitigacion implementada y riesgo "
             "residual; resumen por nivel",
             "Incluido en este documento"],
            ["D", "Evidencias Objetivas — codigo fuente verificado",
             "EV-D-001: verifyToken; EV-D-002: requireRole con ROLE_GROUPS; "
             "EV-D-003: auditMiddleware; EV-D-004: cadena de middlewares app.js; "
             "EV-D-005: publicPaths.js completo",
             "Incluido en este documento"],
            ["E", "Registros de Capacitacion y Entrenamiento",
             "Plan de entrenamiento por perfil (4 perfiles); formularios de "
             "lista de asistencia listos para firma: solicitante, aprobador, "
             "finanzas, TI administrador",
             "Incluido en este documento"],
        ],
        widths=[1.2, 3.8, 10, 2.5], last_estado=False)

    doc.add_paragraph()
    fin = doc.add_paragraph()
    fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fin.add_run(
        "— Fin del Protocolo Maestro / Inicio de Anexos —"
    )
    r.bold = True; r.font.size = Pt(9)
    r.font.name = "Calibri"; r.font.color.rgb = GRIS
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  ANEXO A — URS y FRS
# ═══════════════════════════════════════════════════════════════════════════════

def _portada_anexo(doc, letra, titulo):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"ANEXO {letra} — {titulo}")
    r.bold = True; r.font.size = Pt(16)
    r.font.name = "Calibri"; r.font.color.rgb = AZUL
    doc.add_paragraph()


def anexo_a(doc):
    _portada_anexo(doc, "A", "ESPECIFICACION DE REQUERIMIENTOS (URS) Y ESPECIFICACION FUNCIONAL (FS/FRS)")
    _p(doc, f"Fuente: docs/validation/URS/ y docs/validation/FRS/ — Extraccion: {FECHA}", italic=True, size=9)

    _label(doc, "Modulo: Autenticacion y Sesiones  |  URS_modulo_autenticacion_sesiones.md")
    _table(doc,
        ["REQ-ID", "Actor", "Requerimiento", "Resultado esperado", "Criticidad"],
        [
            ["REQ-AUTH-001", "Colaborador interno", "Login via Google OAuth2 con cuenta corporativa", "Usuario autenticado redirigido al dashboard segun su rol", "Alta"],
            ["REQ-AUTH-002", "Sistema", "Validar que el dominio del correo pertenezca al dominio autorizado", "Cuentas de dominio externo rechazadas con HTTP 403", "Alta"],
            ["REQ-AUTH-003", "Sistema", "Crear usuario automaticamente en primer login", "Usuario creado con rol inicial y datos de identidad", "Alta"],
            ["REQ-AUTH-005", "Colaborador interno", "Emitir accessToken y refreshToken tras login exitoso", "Frontend recibe ambos tokens y habilita sesion autenticada", "Alta"],
            ["REQ-AUTH-006", "Colaborador interno", "Renovar sesion sin nuevo login mientras el refresh token sea valido", "Nuevos tokens emitidos; user_sessions actualizado", "Media"],
            ["REQ-AUTH-007", "Colaborador interno", "Cerrar sesion de forma explicita", "Sesion invalidada; token anterior rechazado con HTTP 401", "Media"],
            ["REQ-AUTH-010", "TI / Gerencia", "Consultar historial de sesiones y usuarios activos", "Solo roles autorizados acceden; datos correctos retornados", "Alta"],
            ["REQ-AUTH-012", "Sistema", "Detectar accesos fuera de horario laboral y notificar", "Evento registrado; notificacion emitida al canal de TI", "Media"],
        ],
        widths=[2.2, 2.8, 6, 4.5, 1.8])

    _label(doc, "Modulo: Talento Humano  |  URS_modulo_talento_humano.md")
    _table(doc,
        ["REQ-ID", "Actor", "Requerimiento", "Resultado esperado", "Criticidad"],
        [
            ["REQ-TH-001", "Talento Humano", "Administrar perfiles de colaboradores con trazabilidad", "Perfiles actualizados con historial de cambios auditable", "Alta"],
            ["REQ-TH-002", "Colaborador y aprobadores", "Gestionar solicitudes de permisos y vacaciones por flujo jerarquico", "Solicitud avanza por estados validos con trazabilidad completa", "Alta"],
            ["REQ-TH-003", "TH / Gerencia", "Gestionar solicitudes de personal con expediente completo", "Expediente de vacante completo y consistente", "Media"],
            ["REQ-TH-004", "Colaborador autenticado", "Registrar marcaciones de asistencia y overtime", "Registros consolidados para reporte y control de nomina", "Media"],
            ["REQ-TH-005", "Talento Humano", "Administrar departamentos organizacionales", "Estructura departamental vigente para otros modulos", "Baja"],
            ["REQ-TH-006", "Sistema", "Emitir links de verificacion legal para firmas", "Verificaciones legales accesibles por token publico sin JWT", "Alta"],
        ],
        widths=[2.2, 2.8, 6, 4.5, 1.8])

    _label(doc, "Modulo: Finanzas y Viaticos  |  URS_modulo_finanzas_viaticos.md")
    _table(doc,
        ["REQ-ID", "Actor", "Requerimiento", "Resultado esperado", "Criticidad"],
        [
            ["REQ-VT-001", "Solicitante", "Crear workspace de viatico vinculado a salida operacional", "Workspace creado con estado borrador; vinculado al colaborador", "Alta"],
            ["REQ-VT-003", "Solicitante", "Cargar facturas SRI desde archivo TXT de 14 columnas", "Facturas parseadas, validadas y asociadas al workspace", "Alta"],
            ["REQ-VT-006", "Solicitante", "Asistente guiado de 4 pasos para completar el viatico", "Wizard guia carga de soportes, categorias, notas y envio", "Alta"],
            ["REQ-VT-008", "Sistema", "Gestionar 9 estados del flujo de aprobacion de viaticos", "Cada transicion valida actor, permisos y estado previo", "Alta"],
            ["REQ-VT-012", "Finanzas", "Categorizar facturas por tipo de gasto", "Cada factura categorizada para reportes y exportacion ATS", "Alta"],
            ["REQ-VT-016", "Sistema", "Registrar trazabilidad completa de cada cambio de estado", "Trail auditable en auditoria.logs con actor y timestamp", "Alta"],
        ],
        widths=[2.2, 2.8, 6, 4.5, 1.8])

    _label(doc, "Estado documental URS — todos los modulos", "37490D")
    _table(doc,
        ["Modulo", "N REQ", "Archivo fuente", "Estado"],
        [
            ["Autenticacion y Sesiones", "12", "URS_modulo_autenticacion_sesiones.md", "Completo"],
            ["Usuarios y Perfiles", "13", "URS_modulo_usuarios_perfiles.md", "Completo"],
            ["Talento Humano", "6", "URS_modulo_talento_humano.md", "Completo"],
            ["Comercial y Clientes", "18", "URS_modulo_comercial_clientes.md", "v2.0 (2026-06)"],
            ["Business Case", "5+", "URS_modulo_business_case.md", "En revision"],
            ["Finanzas y Viaticos", "16", "URS_modulo_finanzas_viaticos.md", "v2.0 (2026-06)"],
            ["Servicio Tecnico", "11", "URS_modulo_servicio_tecnico_mantenimientos.md", "v2.0 (2026-06)"],
            ["Documentos y Firma", "11", "URS_modulo_documentos_firma.md", "v2.0 (2026-06)"],
            ["Notificaciones", "9+", "URS_modulo_notificaciones_comunicaciones.md", "Completo"],
            ["Reportes y Auditoria", "8", "URS_modulo_reportes_auditoria.md", "v2.0 (2026-06)"],
            ["TI Soporte y Tickets", "11", "URS_modulo_ti_soporte_tickets.md", "v2.0 (2026-06)"],
            ["Inventario y Equipos", "19", "URS_modulo_inventario_equipos.md", "v2.0 (2026-06)"],
        ],
        widths=[5, 1.8, 8, 2.5], last_estado=False)

    _h(doc, "A.2 Especificaciones Funcionales (FS/FRS) criticas", level=2)
    _label(doc, "Autenticacion — FRS verificados en produccion")
    _table(doc,
        ["FRS-ID", "Endpoint", "Entradas", "Proceso ejecutado", "Salida HTTP"],
        [
            ["FRS-AUTH-001", "GET /api/v1/auth/google", "—",
             "Redirige a Google OAuth consent screen con scope openid+email+profile", "302 hacia Google"],
            ["FRS-AUTH-002", "GET /api/v1/auth/google/callback", "code, state (OAuth params)",
             "Intercambia code por tokens Google; valida dominio; upsert usuario; emite JWT accessToken + refreshToken",
             "200 + {accessToken, refreshToken, user}"],
            ["FRS-AUTH-003", "POST /api/v1/auth/refresh", "Header x-refresh-token",
             "Verifica refreshToken; emite nuevo accessToken; actualiza user_sessions", "200 + {accessToken}"],
            ["FRS-AUTH-004", "POST /api/v1/auth/logout", "Bearer accessToken",
             "Invalida sesion activa en user_sessions; registra en auditoria.logs", "200 + {ok: true}"],
            ["FRS-AUTH-005", "GET /api/v1/auth/me", "Bearer accessToken",
             "Decodifica JWT; retorna perfil del usuario con rol, scope e id", "200 + {user}"],
        ],
        widths=[2, 4, 3, 5.5, 3])

    _label(doc, "Talento Humano — FRS verificados en produccion")
    _table(doc,
        ["FRS-ID", "Endpoint", "Roles permitidos", "Proceso ejecutado", "Salida HTTP"],
        [
            ["FRS-TH-001", "POST /api/v1/talento-humano/permisos", "Todos los roles internos",
             "Valida campos obligatorios; crea registro con estado pendiente; notifica aprobador",
             "201 + {id, estado: pendiente}"],
            ["FRS-TH-002", "GET /api/v1/talento-humano/permisos", "talento_humano, gerencia, jefaturas",
             "Consulta solicitudes con filtros de estado, fecha, usuario; pagina resultados",
             "200 + {items, total}"],
            ["FRS-TH-003", "PUT /api/v1/talento-humano/permisos/:id/estado", "Aprobadores segun area",
             "Valida actor autorizado; verifica estado actual; actualiza estado; registra en auditoria.logs",
             "200 + {id, estado_nuevo}"],
            ["FRS-TH-004", "GET /api/v1/talento-humano/permisos/:id", "Solicitante, jefaturas, TH",
             "Retorna detalle completo con historial de estados y actor de cada transicion",
             "200 + {permiso, historial}"],
        ],
        widths=[2, 4.5, 3.5, 5.5, 2])
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  ANEXO B — DDS
# ═══════════════════════════════════════════════════════════════════════════════

def anexo_b(doc):
    _portada_anexo(doc, "B", "ESPECIFICACION DE DISENO DEL SISTEMA (DDS)")
    _p(doc, f"Fuente: backend/src/app.js, routes/registerRoutes.js, routes/publicPaths.js, config/db.js — Extraccion: {FECHA}", italic=True, size=9)

    _h(doc, "B.1 Arquitectura del sistema", level=2)
    _table(doc,
        ["Capa", "Tecnologia", "Version", "Archivo de config", "Responsabilidad"],
        [
            ["Presentacion (SPA)", "React 19 + Tailwind + Bootstrap", "19.x",
             "spi_front/vite.config.js", "Interfaces por rol; routing React Router; state local"],
            ["API / Logica de negocio", "Node.js + Express.js", "5.1.0",
             "backend/src/app.js", "Rutas REST; middlewares en cadena; controladores; servicios"],
            ["Persistencia", "PostgreSQL — Neon serverless", "pg 8.x",
             "backend/src/config/db.js", "Pool 20 conexiones; transacciones ACID; retry en errores transitorios"],
            ["Autenticacion", "Google OAuth 2.0 + JWT", "google-auth-library 10.x",
             "backend/src/middlewares/auth.js", "Sessions con accessToken corto; claims iss/aud/sub validados"],
            ["Seguridad HTTP", "Helmet + CORS + Rate Limit", "helmet 8.x",
             "backend/src/config/security.js", "CSP, HSTS, X-Frame; CORS dinamico; 3000 req/15min produccion"],
            ["Archivos en nube", "Google Drive API", "googleapis 164.x",
             "backend/src/modules/documents/", "Documentos, adjuntos y actas via service account"],
            ["Proceso gestor", "PM2", "Latest",
             "ecosystem.config.js", "Reinicio automatico; logs de proceso; instancias separadas"],
        ],
        widths=[2.8, 3.5, 2, 4.5, 4.7], last_estado=False)

    _h(doc, "B.2 Cadena de middleware en app.js (orden de ejecucion verificado)", level=2)
    _code(doc, [
        "// backend/src/app.js — cadena completa verificada",
        "app.use(helmet(helmetConfig))           // 1. Seguridad HTTP (CSP, HSTS, X-Frame)",
        "app.use(rateLimit({...}))               // 2. Rate limiting (3000 req/15min produccion)",
        "app.use(cors(corsConfig))               // 3. CORS dinamico por lista de origenes",
        "app.use(express.json({ limit:'5mb' })) // 4. Parsing JSON del body",
        "app.use(requestContextMiddleware)       // 5. Correlacion de requests",
        "app.use(mLogger)                        // 6. Logging estructurado",
        "",
        "mountPublicRoutes(app)                  // 7. Rutas publicas (auth, webhook)",
        "",
        "app.use((req, res, next) => {           // 8. Guardia JWT global:",
        "  if (isPublicPath(req.path)) return next()",
        "  return verifyToken(req, res, next)    //    HTTP 401 sin token valido",
        "})",
        "",
        "app.use(normalizeApiPayloads)           //  9. Normalizacion",
        "app.use(moduleAccessGuard)              // 10. Control de acceso a modulos",
        "app.use(auditMiddleware)                // 11. AUDITORIA de mutaciones",
        "mountPrivateRoutes(app)                 // 12. Logica de negocio",
    ])

    _h(doc, "B.3 Rutas publicas (publicPaths.js) — sin JWT requerido", level=2)
    _table(doc,
        ["Ruta / Prefijo", "Justificacion de excepcion"],
        [
            ["/api/v1/auth/google*", "Inicio del flujo OAuth2 — usuario no tiene token todavia"],
            ["/api/v1/auth/google/callback", "Callback de Google — procesa el code y emite el JWT"],
            ["/api/v1/permisos/legal-verification/*", "Verificacion publica de firma de permiso por QR"],
            ["/api/v1/vacaciones/legal-verification/*", "Verificacion publica de firma de vacaciones por QR"],
            ["/api/verificar, /api/verify, /api/signature/verificar*", "Verificacion publica de documentos firmados por QR"],
            ["/api/v1/integrations/crm/webhook", "Webhook de EspoCRM validado por X-Hook-Secret (no JWT)"],
            ["/health", "Health check para el proceso gestor PM2"],
            ["/ws", "WebSocket (handshake inicial sin autenticacion)"],
        ],
        widths=[6.5, 11], last_estado=False)

    _h(doc, "B.4 Configuracion del pool de base de datos (db.js)", level=2)
    _table(doc,
        ["Parametro", "Valor", "Evidencia en codigo"],
        [
            ["Motor de base de datos", "PostgreSQL via Neon serverless (cloud)", "const { Pool } = require('pg')"],
            ["Pool maximo de conexiones", "20 conexiones simultaneas", "max: intFromEnv('DB_POOL_MAX', 20)"],
            ["Pool minimo activo", "2 conexiones calientes permanentes", "min: intFromEnv('DB_POOL_MIN', 2)"],
            ["Timeout conexion inactiva", "30,000 ms", "idleTimeoutMillis: intFromEnv('DB_IDLE_TIMEOUT_MS', 30000)"],
            ["Timeout para establecer conexion", "5,000 ms", "connectionTimeoutMillis: intFromEnv('DB_CONN_TIMEOUT_MS', 5000)"],
            ["Max usos por conexion", "7,500 queries antes de reciclar", "maxUses: intFromEnv('DB_POOL_MAX_USES', 7500)"],
            ["SSL", "Configurable via DB_SSL=true", "ssl: process.env.DB_SSL === 'true' ? { rejectUnauthorized: ... }"],
            ["Retry automatico", "2 intentos con 250ms entre cada uno", "maxAttempts = intFromEnv('DB_QUERY_RETRY_ATTEMPTS', 2)"],
            ["Errores con retry", "ETIMEDOUT, ECONNRESET, 57P01, 57P03, 53300", "isTransientDbError() en db.js"],
            ["Shutdown graceful", "Cierra el pool en SIGINT y SIGTERM", "setupGracefulShutdown() — pool.end()"],
        ],
        widths=[5, 4, 8.5], last_estado=False)
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  ANEXO C — FMEA
# ═══════════════════════════════════════════════════════════════════════════════

def anexo_c(doc):
    _portada_anexo(doc, "C", "MATRIZ HLRA / FRA COMPLETA — EVALUACION DE RIESGOS FMEA")
    _p(doc, "Metodologia: RPN = Severidad x Probabilidad x Detectabilidad (escala 1-5). Fuente: general/07A_evaluacion_riesgos_fmea.md", italic=True, size=9)

    _h(doc, "C.1 Escala de evaluacion", level=2)
    _table(doc,
        ["Valor", "Severidad", "Probabilidad", "Detectabilidad"],
        [
            ["1", "Insignificante", "Muy baja (menos de 1% de los ciclos)", "Detectado inmediatamente — control automatico"],
            ["2", "Menor", "Baja (1-5% de los ciclos)", "Detectado rapidamente en la misma sesion"],
            ["3", "Moderado", "Media (5-20% de los ciclos)", "Detectado en revision periodica"],
            ["4", "Mayor", "Alta (20-50% de los ciclos)", "Dificil de detectar sin auditoria especifica"],
            ["5", "Critico / Catastrofico", "Muy alta (mas de 50%)", "Muy dificil de detectar — sin control automatico"],
        ],
        widths=[1.5, 3.5, 4.5, 8], last_estado=False)

    _h(doc, "C.2 Matriz FMEA — 15 modos de falla", level=2)
    _table(doc,
        ["ID", "Modulo", "Modo de falla", "S", "P", "D", "RPN", "Nivel", "Mitigacion", "Estado"],
        [
            ["R-001", "Infraestructura", "Interrupcion del proceso — indisponibilidad total", "5", "2", "2", "20", "Medio",
             "PM2 con reinicio automatico; watchdog de proceso", "Mitigado"],
            ["R-002", "auth", "Compromiso de credenciales OAuth2 / replay de token JWT", "5", "2", "3", "30", "ALTO",
             "JWT corta duracion; claims iss/aud/sub; dominio restringido; refresh rotation", "ACTIVO — OQ-014"],
            ["R-003", "auth", "Sesion no invalidada tras logout — acceso residual", "4", "2", "3", "24", "Medio",
             "user_sessions invalidado en logout; accessToken de corta duracion", "Mitigado"],
            ["R-004", "security", "Acceso fuera de horario no detectado ni notificado", "4", "3", "2", "24", "Medio",
             "Modulo security detecta off-hours; notificacion automatica a TI", "Mitigado"],
            ["R-005", "auditoria", "Perdida de registros de audit trail", "5", "1", "2", "10", "Bajo",
             "auditoria.logs en PostgreSQL; usuarios no pueden eliminar; exportacion protegida", "Mitigado"],
            ["R-006", "Base de datos", "Corrupcion o perdida total de datos", "5", "1", "3", "15", "Bajo",
             "Transacciones ACID; retry en errores transitorios; backups Neon", "Mitigado"],
            ["R-007", "signature", "Firma avanzada invalidada o no generada", "5", "2", "2", "20", "Medio",
             "Hash SHA-256; QR de verificacion publica; cadena inmutable en audit log", "Mitigado"],
            ["R-008", "documents", "Documento generado corrupto (PDF/ZIP)", "3", "2", "2", "12", "Bajo",
             "Validacion post-generacion; reintentos automaticos; registro del resultado", "Mitigado"],
            ["R-009", "Usuarios RBAC", "Asignacion incorrecta de rol — acceso no autorizado", "4", "2", "3", "24", "Medio",
             "requireRole() en cada endpoint; SUPER_ROLES solo admin; guards assertX()", "Mitigado"],
            ["R-010", "Gmail/SMTP", "Falla en envio de notificacion critica", "2", "3", "2", "12", "Bajo",
             "Cola de dispatch con reintentos; canal in-app siempre disponible", "Mitigado"],
            ["R-011", "permisos/vacaciones", "Calculo incorrecto de saldo — impacto en nomina", "4", "2", "3", "24", "Medio",
             "Validacion unitaria; pruebas OQ especificas; historial auditable de transiciones", "Bajo seguimiento"],
            ["R-012", "attendance", "Marcacion duplicada en asistencia", "3", "2", "3", "18", "Bajo",
             "Idempotencia en API; validacion de timestamp unico por colaborador", "Mitigado"],
            ["R-013", "personnel-requests", "Expediente de candidato incompleto", "3", "3", "3", "27", "Medio",
             "Validacion de campos obligatorios; checklist de documentos por estado", "Mitigado"],
            ["R-014", "Entorno produccion", "Despliegue en entorno incorrecto", "5", "1", "2", "10", "Bajo",
             "Separacion por NODE_ENV; variables por ambiente; IQ verifica entorno", "Mitigado"],
            ["R-015", "LOPDP", "Exposicion de datos personales sensibles", "5", "1", "3", "15", "Bajo",
             "RBAC; HTTPS/TLS; aceptacion LOPDP en primer login; datos sensibles no en logs", "Mitigado"],
        ],
        widths=[1.0, 2.2, 4, 0.5, 0.5, 0.5, 0.8, 1.2, 5.5, 1.8], font=7)
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  ANEXO D — EVIDENCIAS OBJETIVAS
# ═══════════════════════════════════════════════════════════════════════════════

def anexo_d(doc):
    _portada_anexo(doc, "D", "EVIDENCIAS OBJETIVAS — CODIGO FUENTE VERIFICADO")
    _p(doc, f"Extractos reales del codigo fuente de FamSPI v1.0.0. Evidencia de controles validados. Extraccion: {FECHA}", italic=True, size=9)

    _h(doc, "EV-D-001 — verifyToken (backend/src/middlewares/auth.js)", level=2)
    _p(doc, "Evidencia de REQ-AUTH-001 a REQ-AUTH-007 y mitigacion de R-002 (RPN=30).", size=9)
    _table(doc, ["Aspecto verificado", "Implementacion confirmada"], [
        ["Token requerido en cabecera", "Verifica Authorization, authorization o x-access-token; HTTP 401 si ausente"],
        ["Soporte formato Bearer", "Extrae token correctamente de Bearer <token>"],
        ["Verificacion criptografica", "jwt.verify(token, process.env.SECRET_KEY); HTTP 401 si invalido o expirado"],
        ["Validacion de claims", "Verifica iss=spi-fam-backend, aud=spi-fam-frontend y sub; HTTP 403 si falla"],
        ["Propagacion de identidad", "req.user = { ...decoded, ip, userAgent } disponible para middlewares siguientes"],
    ], widths=[5, 12.5], last_estado=False)
    _code(doc, [
        "// backend/src/middlewares/auth.js",
        "const verifyToken = (req, res, next) => {",
        "  const headerAuth = req.headers['authorization'] || req.headers['x-access-token'];",
        "  if (!headerAuth) return res.status(401).json({ ok: false, code: 'NO_TOKEN' });",
        "  const token = headerAuth.startsWith('Bearer ') ? headerAuth.split(' ')[1] : headerAuth;",
        "  let decoded;",
        "  try { decoded = jwt.verify(token, process.env.SECRET_KEY); }",
        "  catch (err) { return res.status(401).json({ ok: false, code: 'INVALID_TOKEN' }); }",
        "  if (decoded.iss !== 'spi-fam-backend' ||",
        "      decoded.aud !== 'spi-fam-frontend' || !decoded.sub) {",
        "    return res.status(403).json({ ok: false, code: 'INVALID_CLAIMS' });",
        "  }",
        "  req.user = { ...decoded, ip: req.headers['x-forwarded-for']?.split(',')[0] || req.ip };",
        "  next();",
        "};",
    ])

    _h(doc, "EV-D-002 — requireRole (backend/src/middlewares/roles.js)", level=2)
    _p(doc, "Evidencia del control RBAC — mitigacion R-009 (RPN=24).", size=9)
    _code(doc, [
        "const ROLE_GROUPS = {",
        "  ti:       ['ti', 'jefe_ti', 'desarrollador', 'soporte'],",
        "  finanzas: ['finanzas', 'financiero', 'jefe_finanzas', 'contador'],",
        "  gerencia: ['gerencia', 'gerencia_general', 'director', 'gerente'],",
        "  // ... 12 grupos adicionales",
        "};",
        "const SUPER_ROLES = new Set(['admin', 'administrador']);",
        "",
        "function requireRole(allowedRoles = []) {",
        "  const expanded = expandRoles(allowedRoles);",
        "  return (req, res, next) => {",
        "    if (!req.user) return res.status(401).json({ ok: false, error: 'No autenticado.' });",
        "    const candidates = collectUserRoles(req.user);",
        "    for (const role of candidates) {",
        "      if (SUPER_ROLES.has(role)) return next();",
        "      if (expanded.has(role)) return next();",
        "    }",
        "    return res.status(403).json({ ok: false,",
        "      error: `Acceso denegado. Roles: ${[...expanded].join(', ')}` });",
        "  };",
        "}",
    ])

    _h(doc, "EV-D-003 — auditMiddleware (backend/src/middlewares/auditMiddleware.js)", level=2)
    _p(doc, "Evidencia del Audit Trail — REQ-AUD-001, mitigacion R-005 (RPN=10).", size=9)
    _code(doc, [
        "async function auditMiddleware(req, res, next) {",
        "  const method = req.method.toUpperCase();",
        "  if (!['POST','PUT','PATCH','DELETE'].includes(method)) return next();",
        "  if (req.originalUrl.startsWith('/api/v1/auth')) return next();",
        "  const start = Date.now();",
        "  res.on('finish', async () => {",
        "    const modulo = resolveModuleFromPath(pathParts);",
        "    const accion = mapHttpMethodToAction(method, pathParts);",
        "    await logAction({",
        "      usuario_email, rol, modulo,",
        "      accion: success ? accion : `${accion}_failed`,",
        "      datos_nuevos: truncateBody(req.body),",
        "      ip, duracion_ms: Date.now() - start,",
        "    });",
        "  });",
        "  next();",
        "}",
    ])

    _h(doc, "EV-D-004 — Guardia JWT global (backend/src/app.js)", level=2)
    _code(doc, [
        "// JWT global — aplica a TODO lo que no este en publicPaths",
        "app.use((req, res, next) => {",
        "  if (isPublicPath(req.path)) return next();",
        "  return verifyToken(req, res, next);",
        "});",
        "app.use(auditMiddleware);     // AUDITORIA despues de auth",
        "mountPrivateRoutes(app);",
    ])

    _h(doc, "EV-D-005 — publicPaths.js completo (backend/src/routes/publicPaths.js)", level=2)
    _code(doc, [
        "const PUBLIC_PATH_PREFIXES = [",
        "  '/ws', '/internal/jobs',",
        "  '/api/v1/equipment-purchases/events',",
        "  '/api/v1/private-purchases/events',",
        "  '/api/v1/auth/google',",
        "  '/api/v1/gmail/auth/callback',",
        "  '/api/v1/permisos/legal-verification/',",
        "  '/api/v1/vacaciones/legal-verification/',",
        "  '/health', '/api/verificar', '/api/verify',",
        "  '/api/signature/verificar', '/api/signature/verify',",
        "  '/api/v1/signature/verificar', '/api/v1/signature/verify'",
        "];",
        "// TODAS las demas rutas requieren JWT valido.",
    ])
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  ANEXO E — ENTRENAMIENTO
# ═══════════════════════════════════════════════════════════════════════════════

def anexo_e(doc):
    _portada_anexo(doc, "E", "REGISTROS DE CAPACITACION Y ENTRENAMIENTO")
    _p(doc, (
        "El entrenamiento formal documentado es prerequisito irrenunciable para la ejecucion "
        "de PQ/UAT segun WHO TRS 1019 Annex 3, Sec. 12. Las listas de asistencia firmadas "
        "son la evidencia de cumplimiento de este requisito regulatorio."
    ), size=9)

    _h(doc, "E.1 Plan de entrenamiento por perfil", level=2)
    _table(doc,
        ["Perfil", "Modulos cubiertos", "Duracion estimada", "Estado", "Fecha"],
        [
            ["Solicitante general (comercial, tecnico, RRHH, backoffice)",
             "Dashboard, permisos, vacaciones, workspace de viaticos, notificaciones",
             "2-4 horas", "Pendiente", "_______________"],
            ["Aprobador / Jefe de area",
             "Todo lo del solicitante + flujos de aprobacion y rechazo por area",
             "3-5 horas", "Pendiente", "_______________"],
            ["Finanzas",
             "Viaticos financiero, categorizacion de facturas, exportacion ATS/XML, reportes",
             "4-6 horas", "Pendiente", "_______________"],
            ["TI Administrador",
             "Sistema completo: usuarios, roles, auditoria, soporte, activos, respaldo, monitoreo",
             "8-12 horas (sesiones)", "Parcial", "_______________"],
        ],
        widths=[4.5, 6.5, 2.5, 2, 2], last_estado=False)

    for perfil, modulos, n_rows in [
        ("Solicitante general", "Dashboard + Permisos/Vacaciones + Viaticos + Notificaciones", 8),
        ("Aprobador / Jefe de area", "Flujos de aprobacion y rechazo + historial", 5),
        ("Finanzas", "Viaticos financiero + Reportes + ATS/XML", 3),
        ("TI Administrador", "Sistema completo con acceso admin", 2),
    ]:
        _h(doc, f"Lista de asistencia — {perfil}", level=2)
        _p(doc, f"Modulo: {modulos}", size=9, after=2)
        _p(doc, "Instructor: ___________________________   Fecha: _______________   Duracion: _____ horas", size=9, after=3)
        _table(doc,
            ["#", "Nombre completo", "Cargo / Departamento", "Correo corporativo", "Firma"],
            [[str(i), "", "", "", ""] for i in range(1, n_rows + 1)],
            widths=[0.8, 4.5, 3.8, 5, 3.5], last_estado=False)
        _p(doc, "Firma del instructor: _________________________   Fecha: _______________", size=9)
        doc.add_paragraph()

    doc.add_paragraph()
    fin = doc.add_paragraph()
    fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fin.add_run("— Fin del Protocolo Maestro de Validacion FamSPI v1.0.0 + Anexos A-E —")
    r.bold = True; r.font.size = Pt(10)
    r.font.name = "Calibri"; r.font.color.rgb = AZUL


# ═══════════════════════════════════════════════════════════════════════════════
#  CONSTRUCCION PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════════════

def build(path=None):
    target = path or OUTPUT
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    portada(doc)
    s1_control(doc)
    s2_glosario(doc)
    s3_intro(doc)
    s4_qms(doc)
    s5_specs(doc)
    s6_riesgos(doc)
    s7_plan(doc)
    s8_rtm(doc)
    s9_anexos(doc)
    anexo_a(doc)
    anexo_b(doc)
    anexo_c(doc)
    anexo_d(doc)
    anexo_e(doc)

    doc.save(target)
    print(f"Documento generado: {target}")


if __name__ == "__main__":
    build()
