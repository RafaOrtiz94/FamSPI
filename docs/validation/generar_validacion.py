"""Protocolo Maestro de Validacion FamSPI v1.0.0 — alineado WHO TRS 1019 Annex 3, Appendix 5.

Estructura de 9 secciones:
  1. Control Documental y Aprobaciones
  2. Glosario y Definiciones Tecnicas
  3. Introduccion, Alcance y Estrategia
  4. Especificaciones del Sistema (URS / FRS / DDS)  — narrativa, tablas en Anexos A y B
  5. Analisis de Riesgos (FMEA)                     — resumen criticos + referencia a Anexo C
  6. Plan de Calificacion y Protocolos de Prueba     — 6.1 DQ / 6.2 IQ / 6.3 OQ / 6.4 PQ
  7. Matriz de Trazabilidad (RTM)
  8. Procedimientos Operativos (SOPs) y Entrenamiento
  9. Indice de Anexos
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = (
    r"c:\Users\Departamento de TI\Desktop\PROYECTOS\FamSPI"
    r"\docs\validation\PROTOCOLO_VALIDACION_FAMSPI_V1_0_0.docx"
)
FECHA   = "18/06/2026"
VERSION = "1.0"
SISTEMA = "FamSPI v1.0.0"
DEPTO   = "Departamento de Tecnologia de la Informacion (TI)"
NORMA   = "WHO TRS 1019 Annex 3, Appendix 5"

AZUL    = RGBColor(0x1F, 0x49, 0x7D)
VERDE   = RGBColor(0x37, 0x86, 0x50)
GRIS    = RGBColor(0x80, 0x80, 0x80)


# ── Primitivas ────────────────────────────────────────────────────────────────

def _shd(cell, hex_color):
    tcp = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcp.append(shd)


def _cell(cell, text, bold=False, size=8, color=None, fill=None, center=False):
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if fill:
        _shd(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _widths(table, widths_cm):
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def _estado_fill(v):
    """Retorna hex de fondo segun valor de estado."""
    if v in ("Conforme", "CONFORME"):
        return "E2EFDA"
    if v in ("Pendiente", "PENDIENTE", "—"):
        return "FFF2CC"
    return None


def _table(doc, headers, rows, widths=None, hbg="1F497D", font=8, estado_col=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        _cell(t.rows[0].cells[i], h, bold=True, color="FFFFFF",
              fill=hbg, center=True, size=font)
    for ri, rd in enumerate(rows):
        row = t.add_row()
        base = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        last = len(rd) - 1
        for ci, v in enumerate(rd):
            sf = _estado_fill(v) if (estado_col and ci == last) else None
            bold = ci == last and sf is not None
            _cell(row.cells[ci], v, size=font,
                  fill=sf or base, bold=bold,
                  center=(ci == last and sf is not None))
    if widths:
        _widths(t, widths)
    doc.add_paragraph()
    return t


def _kv(doc, rows, widths=(5, 11)):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        _cell(t.rows[i].cells[0], k, bold=True,
              fill="1F497D", color="FFFFFF", size=9)
        _cell(t.rows[i].cells[1], v, size=9, fill="F8F8F8")
    _widths(t, widths)
    doc.add_paragraph()


def _h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.color.rgb = AZUL if level == 1 else VERDE
    return p


def _p(doc, text, italic=False, size=10, after=5):
    para = doc.add_paragraph()
    para.style = doc.styles["Normal"]
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.italic = italic
    return para


def _bullet(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10)


def _numbered(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10)


# ── Portada ───────────────────────────────────────────────────────────────────

def portada(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Protocolo Maestro de Validacion\nFamSPI v1.0.0")
    r.bold = True; r.font.name = "Calibri"
    r.font.size = Pt(22); r.font.color.rgb = AZUL

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run(f"{DEPTO}  |  {FECHA}")
    rs.italic = True; rs.font.size = Pt(10)
    rs.font.name = "Calibri"; rs.font.color.rgb = GRIS

    _kv(doc, [
        ("Sistema",               SISTEMA),
        ("Marco normativo",       NORMA),
        ("Version del protocolo", VERSION),
        ("Ambito",                "12 modulos / 6 areas de negocio — Ecuador"),
        ("Clasificacion",         "INTERNO — Uso restringido"),
        ("Estado",                "Para aprobacion — liberacion condicionada Areas 01 y 02"),
    ], widths=[5, 11])
    doc.add_page_break()


# ── 1. Control Documental y Aprobaciones ─────────────────────────────────────

def seccion_1(doc):
    _h(doc, "1. Control Documental y Aprobaciones", level=1)

    _kv(doc, [
        ("Titulo",                   "Protocolo Maestro de Validacion de FamSPI v1.0.0"),
        ("Codigo del documento",     "TI-VAL-PROTO-001"),
        ("Version",                  VERSION),
        ("Fecha de elaboracion",     FECHA),
        ("Fecha de revision",        "_______________"),
        ("Fecha de aprobacion",      "_______________"),
        ("Alcance resumido",         "12 modulos operativos / 6 areas / 120+ endpoints — FamSPI v1.0.0"),
        ("Departamento responsable", DEPTO),
        ("Marco normativo",          NORMA),
        ("Referencias adicionales",  "21 CFR Part 11 (referencia informativa); LOPDP Ecuador"),
        ("Estado",                   "Para aprobacion — liberacion condicionada a Areas 01 y 02"),
        ("Localizacion",             "docs/validation/ — repositorio FamSPI"),
    ], widths=[4.5, 12])

    _h(doc, "Historial de revisiones", level=2)
    _table(doc,
        ["Version", "Fecha", "Autor", "Descripcion del cambio"],
        [["1.0", FECHA, "Departamento de TI",
          "Version inicial — ciclo completo DQ/IQ/OQ/PQ alineado WHO TRS 1019"]],
        widths=[1.5, 2.5, 4, 9], estado_col=False)

    _h(doc, "Firmas de aprobacion", level=2)
    t = doc.add_table(rows=3, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(["Realizado por", "Revisado por", "Aprobado por"]):
        _cell(t.rows[0].cells[ci], h, bold=True,
              fill="1F497D", color="FFFFFF", center=True, size=9)
    for ri in range(1, 3):
        lbl = ["Nombre:", "Firma / Fecha:"][ri - 1]
        for ci in range(3):
            _cell(t.rows[ri].cells[ci],
                  lbl + "  ________________________", size=9, fill="FAFAFA")
    _widths(t, [5.5, 5.5, 5.5])
    doc.add_paragraph()
    doc.add_page_break()


# ── 2. Glosario y Definiciones Tecnicas ──────────────────────────────────────

def seccion_2(doc):
    _h(doc, "2. Glosario y Definiciones Tecnicas", level=1)
    _p(doc, (
        "Este glosario estandariza los terminos regulatorios de la Organizacion Mundial "
        "de la Salud (OMS) y los componentes tecnicos del sistema FamSPI v1.0.0. "
        "Su proposito es garantizar una interpretacion uniforme de las pruebas, criterios "
        "de aceptacion y evidencias documentadas en este protocolo, evitando ambiguedades "
        "durante la revision por parte de auditores internos o externos."
    ))

    _h(doc, "2.1 Terminos regulatorios (OMS / WHO TRS 1019)", level=2)
    _table(doc,
        ["Termino", "Definicion"],
        [
            ["Audit trail (Bitacora de auditoria)",
             "Registro electronico seguro, cronologico y trazable de todas las acciones "
             "criticas realizadas en el sistema (creaciones, modificaciones, eliminaciones). "
             "En FamSPI se implementa en la tabla auditoria.logs con campos: usuario_id, "
             "modulo, accion, ip, duracion_ms y timestamp inmutable."],
            ["Data integrity (Integridad de datos)",
             "Garantia de que los datos son completos, consistentes, precisos y no han "
             "sido alterados de manera no autorizada durante todo su ciclo de vida "
             "(creacion, procesamiento, almacenamiento, recuperacion y archivo)."],
            ["Metadata",
             "Datos que describen el contexto, el contenido y la estructura de los datos "
             "principales. Incluye timestamps de creacion/modificacion, identificadores de "
             "actor y estado de los registros en FamSPI."],
            ["Archiving (Archivo)",
             "Proceso de trasladar datos que ya no estan en uso activo a un almacenamiento "
             "de largo plazo, garantizando su recuperacion e integridad futura. "
             "Aplica a bitacoras de auditoria y documentos finalizados en FamSPI."],
            ["Backup (Copia de seguridad)",
             "Copia redundante de los datos del sistema realizada con periodicidad definida "
             "para garantizar la recuperacion ante fallos. En FamSPI provisto por el "
             "plan de respaldo de Neon PostgreSQL (cloud)."],
            ["Validation (Validacion)",
             "Proceso documentado que establece mediante evidencia objetiva que los "
             "requerimientos para un uso especifico previsto se han cumplido de forma "
             "consistente (WHO TRS 1019, Sec. 3.1)."],
            ["Qualification (Calificacion)",
             "Accion de demostrar y documentar que los equipos, sistemas o instalaciones "
             "estan correctamente instalados y/u operan conforme a especificacion. "
             "Incluye DQ, IQ, OQ y PQ."],
        ],
        widths=[4.5, 13], estado_col=False)

    _h(doc, "2.2 Terminos del sistema FamSPI", level=2)
    _table(doc,
        ["Termino", "Definicion"],
        [
            ["API REST",
             "Interfaz de programacion de aplicaciones que utiliza el protocolo HTTP "
             "con verbos estandar (GET, POST, PUT, PATCH, DELETE) para exponer "
             "los recursos del backend de FamSPI a los clientes autorizados."],
            ["JWT (JSON Web Token)",
             "Estandar abierto (RFC 7519) para transmitir informacion de autenticacion "
             "de forma segura entre partes como objeto JSON firmado criptograficamente. "
             "En FamSPI: accessToken (corta duracion) y refreshToken, con claims "
             "iss=spi-fam-backend, aud=spi-fam-frontend, sub=usuario_id."],
            ["Token de acceso (accessToken)",
             "Credencial de corta duracion emitida tras autenticacion exitosa. "
             "Debe enviarse en la cabecera Authorization: Bearer <token> en cada "
             "solicitud a endpoints privados de FamSPI."],
            ["OAuth2 (Google)",
             "Protocolo de autorizacion delegada utilizado para autenticar a los usuarios "
             "de FamSPI mediante sus cuentas Google corporativas. El sistema valida "
             "que el dominio del correo pertenezca al dominio corporativo autorizado."],
            ["RBAC (Role-Based Access Control)",
             "Modelo de control de acceso en el que los permisos se asignan a roles "
             "y los usuarios heredan los permisos del rol asignado. En FamSPI implementado "
             "via requireRole() con 15 grupos de roles definidos en roles.js."],
            ["Endpoint",
             "URL especifica de la API REST de FamSPI que acepta solicitudes HTTP. "
             "Ejemplo: POST /api/v1/talento-humano/permisos para crear una solicitud "
             "de permiso. Los endpoints privados requieren JWT valido."],
            ["PostgreSQL / Neon",
             "Sistema de gestion de bases de datos relacional (RDBMS) utilizado por "
             "FamSPI a traves del proveedor serverless Neon. Pool de 20 conexiones "
             "maximo; SSL habilitado; transacciones ACID garantizadas."],
            ["PM2",
             "Gestor de procesos para Node.js que mantiene activa la aplicacion FamSPI "
             "en produccion con reinicio automatico ante fallos y separacion de instancias."],
            ["Express.js",
             "Framework de servidor HTTP para Node.js. FamSPI usa la version 5.1.0 "
             "con enrutamiento modular y cadena de middlewares en orden definido."],
        ],
        widths=[4.5, 13], estado_col=False)

    _h(doc, "2.3 Terminos de negocio", level=2)
    _table(doc,
        ["Termino", "Definicion"],
        [
            ["Permiso",
             "Solicitud formal de ausencia laboral justificada de corto plazo gestionada "
             "mediante el flujo de aprobacion jerarquica en FamSPI: creacion → pendiente → "
             "aprobado / rechazado, con link de verificacion legal generado por QR."],
            ["Vacaciones",
             "Solicitud formal de dias de descanso remunerado gestionada en FamSPI con "
             "el mismo flujo de aprobacion que los permisos, incluyendo verificacion legal "
             "y descuento del saldo de dias disponibles del colaborador."],
            ["Saldo de vacaciones",
             "Cantidad de dias de vacaciones acumulados y disponibles para un colaborador, "
             "calculado por el sistema en funcion de la antiguedad y los dias ya utilizados. "
             "Su integridad es critica para nomina y control laboral (FMEA R-011)."],
            ["Viatico",
             "Anticipo o reembolso de gastos de viaje gestionado mediante un workspace "
             "en FamSPI con asistente de 4 pasos, carga de facturas electronicas SRI "
             "formato TXT 14 columnas y flujo de 9 estados de aprobacion."],
            ["Business Case (BC)",
             "Documento de analisis de oportunidad comercial gestionado en FamSPI "
             "vinculado a un cliente y oportunidad, con campos, estados y flujo "
             "de revision definidos por el area comercial."],
            ["Acta de calificacion",
             "Documento formal generado por FamSPI (PDF) que certifica la entrega, "
             "asignacion o baja de un activo TI, firmado con hash SHA-256 y "
             "verificable mediante QR publico."],
        ],
        widths=[4.5, 13], estado_col=False)

    doc.add_page_break()


# ── 3. Introduccion, Alcance y Estrategia ─────────────────────────────────────

def seccion_3(doc):
    _h(doc, "3. Introduccion, Alcance y Estrategia de Validacion", level=1)

    _p(doc, (
        "El objetivo de este protocolo es establecer la planificacion, ejecucion y "
        "registro de la validacion del sistema computarizado FamSPI v1.0.0, desarrollado "
        "internamente por el Departamento de TI de FAM Ecuador. La validacion aplica un "
        "enfoque basado en riesgo (RPN = Severidad x Probabilidad x Detectabilidad) para "
        "priorizar los casos de prueba segun el impacto potencial de cada modo de falla, "
        "siguiendo el ciclo de vida de validacion definido en WHO TRS 1019 Annex 3, "
        "Appendix 5: DQ → IQ → OQ → PQ."
    ))

    _h(doc, "3.1 Modulos y areas en alcance", level=2)
    _table(doc,
        ["Area", "Modulos incluidos", "Estado"],
        [
            ["Area 01 — Gobierno, Seguridad y Acceso",
             "Autenticacion Google OAuth2, gestion de usuarios y roles (RBAC), "
             "bitacora de auditoria, deteccion de accesos fuera de horario (security)",
             "DQ/IQ/OQ/PQ — Conforme"],
            ["Area 02 — Personas y Talento Humano",
             "Colaboradores y perfiles, permisos laborales, vacaciones, "
             "solicitudes de personal, asistencia, departamentos",
             "DQ/IQ/OQ/PQ — Conforme"],
            ["Area 03 — Comercial y Business Case",
             "Clientes, oportunidades comerciales, business cases, "
             "entregas de colaboradores, integracion CRM",
             "DQ/IQ conforme — OQ/PQ pendientes"],
            ["Area 04 — Servicio Tecnico y Operaciones",
             "Mantenimientos preventivos y correctivos, capacitaciones, "
             "proyectos externos GoApp, certificaciones de personal",
             "DQ/IQ conforme — OQ/PQ pendientes"],
            ["Area 05 — Compras, Inventario y Logistica",
             "Activos TI (alta, asignacion, baja), tickets de soporte, "
             "gestion de equipos, compras privadas",
             "DQ/IQ conforme — OQ/PQ pendientes"],
            ["Area 06 — Finanzas y Transacciones",
             "Viaticos (wizard 4 pasos, 9 estados), firma digital de documentos "
             "con SHA-256, notificaciones, dashboard de KPIs, reportes y auditoria",
             "DQ/IQ conforme — OQ/PQ pendientes"],
        ],
        widths=[3.5, 8, 3.5], estado_col=False)

    _h(doc, "3.2 Exclusiones del alcance", level=2)
    _p(doc, "Los siguientes elementos quedan explicitamente excluidos de esta validacion v1.0.0:")
    _bullet(doc, [
        "Infraestructura fisica del servidor: la validacion no cubre la calificacion del hardware, "
        "centros de datos ni conectividad de red provista por terceros.",
        "Servicios de terceros (Google OAuth, Neon, Gmail API): se consideran proveedores "
        "calificados; la validacion verifica la integracion pero no el servicio en si.",
        "Modulos en desarrollo activo no desplegados en produccion al corte v1.0.0.",
        "Funcionalidades de reportes avanzados y Business Intelligence pendientes de implementacion.",
        "Aplicacion movil (no existe en v1.0.0 — SPA web unicamente).",
        "Procesos de negocio externos al sistema (flujos en papel, aprobaciones por correo "
        "fuera de FamSPI).",
    ])
    doc.add_page_break()


# ── 4. Especificaciones del Sistema ───────────────────────────────────────────

def seccion_4(doc):
    _h(doc, "4. Especificaciones del Sistema (URS / FRS / DDS)", level=1)
    _p(doc, (
        "Las Especificaciones de Requerimientos de Usuario (URS) cubren los 12 modulos "
        "del sistema con un total de 120+ requerimientos catalogados mediante identificador "
        "unico (REQ-{MODULO}-NNN), nivel de criticidad (Alta / Media / Baja) y actor "
        "responsable. Las Especificaciones Funcionales (FRS) describen el endpoint exacto, "
        "las entradas requeridas, el proceso ejecutado y la salida HTTP esperada para cada "
        "requerimiento critico. Ambas especificaciones estan consolidadas formalmente en el "
        "Anexo A del expediente de validacion."
    ))
    _p(doc, (
        "La Especificacion de Diseno del Sistema (DDS) documenta la arquitectura de tres "
        "capas verificada en codigo fuente: capa de presentacion (React 19 + Tailwind), "
        "capa de logica de negocio (Express 5.1.0 + middlewares en cadena) y capa de "
        "persistencia (PostgreSQL Neon con pool de 20 conexiones, retry en errores "
        "transitorios y SSL). La DDS incluye la cadena completa de middlewares en orden "
        "de ejecucion (Helmet, Rate Limit, CORS, JWT, Audit), la segmentacion explicita "
        "de rutas publicas y privadas verificada en publicPaths.js y registerRoutes.js, "
        "y la lista de dependencias criticas con sus versiones fijadas. "
        "El DDS completo con extractos de codigo fuente como evidencia objetiva se "
        "encuentra en el Anexo B."
    ))
    doc.add_page_break()


# ── 5. Analisis de Riesgos (FMEA) ────────────────────────────────────────────

def seccion_5(doc):
    _h(doc, "5. Analisis de Riesgos (FMEA)", level=1)
    _p(doc, (
        "Se aplica la metodologia FMEA (Failure Mode and Effects Analysis) para priorizar "
        "los casos de prueba segun el nivel de riesgo de cada modo de falla potencial. "
        "El indice de prioridad de riesgo se calcula como: RPN = Severidad (S) x "
        "Probabilidad (P) x Detectabilidad (D), con escala de 1 a 5 en cada dimension. "
        "Los modos de falla con RPN >= 30 o Severidad = 5 sin mitigacion documentada "
        "requieren verificacion obligatoria durante la fase OQ antes de poder declarar "
        "el cierre del ciclo de validacion. La matriz completa con los 15 modos de falla "
        "identificados (R-001 a R-015), controles implementados y riesgo residual "
        "se encuentra en el Anexo C."
    ))

    _h(doc, "Riesgos criticos que requieren verificacion obligatoria en OQ", level=2)
    _table(doc,
        ["ID", "Modulo", "Modo de falla", "S", "P", "D", "RPN", "Nivel", "Caso OQ obligatorio"],
        [
            ["R-002", "auth",
             "Compromiso de credenciales OAuth2 / robo o replay de token JWT",
             "5", "2", "3", "30", "ALTO", "OQ-014"],
            ["R-003", "auth",
             "Sesion no invalidada tras logout — acceso residual con token robado",
             "4", "2", "3", "24", "Medio", "OQ-004"],
            ["R-004", "security",
             "Acceso fuera de horario laboral no detectado ni notificado",
             "4", "3", "2", "24", "Medio", "OQ-001"],
            ["R-009", "Usuarios (RBAC)",
             "Asignacion incorrecta de rol — acceso a funciones no autorizadas",
             "4", "2", "3", "24", "Medio", "OQ-005"],
            ["R-011", "permisos / vacaciones",
             "Calculo incorrecto de saldo de dias — impacto en nomina",
             "4", "2", "3", "24", "Medio", "OQ-007 / OQ-008"],
        ],
        widths=[1.2, 2, 5, 0.6, 0.6, 0.6, 0.9, 1.3, 2.8],
        estado_col=False)

    _p(doc, (
        "Nota: R-002 es el unico riesgo de nivel ALTO en FamSPI v1.0.0. Su verificacion "
        "en OQ-014 (replay de token expirado rechazado) es condicion necesaria para "
        "declarar el cierre del Area 01. Los demas riesgos del nivel Medio han sido "
        "mitigados mediante controles implementados en codigo y se verifican en las "
        "pruebas OQ correspondientes."
    ), italic=True, size=9)
    doc.add_page_break()


# ── 6. Plan de Calificacion y Protocolos de Prueba ───────────────────────────

# Columnas comunes DQ / IQ / OQ
QUAL_H = ["ID", "Descripcion de la Prueba",
           "Criterio de Aceptacion", "Ref. de Evidencia", "Resultado"]
QUAL_W = [1.4, 5, 4, 3.5, 3.6]

# Columnas PQ / UAT
PQ_H = ["ID PQ", "Escenario de Uso Real", "Participantes",
         "Resultado Esperado", "Ref. Evidencia", "Resultado"]
PQ_W = [1.2, 5, 3, 3.5, 2.5, 2.3]


def seccion_6(doc):
    _h(doc, "6. Plan de Calificacion y Protocolos de Prueba", level=1)
    _p(doc, (
        "Esta seccion contiene los protocolos de las cuatro fases de calificacion. "
        "Cada subseccion incluye el contexto narrativo de la fase seguido de la tabla "
        "de ejecucion. Las Areas 01 y 02 tienen resultado conforme (13/05/2026). "
        "Las Areas 03 a 06 tienen ejecucion pendiente — los campos de resultado "
        "se completaran durante la ejecucion formal del protocolo."
    ), after=4)

    # ── 6.1 DQ ────────────────────────────────────────────────────────────────
    _h(doc, "6.1 Calificacion de Diseno (DQ)", level=2)
    _p(doc, (
        "La Calificacion de Diseno (DQ) verifica, mediante revision documental, que el "
        "diseno tecnico propuesto y la configuracion de FamSPI son adecuados para su "
        "proposito previsto y garantizan la integridad de los datos, cumpliendo con las "
        "URS y FRS antes de la instalacion."
    ), size=9)
    _table(doc, QUAL_H, [
        ["DQ-001",
         "Arquitectura de tres capas documentada con tecnologias, versiones y archivo "
         "de configuracion de cada capa",
         "Existe DDS con descripcion completa de presentacion (React 19), logica "
         "(Express 5.1.0) y persistencia (PostgreSQL Neon)",
         "Anexo B, Sec. B.1 — tabla de arquitectura verificada en codigo fuente",
         "Conforme"],
        ["DQ-002",
         "Segmentacion explicita de rutas publicas y privadas declarada en codigo",
         "publicPaths.js contiene lista de prefijos; mountPublicRoutes() y "
         "mountPrivateRoutes() separados en registerRoutes.js",
         "Anexo B, Sec. B.3; Anexo D, EV-D-005",
         "Conforme"],
        ["DQ-003",
         "URS documentada para los 12 modulos del sistema con REQ-IDs unicos "
         "y niveles de criticidad definidos",
         "Existen 12 archivos URS bajo docs/validation/URS/ con al menos "
         "5 REQ catalogados por modulo",
         "Anexo A — tabla resumen de 12 modulos con estado documental",
         "Conforme"],
        ["DQ-004",
         "Evaluacion FMEA completada: 15 modos de falla con S/P/D/RPN, "
         "mitigacion y estado de cada riesgo",
         "Matriz FMEA con R-001 a R-015; R-002 identificado como ALTO (RPN=30) "
         "con control documentado",
         "Anexo C — Matriz FMEA completa; general/07A_evaluacion_riesgos_fmea.md",
         "Conforme"],
        ["DQ-005",
         "Todos los requerimientos de criticidad Alta tienen diseno tecnico implementado "
         "y endpoint FRS correspondiente verificable en codigo",
         "REQ de criticidad Alta con FRS asociado en Anexo A; "
         "implementacion activa en produccion",
         "Anexo A, Tablas FRS-AUTH y FRS-TH; Anexo D, EV-D-001 a EV-D-005",
         "Conforme"],
    ], widths=QUAL_W)

    # ── 6.2 IQ ────────────────────────────────────────────────────────────────
    _h(doc, "6.2 Calificacion de Instalacion (IQ)", level=2)
    _p(doc, (
        "La Calificacion de Instalacion (IQ) provee evidencia documentada de que el "
        "software, la infraestructura subyacente (Node.js, PM2) y la base de datos "
        "(PostgreSQL / Neon) estan correctamente instalados y configurados en el entorno "
        "de produccion segun el DDS aprobado."
    ), size=9)
    _table(doc, QUAL_H, [
        ["IQ-001",
         "Node.js instalado en servidor de produccion con version compatible (v18.x o superior)",
         "node --version retorna v18.x o superior en el servidor productivo",
         "Captura de terminal: node --version en servidor produccion",
         "Conforme"],
        ["IQ-002",
         "Express 5.1.0 instalado segun package.json; version activa verificada",
         "package.json declara 'express': '^5.1.0'; npm list retorna 5.1.x instalado",
         "backend/package.json; salida npm list express",
         "Conforme"],
        ["IQ-003",
         "Variables de entorno criticas presentes y no vacias en produccion",
         "SECRET_KEY, DATABASE_URL, GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET "
         "y NODE_ENV=production presentes",
         "Verificacion con script de arranque (valores no registrados por seguridad)",
         "Conforme"],
        ["IQ-004",
         "Conexion a base de datos PostgreSQL/Neon activa; pool inicializado correctamente",
         "Aplicacion inicia sin error de conexion; /health retorna db_status: connected",
         "Log de arranque PM2; respuesta de /health en produccion",
         "Conforme"],
        ["IQ-005",
         "Proceso PM2 activo y configurado para reinicio automatico ante fallos",
         "pm2 list muestra proceso en estado 'online'; restart_policy activa",
         "Captura de pm2 list y pm2 show en servidor de produccion",
         "Conforme"],
        ["IQ-006",
         "HTTPS/TLS activo en produccion; HTTP redirige a HTTPS",
         "Acceso https:// retorna 200; http:// redirige con HTTP 301",
         "Verificacion via curl del dominio productivo en ambos protocolos",
         "Conforme"],
        ["IQ-007",
         "Rate limiting activo: maximo 3000 solicitudes por 15 minutos en produccion",
         "Cabecera X-RateLimit-Limit: 3000 presente en respuestas de la API",
         "Anexo D, EV-D-004; backend/src/config/security.js",
         "Conforme"],
        ["IQ-008",
         "Cabeceras de seguridad Helmet activas: CSP, HSTS, X-Frame-Options",
         "Respuesta HTTP contiene Content-Security-Policy, "
         "Strict-Transport-Security y X-Frame-Options: DENY",
         "Inspeccion de cabeceras HTTP en produccion via curl -I",
         "Conforme"],
        ["IQ-009",
         "Tablas principales de base de datos creadas y accesibles en produccion",
         "Consulta a information_schema.tables retorna todas las tablas principales "
         "sin errores",
         "Script de verificacion de esquema ejecutado en produccion",
         "Conforme"],
        ["IQ-010",
         "Frontend React compilado y servido correctamente en produccion",
         "Acceso al dominio carga la SPA React; sin errores HTTP 4xx/5xx en carga inicial",
         "Captura de pantalla del dashboard cargado; sin errores en consola",
         "Conforme"],
        ["IQ-011",
         "Usuario administrador inicial creado y funcional en produccion",
         "Al menos un usuario con rol 'administrador' puede autenticarse y acceder "
         "al modulo de administracion",
         "Captura de sesion activa de usuario admin en produccion",
         "Conforme"],
        ["IQ-012",
         "Procedimiento de rollback documentado; restauracion desde respaldo verificada",
         "Existe procedimiento escrito de rollback; al menos una restauracion de prueba "
         "ejecutada y documentada",
         "docs/validation/general/10_calificacion_instalacion_iq.md — seccion rollback",
         "Pendiente"],
    ], widths=QUAL_W)

    # ── 6.3 OQ ────────────────────────────────────────────────────────────────
    _h(doc, "6.3 Calificacion Operacional (OQ)", level=2)
    _p(doc, (
        "La Calificacion Operacional (OQ) somete al sistema a pruebas funcionales para "
        "demostrar que opera conforme a lo esperado en sus rangos operativos. Esto incluye "
        "retar los controles de seguridad (autenticacion OAuth2, roles RBAC, accesos "
        "invalidos) y verificar el correcto registro en la bitacora de auditoria "
        "(Audit Trail), ademas de los flujos de negocio criticos de cada area."
    ), size=9)
    _table(doc, QUAL_H, [
        ["OQ-001",
         "Autenticacion OAuth2 con cuenta Google del dominio corporativo autorizado; "
         "JWT emitido con claims correctos",
         "Usuario de dominio autorizado se autentica; JWT con iss=spi-fam-backend, "
         "aud=spi-fam-frontend, sub=usuario_id; sesion registrada en user_sessions",
         "Registro en user_sessions; token decodificado; log de auditMiddleware",
         "Conforme"],
        ["OQ-002",
         "Rechazo de cuenta Google con dominio no autorizado; sin creacion de sesion",
         "Cuenta de dominio externo recibe HTTP 403; user_sessions sin nuevo registro; "
         "mensaje de dominio no permitido en respuesta",
         "Intento de login con cuenta gmail.com; respuesta HTTP capturada",
         "Conforme"],
        ["OQ-003",
         "Renovacion de sesion via refresh token antes del vencimiento del access token",
         "POST /api/v1/auth/refresh con refresh token valido retorna nuevo accessToken; "
         "user_sessions actualizado",
         "Respuesta HTTP 200 con nuevo accessToken; registro en user_sessions",
         "Conforme"],
        ["OQ-004",
         "Logout invalida sesion activa; token anterior rechazado post-logout",
         "POST /api/v1/auth/logout invalida sesion; uso del token previo retorna HTTP 401 "
         "con code: INVALID_TOKEN",
         "Estado de sesion en user_sessions; intento de reuso del token con respuesta 401",
         "Conforme"],
        ["OQ-005",
         "Endpoint protegido rechaza con HTTP 403 a usuario con rol sin permiso",
         "Usuario con rol 'comercial' intenta GET /api/v1/auditoria: HTTP 403 con "
         "mensaje de roles permitidos; sin datos filtrados en respuesta",
         "Respuesta HTTP 403 con body correcto; sin exposicion de datos del sistema",
         "Conforme"],
        ["OQ-006",
         "Accion de mutacion (POST/PUT/DELETE) registrada en auditoria.logs con "
         "todos los campos requeridos por WHO",
         "auditoria.logs contiene: usuario_id, usuario_email, rol, modulo, accion, "
         "ip, duracion_ms, timestamp — inmediatamente tras la accion",
         "Consulta SELECT a auditoria.logs post-ejecucion; verificacion de campos",
         "Conforme"],
        ["OQ-007",
         "Creacion de solicitud de permiso con validaciones de campos obligatorios",
         "POST /api/v1/talento-humano/permisos: HTTP 201 con estado 'pendiente'; "
         "campos obligatorios ausentes retornan HTTP 400",
         "Registro en DB con estado correcto; prueba negativa con campos vacios",
         "Conforme"],
        ["OQ-008",
         "Aprobacion de solicitud por responsable funcional autorizado; "
         "registro en auditoria.logs",
         "PUT /api/v1/talento-humano/permisos/:id/estado: estado actualizado a "
         "'aprobado'; historial con actor y timestamp; registro en auditoria",
         "DB con nuevo estado; historial de estados; registro en auditoria.logs",
         "Conforme"],
        ["OQ-009",
         "Rechazo de vacaciones con motivo obligatorio; notificacion emitida "
         "al solicitante",
         "Estado actualizado a 'rechazado' con motivo en DB; notificacion en cola "
         "de dispatchers; registro en auditoria.logs",
         "DB con estado y motivo; cola de notificaciones; registro de auditoria",
         "Conforme"],
        ["OQ-010",
         "Link de verificacion legal accesible publicamente sin JWT; "
         "datos correctos retornados",
         "GET /api/v1/permisos/legal-verification/:token: HTTP 200 con datos del "
         "documento verificado; sin necesidad de Bearer token",
         "Acceso al link sin autenticacion; respuesta HTTP 200 con datos correctos",
         "Conforme"],
        ["OQ-011",
         "Creacion de Business Case vinculado a cliente y oportunidad comercial",
         "POST /api/v1/business-case: HTTP 201; BC con estado inicial correcto; "
         "vinculado al cliente y oportunidad en DB",
         "Registro en DB; vinculacion verificada; respuesta HTTP 201 con ID",
         "Pendiente"],
        ["OQ-012",
         "Flujo de viatico: borrador → revision jefe → aprobacion financiero → pago; "
         "cada transicion registrada en auditoria",
         "Cada cambio de estado valida actor, permisos y estado previo; "
         "estado final 'pagado' en DB; historial completo en auditoria.logs",
         "Registros de auditoria por transicion; estado final en DB",
         "Pendiente"],
        ["OQ-013",
         "Ticket de soporte TI creado, asignado a tecnico y cerrado con "
         "trazabilidad completa",
         "Flujo completo sin errores; auditoria registra cada transicion; "
         "cierre con causa, tiempo y datos completos",
         "DB con todos los estados; auditoria de transiciones; cierre documentado",
         "Pendiente"],
        ["OQ-014",
         "R-002 FMEA — VERIFICACION OBLIGATORIA: replay de access token expirado "
         "rechazado correctamente",
         "Token expirado o de sesion invalida retorna HTTP 401 con "
         "code: INVALID_TOKEN; sin acceso concedido; sin datos en respuesta",
         "Respuesta HTTP 401 capturada; body con code correcto; "
         "verificacion en auth.js lineas criticas",
         "Pendiente"],
    ], widths=QUAL_W)

    # ── 6.4 PQ ────────────────────────────────────────────────────────────────
    _h(doc, "6.4 Calificacion de Desempeno y Pruebas de Usuario (PQ / UAT)", level=2)
    _p(doc, (
        "La Calificacion de Desempeno (PQ) y UAT se ejecuta en el entorno en vivo con "
        "usuarios funcionales previamente capacitados segun el Anexo E. El objetivo es "
        "confirmar la idoneidad del sistema ejecutando flujos de negocio reales de "
        "extremo a extremo, como la solicitud y aprobacion completa de permisos y "
        "vacaciones, hasta el registro de pago de viaticos, con participacion de los "
        "actores reales de cada proceso."
    ), size=9)
    _table(doc, PQ_H, [
        ["PQ-001",
         "Permiso punta a punta: creacion por solicitante → aprobacion → verificacion QR legal",
         "Solicitante + Aprobador + TI",
         "Flujo completo sin errores; trazabilidad completa; QR funcional; estado final correcto",
         "Acta UAT Area 02; capturas de pantalla del flujo",
         "Conforme"],
        ["PQ-002",
         "Vacaciones punta a punta: creacion → aprobacion → link legal accesible",
         "Solicitante + Aprobador",
         "Flujo completo sin errores; estado final correcto en DB; link de verificacion activo",
         "Acta UAT Area 02; capturas del flujo completo",
         "Conforme"],
        ["PQ-003",
         "Aprobacion y rechazo con distintos motivos por responsable funcional",
         "Responsable funcional + TI",
         "Ambas decisiones ejecutadas correctamente; notificaciones emitidas; historial auditable",
         "Capturas de ambas decisiones; notificaciones en DB; registros de auditoria",
         "Conforme"],
        ["PQ-004",
         "Consulta de historial con filtros de estado, fecha y usuario; paginacion",
         "TI + Funcional",
         "Filtros retornan resultados correctos; paginacion operativa; consistencia con DB",
         "Capturas de consultas filtradas; comparacion con datos en DB",
         "Conforme"],
        ["PQ-005",
         "Aceptacion formal de Areas 01 y 02 por usuario clave y gerencia",
         "Responsable funcional + Gerencia",
         "Acta de aceptacion firmada con declaracion de conformidad del uso previsto",
         "Acta firmada en docs/validation/areas/area_01_*/ y area_02_*/",
         "Conforme"],
        ["PQ-006",
         "Business Case completo: creacion → seguimiento → cierre",
         "Comercial + Gerencia",
         "BC pasa todos los estados; documentacion completa; vinculacion a cliente activa",
         "Capturas del flujo; datos en DB verificados",
         "Pendiente"],
        ["PQ-007",
         "Asignacion de cliente y registro de visita comercial con datos completos",
         "Comercial + Administrador",
         "Cliente asignado; visita registrada con fecha, actor y resultado",
         "Capturas de registro; datos en DB; auditoria del evento",
         "Pendiente"],
        ["PQ-008",
         "Aceptacion formal del modulo comercial por Jefe Comercial y Gerencia",
         "Jefe Comercial + Gerencia",
         "Acta de aceptacion firmada declarando conformidad",
         "Acta de aceptacion Area 03 — pendiente emision",
         "Pendiente"],
        ["PQ-009",
         "Capacitacion punta a punta: planificacion → ejecucion → asistencia → certificado",
         "TH + Instructor + Tecnico",
         "Flujo completo sin errores; certificado generado en PDF con datos correctos",
         "Capturas del flujo; certificado PDF generado; datos en DB",
         "Pendiente"],
        ["PQ-010",
         "Orden de mantenimiento preventivo: creacion → asignacion → ejecucion → cierre",
         "Tecnico + Supervisor",
         "Orden pasa todos los estados validos; cierre con evidencia de ejecucion",
         "Capturas del flujo; datos en DB; auditoria de transiciones",
         "Pendiente"],
        ["PQ-011",
         "Proyecto externo GoApp con hitos: creacion, seguimiento y cierre de hitos",
         "Tecnico + Cliente externo",
         "Hitos creados y actualizados; trazabilidad de avance disponible",
         "Capturas del flujo de proyecto; datos en DB",
         "Pendiente"],
        ["PQ-012",
         "Aceptacion formal del modulo servicio tecnico por Jefe Tecnico y Gerencia",
         "Jefe Tecnico + Gerencia",
         "Acta de aceptacion firmada declarando conformidad",
         "Acta de aceptacion Area 04 — pendiente emision",
         "Pendiente"],
        ["PQ-013",
         "Alta, asignacion y baja de activo TI con acta generada en cada transicion",
         "TI + Usuario receptor",
         "Activo pasa por alta → asignado → dado_de_baja; acta PDF en cada estado",
         "Capturas de cada estado; datos en DB; actas generadas",
         "Pendiente"],
        ["PQ-014",
         "Ticket de soporte TI de creacion a cierre con SLA registrado",
         "Usuario + Tecnico TI",
         "Ticket pasa por creacion → asignado → en_progreso → resuelto → cerrado; "
         "SLA calculado y visible",
         "Capturas del flujo; datos en DB; SLA verificado",
         "Pendiente"],
        ["PQ-015",
         "Aceptacion formal del modulo inventario/TI por Jefe TI y Gerencia",
         "Jefe TI + Gerencia",
         "Acta de aceptacion firmada declarando conformidad",
         "Acta de aceptacion Area 05 — pendiente emision",
         "Pendiente"],
        ["PQ-016",
         "Viatico punta a punta: borrador → carga facturas TXT → revision → pago registrado",
         "Solicitante + Jefe + Finanzas",
         "Flujo completo sin errores; facturas TXT parseadas; estado final 'pagado' en DB",
         "Capturas de cada paso del wizard; facturas en DB; estado final verificado",
         "Pendiente"],
        ["PQ-017",
         "Categorizacion de facturas y exportacion ATS/XML formato SRI Ecuador",
         "Finanzas",
         "Facturas categorizadas; archivo ATS/XML generado con estructura valida SRI",
         "Archivo ATS/XML exportado; validacion de estructura del archivo",
         "Pendiente"],
        ["PQ-018",
         "Aceptacion formal del modulo finanzas/viaticos por Jefe Finanzas y Gerencia",
         "Jefe Finanzas + Gerencia",
         "Acta de aceptacion firmada declarando conformidad",
         "Acta de aceptacion Area 06 — pendiente emision",
         "Pendiente"],
    ], widths=PQ_W)

    doc.add_page_break()


# ── 7. Matriz de Trazabilidad (RTM) ───────────────────────────────────────────

def seccion_7(doc):
    _h(doc, "7. Matriz de Trazabilidad (RTM)", level=1)
    _p(doc, (
        "Esta matriz asegura que ningun requerimiento del usuario quede sin probar y que "
        "todo riesgo identificado en el FMEA este mitigado y verificado mediante al menos "
        "un caso de prueba. Conecta cada REQ-ID con el riesgo FMEA asociado, el protocolo "
        "de calificacion que lo verifica y el estado actual de dicha verificacion. "
        "El RTM completo del sistema se mantiene actualizado en "
        "docs/validation/RTM/RTM_sistema_spi.md."
    ), after=4)
    _table(doc,
        ["ID URS", "Modulo / Area", "Descripcion resumida",
         "Riesgo FMEA", "Caso de prueba", "Estado"],
        [
            ["REQ-AUTH-001", "Area 01 — Auth",
             "Login via Google OAuth2 con cuenta corporativa",
             "R-002 (RPN=30)", "OQ-001 / PQ-001", "Conforme"],
            ["REQ-AUTH-002", "Area 01 — Auth",
             "Validacion de dominio corporativo en login",
             "R-002 (RPN=30)", "OQ-002", "Conforme"],
            ["REQ-AUTH-005", "Area 01 — Auth",
             "Emision de accessToken y refreshToken",
             "R-003 (RPN=24)", "OQ-003", "Conforme"],
            ["REQ-AUTH-007", "Area 01 — Auth",
             "Cierre de sesion explicito e invalidacion",
             "R-003 (RPN=24)", "OQ-004", "Conforme"],
            ["REQ-USR-002",  "Area 01 — Users",
             "Restriccion de acceso por rol RBAC en endpoints",
             "R-009 (RPN=24)", "OQ-005", "Conforme"],
            ["REQ-AUD-001",  "Area 01 — Auditoria",
             "Registro de toda accion critica en bitacora audit trail",
             "R-005 (RPN=10)", "OQ-006", "Conforme"],
            ["REQ-SEC-001",  "Area 01 — Security",
             "Deteccion de acceso fuera de horario laboral",
             "R-004 (RPN=24)", "OQ-001", "Conforme"],
            ["REQ-TH-001",   "Area 02 — TH",
             "Administracion de perfiles de colaboradores",
             "R-011 (RPN=24)", "OQ-007 / PQ-001", "Conforme"],
            ["REQ-TH-002",   "Area 02 — TH",
             "Flujo jerarquico de permisos y vacaciones",
             "R-011 (RPN=24)", "OQ-008 / OQ-009 / PQ-002", "Conforme"],
            ["REQ-TH-006",   "Area 02 — TH",
             "Links de verificacion legal por QR sin JWT",
             "R-007 (RPN=20)", "OQ-010 / PQ-003", "Conforme"],
            ["REQ-COM-001",  "Area 03 — Comercial",
             "Creacion de Business Case vinculado a oportunidad",
             "R-013 (RPN=27)", "OQ-011 / PQ-006", "Pendiente"],
            ["REQ-COM-005",  "Area 03 — Comercial",
             "Registro de visita y asignacion de cliente",
             "R-013 (RPN=27)", "PQ-007", "Pendiente"],
            ["REQ-SRV-001",  "Area 04 — Servicio",
             "Orden de mantenimiento preventivo punta a punta",
             "R-001 (RPN=20)", "PQ-010", "Pendiente"],
            ["REQ-SRV-003",  "Area 04 — Servicio",
             "Capacitacion con emision de certificado",
             "R-001 (RPN=20)", "PQ-009", "Pendiente"],
            ["REQ-INV-001",  "Area 05 — Inventario",
             "Alta y baja de activo TI con acta generada",
             "R-008 (RPN=12)", "PQ-013", "Pendiente"],
            ["REQ-TI-001",   "Area 05 — TI",
             "Ticket de soporte de creacion a cierre con SLA",
             "R-008 (RPN=12)", "OQ-013 / PQ-014", "Pendiente"],
            ["REQ-VT-001",   "Area 06 — Viaticos",
             "Creacion de workspace de viatico con wizard 4 pasos",
             "R-011 (RPN=24)", "OQ-012 / PQ-016", "Pendiente"],
            ["REQ-VT-003",   "Area 06 — Viaticos",
             "Carga de facturas TXT SRI 14 columnas",
             "R-008 (RPN=12)", "PQ-016", "Pendiente"],
            ["REQ-VT-016",   "Area 06 — Viaticos",
             "Trazabilidad completa de cambios de estado",
             "R-005 (RPN=10)", "OQ-012", "Pendiente"],
            ["REQ-DOC-001",  "Area 06 — Firma",
             "Firma digital de documentos con hash SHA-256 y QR",
             "R-007 (RPN=20)", "OQ-010", "Pendiente"],
            ["REQ-SEC-002",  "Area 01 — Auth",
             "R-002 ALTO: Replay de token expirado rechazado",
             "R-002 ALTO (RPN=30)", "OQ-014", "Pendiente"],
        ],
        widths=[2.2, 3, 4.3, 2.5, 2.8, 2], font=8)
    doc.add_page_break()


# ── 8. Procedimientos Operativos (SOPs) y Entrenamiento ──────────────────────

def seccion_8(doc):
    _h(doc, "8. Procedimientos Operativos (SOPs) y Entrenamiento", level=1)
    _p(doc, (
        "De conformidad con la seccion 11 de WHO TRS 1019 Annex 3, el estado validado de "
        "FamSPI se mantendra mediante procedimientos operativos documentados (SOPs) que "
        "garantizan que cualquier cambio, incidente o nuevo ciclo de uso no degrade la "
        "integridad del sistema ni comprometa la trazabilidad de los datos. "
        "Los documentos que soportan esta seccion son externos a este protocolo y se "
        "referencian a continuacion."
    ))

    _h(doc, "8.1 Procedimientos operativos vigentes", level=2)
    _bullet(doc, [
        "SOP de Control de Cambios (ref. docs/validation/general/14A_control_cambios.md): "
        "define el proceso de evaluacion, aprobacion, implementacion y documentacion de "
        "cualquier cambio al codigo, configuracion o base de datos de FamSPI. "
        "Todo cambio en alcance validado requiere apertura de un registro de cambio, "
        "evaluacion de impacto sobre el estado validado y re-ejecucion de los casos "
        "de prueba OQ/PQ afectados antes de liberar el cambio.",
        "SOP de Revision Periodica y Monitoreo (ref. docs/validation/general/14B_revision_periodica.md): "
        "establece la frecuencia y criterios de la revision periodica del estado validado, "
        "incluyendo revision de bitacora de auditoria, usuarios y roles activos, "
        "disponibilidad del sistema, dependencias con vulnerabilidades (npm audit) "
        "y cumplimiento de los criterios de la seccion 14 de este expediente.",
        "SOP de Respuesta a Incidentes de Seguridad: define los pasos de contencion, "
        "investigacion y documentacion ante eventos de acceso no autorizado, perdida de "
        "datos o deteccion de cambios no controlados, incluyendo la apertura de "
        "desviaciones formales y la notificacion a gerencia.",
        "SOP de Backup y Recuperacion: documenta el procedimiento de verificacion "
        "periodica de los respaldos provistos por Neon PostgreSQL y la prueba anual "
        "de restauracion, con registro del resultado en la bitacora operacional.",
    ])

    _h(doc, "8.2 Entrenamiento del personal", level=2)
    _p(doc, (
        "El personal que ejecuta las pruebas PQ/UAT o que opera FamSPI en roles criticos "
        "debe tener entrenamiento formal documentado antes de la ejecucion de la "
        "Calificacion de Desempeno (PQ), de conformidad con WHO TRS 1019 Annex 3, Sec. 12. "
        "Los registros de entrenamiento son parte integral del expediente de validacion "
        "y se encuentran en el Anexo E."
    ))
    _table(doc,
        ["Perfil de usuario", "Modulos cubiertos", "Duracion estimada", "Estado"],
        [
            ["Solicitante general\n(comercial, tecnico, RRHH, backoffice)",
             "Dashboard, permisos, vacaciones, workspace de viaticos, notificaciones",
             "2-4 horas", "Pendiente"],
            ["Aprobador / Jefe de area",
             "Todo lo del solicitante + flujos de aprobacion y rechazo por area",
             "3-5 horas", "Pendiente"],
            ["Finanzas",
             "Viaticos financiero, categorizacion de facturas, exportacion ATS/XML, reportes",
             "4-6 horas", "Pendiente"],
            ["TI — Administrador del sistema",
             "Sistema completo: usuarios, roles, auditoria, soporte, activos, respaldo, monitoreo",
             "8-12 horas (sesiones)", "Parcial"],
        ],
        widths=[4, 7.5, 2.5, 2.8])
    doc.add_page_break()


# ── 9. Indice de Anexos ────────────────────────────────────────────────────────

def seccion_9(doc):
    _h(doc, "9. Indice de Anexos", level=1)
    _p(doc, (
        "Los documentos adjuntos a continuacion son parte integral del expediente de "
        "validacion de FamSPI v1.0.0. El cuerpo de este protocolo los referencia "
        "en lugar de reproducir su contenido, para facilitar el mantenimiento y "
        "la auditoria del expediente."
    ))
    _table(doc,
        ["Anexo", "Titulo", "Contenido principal", "Archivo"],
        [
            ["A", "URS y FRS completas",
             "Tablas de REQ-IDs por modulo (120+ requerimientos), tablas FRS con "
             "endpoints, entradas, proceso y salida HTTP para todos los modulos criticos",
             "ANEXOS_VALIDACION_FAMSPI_V1_0_0.docx — Anexo A"],
            ["B", "DDS — Especificacion de Diseno",
             "Arquitectura de 3 capas, cadena de middlewares en orden de ejecucion, "
             "segmentacion de rutas, configuracion del pool de BD, dependencias criticas",
             "ANEXOS_VALIDACION_FAMSPI_V1_0_0.docx — Anexo B"],
            ["C", "Matriz FMEA completa",
             "15 modos de falla R-001 a R-015 con S/P/D/RPN, mitigacion implementada, "
             "riesgo residual y estado. R-002 marcado ACTIVO hasta cierre de OQ-014",
             "ANEXOS_VALIDACION_FAMSPI_V1_0_0.docx — Anexo C"],
            ["D", "Evidencias Objetivas — codigo fuente verificado",
             "Extractos reales del codigo: verifyToken, requireRole con ROLE_GROUPS, "
             "auditMiddleware, cadena de middlewares en app.js, publicPaths.js completo",
             "ANEXOS_VALIDACION_FAMSPI_V1_0_0.docx — Anexo D"],
            ["E", "Registros de Capacitacion y Entrenamiento",
             "Plan de entrenamiento por perfil y formularios de lista de asistencia "
             "para 4 perfiles: solicitante, aprobador, finanzas, TI administrador",
             "ANEXOS_VALIDACION_FAMSPI_V1_0_0.docx — Anexo E"],
        ],
        widths=[1.2, 3.8, 8.5, 4.5], estado_col=False)

    doc.add_paragraph()
    fin = doc.add_paragraph()
    fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fin.add_run("— Fin del Protocolo Maestro de Validacion FamSPI v1.0.0 —")
    r.italic = True; r.font.size = Pt(9)
    r.font.name = "Calibri"; r.font.color.rgb = GRIS


# ── Construccion ──────────────────────────────────────────────────────────────

def build_document(path=None):
    target = path or OUTPUT_PATH
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    portada(doc)
    seccion_1(doc)
    seccion_2(doc)
    seccion_3(doc)
    seccion_4(doc)
    seccion_5(doc)
    seccion_6(doc)
    seccion_7(doc)
    seccion_8(doc)
    seccion_9(doc)

    doc.save(target)
    print(f"Documento generado: {target}")


if __name__ == "__main__":
    build_document()
