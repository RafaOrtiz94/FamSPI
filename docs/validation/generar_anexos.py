"""Generador de Anexos del Protocolo de Validacion FamSPI v1.0.0.

Genera un documento Word con los Anexos A-E usando evidencia extraida
directamente del codigo fuente y los archivos de documentacion del proyecto.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement as OE

OUTPUT = (
    r"c:\Users\Departamento de TI\Desktop\PROYECTOS\FamSPI"
    r"\docs\validation\ANEXOS_VALIDACION_FAMSPI_V1_0_0.docx"
)
FECHA = "18/06/2026"
AZUL = RGBColor(0x1F, 0x49, 0x7D)
VERDE = RGBColor(0x37, 0x86, 0x50)
GRIS = RGBColor(0x80, 0x80, 0x80)
NARANJA = RGBColor(0xC5, 0x5A, 0x11)


# ── Helpers ──────────────────────────────────────────────────────────────────


def _shd(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _cell(cell, text, bold=False, size=8, color=None, fill=None, center=False, wrap=True):
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if fill:
        _shd(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _w(doc, widths_cm, table):
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def _table(doc, headers, rows, widths=None, hbg="1F497D", font=8):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        _cell(t.rows[0].cells[i], h, bold=True, color="FFFFFF", fill=hbg, center=True, size=font)
    for ri, rd in enumerate(rows):
        row = t.add_row()
        fill = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, v in enumerate(rd):
            _cell(row.cells[ci], str(v), size=font, fill=fill)
    if widths:
        _w(doc, widths, t)
    doc.add_paragraph()
    return t


def _h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.color.rgb = AZUL
        elif level == 2:
            run.font.color.rgb = VERDE
        else:
            run.font.color.rgb = GRIS
    return p


def _p(doc, text, italic=False, size=10, after=5, color=None):
    para = doc.add_paragraph()
    para.style = doc.styles["Normal"]
    para.paragraph_format.space_after = Pt(after)
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return para


def _code(doc, lines):
    """Bloque de codigo con fondo gris oscuro y fuente monoespaciada."""
    for line in lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.left_indent = Cm(0.5)
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1E1E1E")
        pPr.append(shd)
        run = para.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
    doc.add_paragraph()


def _label(doc, text, color_hex="1F497D"):
    """Etiqueta de seccion con fondo de color."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(4)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _portada_anexo(doc, letra, titulo, descripcion):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"ANEXO {letra}")
    r.bold = True
    r.font.size = Pt(26)
    r.font.name = "Calibri"
    r.font.color.rgb = AZUL
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = t.add_run(titulo)
    rt.bold = True
    rt.font.size = Pt(14)
    rt.font.name = "Calibri"
    rt.font.color.rgb = AZUL
    _p(doc, descripcion, italic=True, size=9, color=GRIS)
    doc.add_paragraph()


# ── ANEXO A — URS y FRS ──────────────────────────────────────────────────────


def anexo_a(doc):
    _portada_anexo(
        doc, "A", "ESPECIFICACION DE REQUERIMIENTOS DEL USUARIO (URS) Y ESPECIFICACION FUNCIONAL (FRS)",
        "Evidencia: extraida de docs/validation/URS/ y docs/validation/FRS/ — archivos por modulo. "
        "Fecha de extraccion: " + FECHA,
    )
    doc.add_page_break()

    _h(doc, "A.1 Requerimientos de usuario (URS) por modulo", level=1)
    _p(doc, (
        "Cada modulo del sistema tiene su documento URS completo en docs/validation/URS/. "
        "La tabla siguiente consolida los requerimientos criticos verificables con su ID, "
        "actor real y resultado esperado. La fuente completa de cada REQ se encuentra "
        "en el archivo URS del modulo correspondiente."
    ))

    # Modulo AUTH — extraido del URS real
    _label(doc, "Modulo: Autenticacion y Sesiones  |  Archivo: URS_modulo_autenticacion_sesiones.md", "1F497D")
    _table(doc,
        ["REQ-ID", "Actor", "Requerimiento", "Resultado esperado", "Criticidad"],
        [
            ["REQ-AUTH-001", "Colaborador interno", "El sistema debe permitir iniciar sesion mediante cuenta Google OAuth2", "Usuario autenticado redirigido al dashboard segun su rol", "Alta"],
            ["REQ-AUTH-002", "Sistema", "El sistema debe validar que el dominio del correo pertenezca al dominio corporativo configurado", "Usuarios de dominio no permitido no pueden ingresar", "Alta"],
            ["REQ-AUTH-003", "Sistema", "El sistema debe crear automaticamente un usuario al autenticar por primera vez", "Usuario registrado con rol inicial 'pendiente' y datos basicos de identidad", "Alta"],
            ["REQ-AUTH-005", "Colaborador interno", "El sistema debe emitir accessToken y refreshToken tras login exitoso", "Frontend recibe ambos tokens y habilita sesion autenticada", "Alta"],
            ["REQ-AUTH-006", "Colaborador interno", "El sistema debe permitir renovar la sesion sin nuevo login mientras el refresh token sea valido", "Se emiten nuevos tokens; registro de sesion actualizado", "Media"],
            ["REQ-AUTH-007", "Colaborador interno", "El sistema debe permitir cerrar sesion de forma explicita", "La sesion activa queda cerrada en user_sessions", "Media"],
            ["REQ-AUTH-010", "TI y Gerencia", "El sistema debe permitir consultar historial de sesiones y usuarios activos", "Solo usuarios con rol autorizado acceden a auditoria de sesiones", "Alta"],
            ["REQ-AUTH-012", "Sistema", "El sistema debe detectar y notificar accesos fuera de horario laboral", "Evento registrado; notificacion emitida al canal de TI", "Media"],
        ],
        widths=[2.2, 3, 6, 4.5, 1.8],
    )

    # Modulo TH — extraido del URS real
    _label(doc, "Modulo: Talento Humano  |  Archivo: URS_modulo_talento_humano.md", "1F497D")
    _table(doc,
        ["REQ-ID", "Actor", "Requerimiento", "Resultado esperado", "Criticidad"],
        [
            ["REQ-TH-001", "Talento Humano", "El sistema debe permitir administrar informacion de colaboradores y perfiles de personal", "Los perfiles quedan actualizados con trazabilidad completa", "Alta"],
            ["REQ-TH-002", "Colaborador y aprobadores", "El sistema debe permitir crear y resolver solicitudes de permisos/vacaciones por flujo jerarquico", "La solicitud avanza por estados validos hasta cierre con trazabilidad", "Alta"],
            ["REQ-TH-003", "TH / Gerencia", "El sistema debe permitir gestionar solicitudes de personal, incluyendo perfil, documentos y contratacion", "El expediente de vacante queda completo y consistente", "Media"],
            ["REQ-TH-004", "Colaborador autenticado", "El sistema debe permitir registrar marcaciones de asistencia, excepciones y overtime", "Los registros quedan consolidados para reporte y control", "Media"],
            ["REQ-TH-005", "Talento Humano", "El sistema debe permitir administrar departamentos organizacionales", "La estructura departamental se mantiene vigente para otros modulos", "Baja"],
            ["REQ-TH-006", "Sistema", "El sistema debe emitir links de verificacion legal para firmas de permisos/vacaciones", "Las validaciones legales quedan verificables por token publico", "Alta"],
        ],
        widths=[2.2, 3, 6, 4.5, 1.8],
    )

    # Modulo Viaticos
    _label(doc, "Modulo: Finanzas y Viaticos  |  Archivo: URS_modulo_finanzas_viaticos.md", "1F497D")
    _table(doc,
        ["REQ-ID", "Actor", "Requerimiento", "Resultado esperado", "Criticidad"],
        [
            ["REQ-VT-001", "Solicitante", "El sistema debe permitir crear un workspace de viaticos vinculado a una salida operacional", "Workspace creado con estado 'borrador'; vinculado al colaborador", "Alta"],
            ["REQ-VT-003", "Solicitante", "El sistema debe permitir cargar facturas electronicas SRI desde archivo TXT de 14 columnas", "Facturas parseadas, validadas y asociadas al workspace", "Alta"],
            ["REQ-VT-006", "Solicitante", "El sistema debe implementar un asistente guiado de 4 pasos para completar el viatico", "El wizard guia la carga de soportes, categorias, notas y envio", "Alta"],
            ["REQ-VT-008", "Sistema", "El sistema debe gestionar 9 estados del flujo: borrador, pendiente_revision, aprobado_jefe, rechazado_jefe, pendiente_financiero, aprobado_financiero, rechazado_financiero, listo_pago, pagado", "Cada transicion valida actor, permisos y estado previo", "Alta"],
            ["REQ-VT-012", "Revision financiera", "El sistema debe permitir categorizar facturas por tipo de gasto (combustible, alimentacion, hospedaje, transporte, movilidad, materiales)", "Cada factura queda categorizada para reportes de gasto", "Alta"],
            ["REQ-VT-016", "Sistema", "El sistema debe registrar trazabilidad completa de cada cambio de estado con actor y timestamp", "Trail auditable disponible en bitacora de auditoria", "Alta"],
        ],
        widths=[2.2, 3, 6, 4.5, 1.8],
    )

    # Tabla resumen de todos los modulos
    _label(doc, "Resumen de todos los modulos — estado de la documentacion URS", "37490D")
    _table(doc,
        ["Modulo", "N REQ", "Archivo fuente", "Estado"],
        [
            ["Autenticacion y Sesiones", "12", "URS_modulo_autenticacion_sesiones.md", "Completo"],
            ["Usuarios y Perfiles", "13", "URS_modulo_usuarios_perfiles.md", "Completo"],
            ["Talento Humano", "6", "URS_modulo_talento_humano.md", "Completo"],
            ["Comercial y Clientes", "18", "URS_modulo_comercial_clientes.md", "v2.0 (2026-06)"],
            ["Business Case", "5+", "URS_modulo_business_case.md", "En revision"],
            ["Finanzas y Viaticos", "16", "URS_modulo_finanzas_viaticos.md", "v2.0 (2026-06)"],
            ["Servicio Tecnico", "11", "URS_modulo_servicio_tecnico_mantenimientos.md", "v2.0 (2026-06)"],
            ["Documentos y Firma", "11", "URS_modulo_documentos_firma.md", "v2.0 (2026-06)"],
            ["Notificaciones", "9+", "URS_modulo_notificaciones_comunicaciones.md", "Completo"],
            ["Reportes y Auditoria", "8", "URS_modulo_reportes_auditoria.md", "v2.0 (2026-06)"],
            ["TI Soporte y Tickets", "11", "URS_modulo_ti_soporte_tickets.md", "v2.0 (2026-06)"],
            ["Inventario y Equipos", "19", "URS_modulo_inventario_equipos.md", "v2.0 (2026-06)"],
        ],
        widths=[5, 1.8, 8, 2.5],
    )

    doc.add_page_break()
    _h(doc, "A.2 Especificaciones funcionales (FRS) por modulo", level=1)
    _p(doc, (
        "Cada FRS documenta el endpoint exacto, las entradas reales, el proceso ejecutado "
        "y la salida esperada. La fuente completa esta en docs/validation/FRS/."
    ))
    _label(doc, "Modulo: Autenticacion  |  FRS-AUTH criticos", "1F497D")
    _table(doc,
        ["FRS-ID", "Endpoint", "Entradas requeridas", "Proceso ejecutado", "Salida HTTP"],
        [
            ["FRS-AUTH-001", "GET /api/v1/auth/google", "—", "Redirige a Google OAuth consent screen con scope openid+email+profile", "302 Redirect a Google"],
            ["FRS-AUTH-002", "GET /api/v1/auth/google/callback", "code, state (OAuth params)", "Intercambia code por tokens Google; valida dominio; upsert usuario en DB; emite JWT accessToken + refreshToken", "200 + { accessToken, refreshToken, user }"],
            ["FRS-AUTH-003", "POST /api/v1/auth/refresh", "Header x-refresh-token", "Verifica refreshToken; emite nuevo accessToken; actualiza user_sessions", "200 + { accessToken }"],
            ["FRS-AUTH-004", "POST /api/v1/auth/logout", "Bearer accessToken", "Invalida sesion activa en user_sessions; registra en auditoria.logs", "200 + { ok: true }"],
            ["FRS-AUTH-005", "GET /api/v1/auth/me", "Bearer accessToken", "Decodifica JWT; retorna perfil del usuario autenticado con rol, scope e id", "200 + { user }"],
        ],
        widths=[2, 4, 3.5, 5.5, 2.5],
    )
    _label(doc, "Modulo: Talento Humano  |  FRS-TH criticos", "1F497D")
    _table(doc,
        ["FRS-ID", "Endpoint", "Roles permitidos", "Proceso ejecutado", "Salida HTTP"],
        [
            ["FRS-TH-001", "POST /api/v1/talento-humano/permisos", "Todos los roles internos", "Valida campos obligatorios; crea registro en DB con estado 'pendiente'; dispara notificacion al aprobador", "201 + { id, estado: 'pendiente' }"],
            ["FRS-TH-002", "GET /api/v1/talento-humano/permisos", "talento_humano, gerencia, jefaturas", "Consulta solicitudes con filtros de estado, fecha y usuario; pagina resultados", "200 + { items, total }"],
            ["FRS-TH-003", "PUT /api/v1/talento-humano/permisos/:id/estado", "Aprobadores jerarquicos segun area", "Valida que el actor es aprobador autorizado; verifica estado actual; actualiza estado; registra en auditoria.logs", "200 + { id, estado_nuevo }"],
            ["FRS-TH-004", "GET /api/v1/talento-humano/permisos/:id", "Solicitante propietario, jefaturas, TH", "Retorna detalle completo con historial de estados y actor de cada transicion", "200 + { permiso, historial }"],
        ],
        widths=[2, 4.5, 3.5, 5.5, 2],
    )
    doc.add_page_break()


# ── ANEXO B — DDS ────────────────────────────────────────────────────────────


def anexo_b(doc):
    _portada_anexo(
        doc, "B", "ESPECIFICACION DE DISENO DEL SISTEMA (DDS)",
        "Evidencia: extraida de backend/src/app.js, routes/registerRoutes.js, "
        "routes/publicPaths.js, config/db.js, config/security.js y estructura de modulos. "
        "Fecha de extraccion: " + FECHA,
    )
    doc.add_page_break()

    _h(doc, "B.1 Arquitectura del sistema", level=1)
    _table(doc,
        ["Capa", "Tecnologia", "Version", "Archivo de configuracion", "Responsabilidad"],
        [
            ["Presentacion (SPA)", "React 19 + Tailwind CSS + Bootstrap", "19.x", "spi_front/vite.config.js", "Interfaces de usuario por rol; routing con React Router; state management local"],
            ["API / Logica de negocio", "Node.js + Express.js", "5.1.0 (pkg)", "backend/src/app.js", "Rutas REST; middlewares en cadena; controladores; servicios de negocio"],
            ["Persistencia", "PostgreSQL — Neon serverless", "pg 8.x", "backend/src/config/db.js", "Pool de 20 conexiones; SQL directo sin ORM; transacciones ACID; retry en errores transitorios"],
            ["Autenticacion", "Google OAuth 2.0 + JWT", "google-auth-library 10.x", "backend/src/middlewares/auth.js", "Sesiones con accessToken (corta duracion) y refreshToken; claims iss/aud/sub validados"],
            ["Seguridad HTTP", "Helmet + CORS + Rate Limit", "helmet 8.x", "backend/src/config/security.js", "CSP, HSTS, X-Frame; CORS dinamico; 3000 req/15min por usuario en produccion"],
            ["Almacenamiento de archivos", "Google Drive API", "googleapis 164.x", "backend/src/modules/documents/", "Documentos, adjuntos y actas; acceso via OAuth service account"],
            ["Proceso gestor", "PM2", "Latest", "ecosystem.config.js (servidor)", "Reinicio automatico; gestion de instancias; separacion de jobs por instancia"],
            ["Build frontend", "Vite", "Latest", "spi_front/vite.config.js", "Compilacion SPA; assets optimizados para produccion"],
        ],
        widths=[3, 3.8, 2.2, 4.5, 4],
    )

    _h(doc, "B.2 Cadena de middleware en app.js (orden de ejecucion)", level=1)
    _p(doc, "El siguiente orden fue verificado directamente en backend/src/app.js — define como cada request es procesado:")
    _code(doc, [
        "// backend/src/app.js — cadena de middleware verificada (lineas 81-138)",
        "",
        "app.use(helmet(helmetConfig))          // 1. Seguridad HTTP (CSP, HSTS, X-Frame)",
        "app.use(rateLimit({...}))              // 2. Rate limiting (3000 req/15min en prod)",
        "app.use(cors(corsConfig))              // 3. CORS dinamico segun lista de origenes",
        "app.use(express.json({ limit:'5mb'})) // 4. Parsing JSON del body",
        "app.use(requestContextMiddleware)      // 5. Correlacion de requests (x-correlation-id)",
        "app.use(mLogger)                       // 6. Logging estructurado de cada request",
        "",
        "mountPublicRoutes(app)                 // 7. Rutas publicas (auth, applicants, webhook CRM)",
        "",
        "app.use((req, res, next) => {          // 8. Guardia JWT global:",
        "  if (isPublicPath(req.path)) return next()",
        "  return verifyToken(req, res, next)   //    verifyToken rechaza con HTTP 401 sin token valido",
        "})",
        "",
        "app.use(normalizeApiPayloads)          // 9. Normalizacion de payloads de API legacy",
        "app.use(moduleAccessGuard)             // 10. Control de acceso a modulos por configuracion",
        "app.use(auditMiddleware)               // 11. Auditoria: registra POST/PUT/PATCH/DELETE",
        "",
        "mountPrivateRoutes(app)                // 12. Rutas privadas (todos los modulos del sistema)",
    ])

    _h(doc, "B.3 Segmentacion de rutas publicas vs. privadas", level=1)
    _p(doc, "Fuente: backend/src/routes/publicPaths.js y backend/src/routes/registerRoutes.js")
    _label(doc, "Rutas publicas (sin JWT requerido) — publicPaths.js", "37490D")
    _table(doc,
        ["Ruta", "Motivo de excepcion"],
        [
            ["/api/v1/auth/google*", "Inicio del flujo OAuth2 — el usuario no tiene token todavia"],
            ["/api/v1/auth/google/callback", "Callback de Google — procesa el code y emite el JWT"],
            ["/api/v1/permisos/legal-verification/*", "Verificacion publica de firma de permiso por QR"],
            ["/api/v1/vacaciones/legal-verification/*", "Verificacion publica de firma de vacaciones por QR"],
            ["/api/verificar, /api/verify, /api/signature/verificar*", "Verificacion publica de documentos firmados por QR"],
            ["/api/v1/integrations/crm/webhook", "Webhook de EspoCRM validado por X-Hook-Secret (no por JWT)"],
            ["/health", "Endpoint de health check para el proceso gestor PM2"],
            ["/ws", "WebSocket (sin autenticacion en el handshake inicial)"],
        ],
        widths=[7, 10.5],
    )

    _label(doc, "Rutas privadas (requieren JWT valido) — registerRoutes.js — mountPrivateRoutes()", "1F497D")
    _table(doc,
        ["Ruta API", "Modulo", "Area de validacion"],
        [
            ["/api/v1/requests", "requests (solicitudes operacionales)", "Area 03"],
            ["/api/v1/talento-humano", "talento_humano (permisos, vacaciones, RRHH)", "Area 02"],
            ["/api/v1/auditoria", "auditoria (bitacora de eventos)", "Area 06"],
            ["/api/v1/security", "security (eventos fuera de horario, alertas)", "Area 01"],
            ["/api/v1/users + /api/v1/collaborators", "users / collaborators (usuarios y perfiles)", "Area 01"],
            ["/api/v1/business-case + /api/v1/clients", "business-case / clients", "Area 03"],
            ["/api/v1/viaticos", "viaticos (workspace, wizard, flujo financiero)", "Area 06"],
            ["/api/v1/servicio + /api/v1/mantenimientos", "servicio tecnico y mantenimientos", "Area 04"],
            ["/api/v1/documents + /api/v1/signature*", "documentos y firma digital", "Area 06"],
            ["/api/v1/notifications", "notificaciones (in-app + cola de dispatch)", "Area 06"],
            ["/api/v1/dashboard", "dashboard (6 KPIs en paralelo)", "Area 06"],
            ["/api/v1/support-tickets + /api/v1/ti-assets", "TI soporte e inventario de activos", "Area 05"],
            ["/api/v1/inventario + /api/v1/equipment-management", "inventario y gestion de equipos", "Area 05"],
        ],
        widths=[5.5, 6, 2.5],
    )

    _h(doc, "B.4 Configuracion de base de datos — db.js", level=1)
    _p(doc, "Fuente: backend/src/config/db.js — verificado en el repositorio")
    _table(doc,
        ["Parametro", "Valor configurado", "Evidencia en codigo"],
        [
            ["Motor", "PostgreSQL via Neon serverless (cloud)", "const { Pool } = require('pg')"],
            ["Pool maximo de conexiones", "20 conexiones simultaneas", "max: intFromEnv('DB_POOL_MAX', 20)"],
            ["Pool minimo de conexiones activas", "2 conexiones calientes permanentes", "min: intFromEnv('DB_POOL_MIN', 2)"],
            ["Timeout de conexion inactiva", "30,000 ms (30 segundos)", "idleTimeoutMillis: intFromEnv('DB_IDLE_TIMEOUT_MS', 30000)"],
            ["Timeout para establecer conexion", "5,000 ms (5 segundos)", "connectionTimeoutMillis: intFromEnv('DB_CONN_TIMEOUT_MS', 5000)"],
            ["Max usos por conexion antes de reciclar", "7,500 queries", "maxUses: intFromEnv('DB_POOL_MAX_USES', 7500)"],
            ["SSL en la conexion", "Configurable via DB_SSL=true", "ssl: process.env.DB_SSL === 'true' ? { rejectUnauthorized: ... }"],
            ["Retry en errores transitorios", "2 intentos con 250ms de espera entre cada uno", "maxAttempts = intFromEnv('DB_QUERY_RETRY_ATTEMPTS', 2)"],
            ["Deteccion de errores transitorios", "ETIMEDOUT, ECONNRESET, 57P01, 57P03, 53300", "isTransientDbError() en db.js lineas 18-30"],
            ["Shutdown graceful", "Cierra el pool en SIGINT y SIGTERM", "setupGracefulShutdown() — pool.end()"],
        ],
        widths=[5, 4.5, 8],
    )

    _h(doc, "B.5 Dependencias criticas del sistema", level=1)
    _p(doc, "Fuente: backend/package.json — version 1.0.0")
    _table(doc,
        ["Paquete", "Version declarada", "Funcion en el sistema"],
        [
            ["express", "^5.1.0", "Framework HTTP del backend — enrutamiento y middlewares"],
            ["jsonwebtoken", "(via auth.js)", "Emision y verificacion de JWT con SECRET_KEY"],
            ["google-auth-library", "^10.5.0", "Verificacion de tokens Google OAuth2"],
            ["googleapis", "^164.1.0", "Integracion con Drive y Gmail APIs"],
            ["pg", "(via db.js)", "Driver PostgreSQL — pool de conexiones"],
            ["helmet", "^8.1.0", "Cabeceras de seguridad HTTP (CSP, HSTS, etc.)"],
            ["express-rate-limit", "^8.2.1", "Limitador de tasa de requests por usuario/IP"],
            ["cors", "^2.8.5", "Control de origenes permitidos (CORS dinamico)"],
            ["csv-stringify", "^6.6.0", "Exportacion de bitacora de auditoria a CSV"],
            ["adm-zip", "^0.5.14", "Manejo de archivos ZIP para exportaciones"],
            ["axios", "^1.13.1", "Cliente HTTP para integraciones externas"],
            ["bcryptjs", "^3.0.2", "Hash de datos sensibles donde aplica"],
        ],
        widths=[4, 3, 10.5],
    )
    doc.add_page_break()


# ── ANEXO C — FMEA ───────────────────────────────────────────────────────────


def anexo_c(doc):
    _portada_anexo(
        doc, "C", "MATRIZ DE EVALUACION DE RIESGOS (FMEA)",
        "Metodologia: RPN = Severidad x Probabilidad x Detectabilidad (escala 1-5). "
        "Fuente: docs/validation/general/07A_evaluacion_riesgos_fmea.md. "
        "Fecha de extraccion: " + FECHA,
    )
    doc.add_page_break()

    _h(doc, "C.1 Escala de evaluacion", level=1)
    _table(doc,
        ["Valor", "Severidad", "Probabilidad", "Detectabilidad"],
        [
            ["1", "Insignificante", "Muy baja (< 1% de los ciclos)", "Detectado inmediatamente"],
            ["2", "Menor", "Baja (1-5% de los ciclos)", "Detectado rapidamente"],
            ["3", "Moderado", "Media (5-20% de los ciclos)", "Detectado en revision"],
            ["4", "Mayor", "Alta (20-50% de los ciclos)", "Dificil de detectar"],
            ["5", "Critico / Catastrofico", "Muy alta (> 50% de los ciclos)", "Muy dificil de detectar"],
        ],
        widths=[1.5, 3.5, 4.5, 4.5],
    )
    _p(doc, "Umbral de accion: RPN >= 30 o Severidad = 5 sin mitigacion bloquea el inicio de IQ.", italic=True)

    _h(doc, "C.2 Matriz FMEA completa", level=1)
    _table(doc,
        ["ID", "Modulo / Area", "Modo de Falla", "Efecto potencial", "S", "P", "D", "RPN", "Mitigacion implementada", "Riesgo residual", "Estado"],
        [
            ["R-001", "Infraestructura", "Interrupcion del servicio de contenedor/proceso", "Indisponibilidad total del sistema para todos los usuarios", "5", "2", "2", "20", "PM2 con reinicio automatico; monitoreo de uptime; watchdog", "Bajo (RPN 10)", "Mitigado"],
            ["R-002", "auth", "Compromiso de credenciales OAuth2 / robo de token JWT", "Acceso no autorizado al sistema por tercero", "5", "2", "3", "30", "OAuth2 Google con dominio restringido; JWT de corta duracion; claims iss/aud/sub validados; refresh token rotation", "Medio (RPN 15)", "ACTIVO — verificar en OQ-014"],
            ["R-003", "auth", "Sesion no invalidada tras logout explicito", "Acceso residual con token robado post-logout", "4", "2", "3", "24", "Registro de sesion en user_sessions; expiracion corta del accessToken", "Bajo (RPN 12)", "Mitigado"],
            ["R-004", "security", "Login fuera de horario laboral no detectado", "Acceso no autorizado sin alerta ni registro", "4", "3", "2", "24", "Modulo security detecta off-hours; notificacion automatica al canal de TI", "Bajo (RPN 8)", "Mitigado"],
            ["R-005", "auditoria", "Perdida de registros de bitacora de auditoria", "No trazabilidad de eventos criticos; incumplimiento regulatorio", "5", "1", "2", "10", "auditoria.logs en PostgreSQL persistente; usuarios no pueden eliminar registros", "Bajo (RPN 5)", "Mitigado"],
            ["R-006", "Base de datos (PostgreSQL)", "Corrupcion o perdida total de datos", "Perdida de informacion operativa critica", "5", "1", "3", "15", "Transacciones ACID; pool con retry en errores transitorios; backups Neon segun plan contratado", "Bajo (RPN 5)", "Mitigado"],
            ["R-007", "signature", "Firma avanzada invalidada o no generada", "Documentos sin validez legal / regulatoria", "5", "2", "2", "20", "Hash SHA-256 del contenido; QR de verificacion publica; cadena inmutable en document_audit_log", "Bajo (RPN 8)", "Mitigado"],
            ["R-008", "documents", "Generacion de documento corrupto (PDF/ZIP)", "Entrega de documento invalido al destinatario", "3", "2", "2", "12", "Validacion post-generacion; reintentos automaticos; registro del resultado en auditoria", "Bajo (RPN 6)", "Mitigado"],
            ["R-009", "Usuarios (RBAC)", "Asignacion incorrecta de rol a un usuario", "Acceso a funciones no autorizadas para ese perfil", "4", "2", "3", "24", "requireRole en cada endpoint; doble capa con guards assertX() en servicios criticos; SUPER_ROLES solo admin", "Bajo (RPN 8)", "Mitigado"],
            ["R-010", "Gmail/SMTP", "Falla en envio de correo de notificacion", "Notificacion critica no entregada al destinatario", "2", "3", "2", "12", "Cola de dispatch con reintentos y backoff; estados pending/sent/failed; canal in-app siempre disponible", "Bajo (RPN 4)", "Mitigado"],
            ["R-011", "permisos / vacaciones", "Calculo incorrecto de saldo de dias", "Error en nomina, control laboral o pago de beneficios", "4", "2", "3", "24", "Validacion unitaria; pruebas OQ/PQ especificas para flujo de saldos; historial auditable", "Medio (RPN 12)", "Bajo seguimiento"],
            ["R-012", "attendance", "Registro de marcacion duplicado en asistencia", "Horas extra incorrectas que afectan nomina", "3", "2", "3", "18", "Idempotencia en API de marcacion; validacion de timestamp unico por colaborador y fecha", "Bajo (RPN 6)", "Mitigado"],
            ["R-013", "personnel-requests", "Datos del candidato incompletos en contratacion", "Expediente laboral incompleto; riesgo legal en contratacion", "3", "3", "3", "27", "Validaciones de campos obligatorios en frontend y backend; checklist de documentos por estado", "Bajo (RPN 9)", "Mitigado"],
            ["R-014", "Entorno produccion", "Despliegue en entorno incorrecto (dev sobre datos reales)", "Ejecucion de pruebas sobre datos productivos no intencionados", "5", "1", "2", "10", "Separacion por NODE_ENV; variables de entorno por ambiente; IQ verifica entorno antes de ejecucion", "Bajo (RPN 5)", "Mitigado"],
            ["R-015", "Confidencialidad LOPDP", "Exposicion de datos personales sensibles de colaboradores", "Incumplimiento de la Ley Organica de Proteccion de Datos (Ecuador)", "5", "1", "3", "15", "RBAC con minimo privilegio; HTTPS/TLS obligatorio; aceptacion LOPDP en primer login; campos sensibles no expuestos en logs", "Bajo (RPN 5)", "Mitigado"],
        ],
        widths=[1.3, 2.5, 3.2, 3, 0.6, 0.6, 0.6, 1, 3.5, 1.8, 2],
        font=7,
    )

    _h(doc, "C.3 Resumen por nivel de riesgo", level=1)
    _table(doc,
        ["Nivel", "Umbral RPN", "Cantidad", "IDs", "Accion requerida"],
        [
            ["Critico", ">= 50 o S=5 sin mitigacion", "0", "Ninguno", "Bloquearia inicio de IQ — no aplica en FamSPI v1.0.0"],
            ["Alto", "30 - 49", "1", "R-002 (OAuth2)", "Verificacion obligatoria en OQ-014 del Area 01 antes de liberacion"],
            ["Medio", "20 - 29", "4", "R-003, R-004, R-009, R-011", "Seguimiento en ejecucion de OQ y PQ; no bloquea inicio"],
            ["Bajo", "< 20", "10", "R-001, R-005 a R-008, R-010, R-012 a R-015", "Monitoreo en operacion normal; sin accion de validacion adicional"],
        ],
        widths=[2, 2.5, 2, 4, 7],
    )

    _p(doc, "Elaborado por: Departamento de TI — Revisado y aprobado: pendiente firma formal")
    doc.add_page_break()


# ── ANEXO D — EVIDENCIAS OBJETIVAS (CODIGO FUENTE) ───────────────────────────


def anexo_d(doc):
    _portada_anexo(
        doc, "D", "EVIDENCIAS OBJETIVAS — CODIGO FUENTE VERIFICADO",
        "Las evidencias de este anexo son extractos reales del codigo fuente de FamSPI v1.0.0. "
        "Cada bloque muestra la implementacion exacta del control validado, con referencia al "
        "archivo fuente y numero de linea. Fecha de extraccion: " + FECHA,
    )
    doc.add_page_break()

    # EV-D-001
    _h(doc, "EV-D-001 — Middleware de autenticacion JWT (verifyToken)", level=1)
    _p(doc, "Archivo: backend/src/middlewares/auth.js | Verificado el: " + FECHA)
    _p(doc, "Evidencia del control URS-S-001 — autenticacion obligatoria en todas las rutas privadas.")
    _table(doc,
        ["Aspecto verificado", "Implementacion confirmada"],
        [
            ["Token requerido en cabecera", "Verifica 'Authorization', 'authorization' o 'x-access-token'; HTTP 401 si ausente"],
            ["Soporte de formato Bearer", "Extrae el token correctamente de 'Bearer <token>'"],
            ["Verificacion criptografica", "jwt.verify(token, process.env.SECRET_KEY) — falla con HTTP 401 si invalido o expirado"],
            ["Validacion de claims", "Verifica iss='spi-fam-backend', aud='spi-fam-frontend' y presencia de sub; HTTP 403 si falla"],
            ["Propagacion de identidad", "req.user = { ...decoded, ip, userAgent } — disponible para todos los middlewares siguientes"],
            ["Manejo de errores", "Try/catch en dos niveles; HTTP 500 en errores internos inesperados"],
        ],
        widths=[5, 12.5],
    )
    _code(doc, [
        "// backend/src/middlewares/auth.js",
        "const verifyToken = (req, res, next) => {",
        "  const headerAuth = req.headers['authorization'] || req.headers['x-access-token'];",
        "  if (!headerAuth) {",
        "    return res.status(401).json({ ok: false, code: 'NO_TOKEN', message: 'Token ausente' });",
        "  }",
        "  const token = headerAuth.startsWith('Bearer ') ? headerAuth.split(' ')[1] : headerAuth;",
        "  let decoded;",
        "  try {",
        "    decoded = jwt.verify(token, process.env.SECRET_KEY);",
        "  } catch (err) {",
        "    return res.status(401).json({ ok: false, code: 'INVALID_TOKEN', message: 'Token invalido o expirado' });",
        "  }",
        "  if (decoded.iss !== 'spi-fam-backend' || decoded.aud !== 'spi-fam-frontend' || !decoded.sub) {",
        "    return res.status(403).json({ ok: false, code: 'INVALID_CLAIMS', message: 'Token no valido para esta aplicacion' });",
        "  }",
        "  req.user = { ...decoded, ip: req.headers['x-forwarded-for']?.split(',')[0] || req.ip };",
        "  next();",
        "};",
    ])

    # EV-D-002
    _h(doc, "EV-D-002 — Middleware de autorizacion por rol (requireRole)", level=1)
    _p(doc, "Archivo: backend/src/middlewares/roles.js | Verificado el: " + FECHA)
    _p(doc, "Evidencia del control URS-S-002 — restriccion por rol en todas las rutas.")
    _table(doc,
        ["Aspecto verificado", "Implementacion confirmada"],
        [
            ["Grupos de roles definidos", "ROLE_GROUPS con 15 grupos: comercial, tecnico, ti, finanzas, talento_humano, gerencia, etc."],
            ["Super roles sin restriccion", "SUPER_ROLES = new Set(['admin', 'administrador']) — bypass total"],
            ["Normalizacion de rol", "normalizeRoleName() — lowercase + reemplazo de espacios/guiones por underscore"],
            ["Expansion de grupos", "expandRoles() expande cada rol recibido a todos los sub-roles del grupo"],
            ["Multi-rol de usuario", "collectUserRoles() recopila de user.role, user.scope, user.role_name, user.roles[]"],
            ["Respuesta de acceso denegado", "HTTP 403 con lista de roles permitidos; nunca revela informacion del sistema"],
        ],
        widths=[5, 12.5],
    )
    _code(doc, [
        "// backend/src/middlewares/roles.js",
        "const ROLE_GROUPS = {",
        "  ti:       ['ti', 'jefe_ti', 'jefe_de_ti', 'desarrollador', 'soporte'],",
        "  finanzas: ['finanzas', 'financiero', 'jefe_finanzas', 'jefe_de_finanzas', 'contador', 'jefe_financiero'],",
        "  gerencia: ['gerencia', 'gerencia_general', 'gerente_general', 'director', 'gerente'],",
        "  // ... 12 grupos adicionales",
        "};",
        "const SUPER_ROLES = new Set(['admin', 'administrador']);",
        "",
        "function requireRole(allowedRoles = []) {",
        "  const expanded = expandRoles(allowedRoles);  // expande grupos a roles individuales",
        "  return (req, res, next) => {",
        "    if (!req.user) return res.status(401).json({ ok: false, error: 'No autenticado.' });",
        "    const candidates = collectUserRoles(req.user);",
        "    for (const role of candidates) {",
        "      if (SUPER_ROLES.has(role)) return next();  // admin siempre pasa",
        "      if (expanded.has(role)) return next();     // rol en lista permitida",
        "    }",
        "    return res.status(403).json({ ok: false, error: `Acceso denegado. Roles permitidos: ${[...expanded].join(', ')}` });",
        "  };",
        "}",
    ])

    # EV-D-003
    _h(doc, "EV-D-003 — Middleware de auditoria (auditMiddleware)", level=1)
    _p(doc, "Archivo: backend/src/middlewares/auditMiddleware.js | Verificado el: " + FECHA)
    _p(doc, "Evidencia del control URS-T-001 — trazabilidad de todas las acciones criticas.")
    _table(doc,
        ["Aspecto verificado", "Implementacion confirmada"],
        [
            ["Metodos auditados", "POST, PUT, PATCH, DELETE — los GET no se auditan (solo lectura)"],
            ["Exclusiones", "Rutas /api/v1/auth/* excluidas (el login no se audita en esta tabla)"],
            ["Datos registrados", "usuario_id, usuario_email, rol, modulo, accion, descripcion, ip, user_agent, duracion_ms, contexto, datos_nuevos"],
            ["Resolucion de modulo", "resolveModuleFromPath() extrae el modulo del path '/api/v1/{modulo}/...'"],
            ["Mapeo de accion", "mapHttpMethodToAction(): POST->crear/aprobar/firmar, PUT->actualizar, DELETE->eliminar"],
            ["Truncado de payload", "truncateBody() limita el cuerpo a 800 chars para evitar registros masivos"],
            ["Contexto relacional", "buildContext() enriquece con request_id, mantenimiento_id, inventario_id si aplica"],
            ["Registro asincrono", "logAction() se ejecuta en res.on('finish') — no bloquea la respuesta al cliente"],
        ],
        widths=[5, 12.5],
    )
    _code(doc, [
        "// backend/src/middlewares/auditMiddleware.js (extracto critico)",
        "async function auditMiddleware(req, res, next) {",
        "  const method = req.method.toUpperCase();",
        "  if (!['POST','PUT','PATCH','DELETE'].includes(method)) return next();",
        "  if (req.originalUrl.startsWith('/api/v1/auth')) return next();",
        "  const start = Date.now();",
        "  res.on('finish', async () => {",
        "    const modulo  = resolveModuleFromPath(pathParts);   // ej: 'viaticos', 'permisos'",
        "    const accion  = mapHttpMethodToAction(method, pathParts); // ej: 'crear', 'aprobar'",
        "    await logAction({",
        "      usuario_id, usuario_email, rol, modulo,",
        "      accion: success ? accion : `${accion}_failed`,",
        "      descripcion, datos_nuevos: truncateBody(req.body),",
        "      ip, user_agent, contexto, duracion_ms: Date.now() - start,",
        "    });",
        "  });",
        "  next();",
        "}",
    ])

    # EV-D-004 - app.js middleware chain
    _h(doc, "EV-D-004 — Aplicacion de middlewares en produccion (app.js)", level=1)
    _p(doc, "Archivo: backend/src/app.js | Verificado el: " + FECHA)
    _p(doc, "Evidencia del orden de procesamiento de cada request y del punto de aplicacion del JWT global.")
    _code(doc, [
        "// backend/src/app.js — aplicacion del JWT global (lineas 125-131)",
        "// Aplica a TODOS los paths que no esten en PUBLIC_PATH_PREFIXES",
        "app.use((req, res, next) => {",
        "  if (isPublicPath(req.path)) return next();",
        "  return verifyToken(req, res, next);  // <-- JWT obligatorio aqui",
        "});",
        "",
        "// Tras el JWT, se aplican en este orden:",
        "app.use(normalizeApiPayloads);   // normalizacion",
        "app.use(moduleAccessGuard);      // control de acceso a modulos",
        "app.use(auditMiddleware);        // AUDITORIA de mutaciones",
        "mountPrivateRoutes(app);         // logica de negocio",
    ])

    # EV-D-005 - publicPaths
    _h(doc, "EV-D-005 — Rutas publicas declaradas (publicPaths.js)", level=1)
    _p(doc, "Archivo: backend/src/routes/publicPaths.js | Verificado el: " + FECHA)
    _p(doc, "Evidencia del control DQ-04 — segmentacion explicita de rutas publicas vs. privadas.")
    _code(doc, [
        "// backend/src/routes/publicPaths.js",
        "const PUBLIC_PATH_PREFIXES = [",
        "  '/ws',",
        "  '/internal/jobs',",
        "  '/api/v1/equipment-purchases/events',",
        "  '/api/v1/private-purchases/events',",
        "  '/api/v1/auth/google',",
        "  '/api/v1/gmail/auth/callback',",
        "  '/api/v1/permisos/legal-verification/',",
        "  '/api/v1/vacaciones/legal-verification/',",
        "  '/health',",
        "  '/api/verificar',",
        "  '/api/verify',",
        "  '/api/signature/verificar',",
        "  '/api/signature/verify',",
        "  '/api/v1/signature/verificar',",
        "  '/api/v1/signature/verify'",
        "];",
        "// Todas las demas rutas REQUIEREN JWT valido.",
    ])

    doc.add_page_break()


# ── ANEXO E — REGISTROS DE ENTRENAMIENTO ─────────────────────────────────────


def anexo_e(doc):
    _portada_anexo(
        doc, "E", "REGISTROS DE CAPACITACION Y ENTRENAMIENTO",
        "La OMS establece que el personal debe estar capacitado con manuales y procedimientos "
        "escritos ANTES de la ejecucion de PQ/UAT. "
        "Las listas de asistencia firmadas son la evidencia de cumplimiento. "
        "Fecha de elaboracion del formulario: " + FECHA,
    )
    doc.add_page_break()

    _h(doc, "E.1 Plan de entrenamiento por perfil", level=1)
    _table(doc,
        ["Perfil", "Modulos cubiertos", "Duracion", "Instructor", "Estado", "Fecha"],
        [
            ["Solicitante general\n(comercial, tecnico, backoffice, RRHH)", "Dashboard, solicitudes de permisos y vacaciones, workspace de viaticos (wizard + carga TXT + notas manuales), notificaciones in-app", "2-4 horas", "TI", "Pendiente", "___________"],
            ["Aprobador / Jefe de area", "Todo lo del solicitante + flujos de aprobacion y rechazo por area, consulta de historial, trazabilidad de decisiones", "3-5 horas", "TI", "Pendiente", "___________"],
            ["Finanzas", "Viaticos: revision financiera, categorizacion de facturas, registro de pago, exportacion ATS/XML, dashboard de reportes", "4-6 horas", "TI", "Pendiente", "___________"],
            ["TI — Administrador del sistema", "Todo el sistema con acceso admin: usuarios/roles, bitacora de auditoria, modo auditoria, soporte tickets, activos TI, respaldo, monitoreo de jobs", "8-12 horas en sesiones", "TI (auto-capacitacion documentada)", "Parcial", "___________"],
        ],
        widths=[4, 6, 2, 2.5, 2, 2],
    )

    _h(doc, "E.2 Registro de asistencia — Perfil: Solicitante general", level=2)
    _p(doc, "Modulo capacitado: Dashboard + Permisos/Vacaciones + Viaticos + Notificaciones")
    _p(doc, "Instructor: ___________________________   Fecha: _______________   Duracion: _____ horas")
    _table(doc,
        ["#", "Nombre completo", "Cargo / Departamento", "Correo corporativo", "Firma"],
        [["1", "", "", "", ""],
         ["2", "", "", "", ""],
         ["3", "", "", "", ""],
         ["4", "", "", "", ""],
         ["5", "", "", "", ""],
         ["6", "", "", "", ""],
         ["7", "", "", "", ""],
         ["8", "", "", "", ""]],
        widths=[1, 4.5, 4, 5, 4],
    )
    _p(doc, "Firma del instructor: _________________________   Fecha: _______________")
    doc.add_paragraph()

    _h(doc, "E.3 Registro de asistencia — Perfil: Aprobador / Jefe de area", level=2)
    _p(doc, "Modulo capacitado: Todo lo del solicitante + flujos de aprobacion por area")
    _p(doc, "Instructor: ___________________________   Fecha: _______________   Duracion: _____ horas")
    _table(doc,
        ["#", "Nombre completo", "Cargo / Area", "Correo corporativo", "Firma"],
        [["1", "", "", "", ""],
         ["2", "", "", "", ""],
         ["3", "", "", "", ""],
         ["4", "", "", "", ""],
         ["5", "", "", "", ""]],
        widths=[1, 4.5, 4, 5, 4],
    )
    _p(doc, "Firma del instructor: _________________________   Fecha: _______________")
    doc.add_paragraph()

    _h(doc, "E.4 Registro de asistencia — Perfil: Finanzas", level=2)
    _p(doc, "Modulo capacitado: Viaticos financiero + Reportes + Exportacion ATS")
    _p(doc, "Instructor: ___________________________   Fecha: _______________   Duracion: _____ horas")
    _table(doc,
        ["#", "Nombre completo", "Cargo / Area", "Correo corporativo", "Firma"],
        [["1", "", "", "", ""],
         ["2", "", "", "", ""],
         ["3", "", "", "", ""]],
        widths=[1, 4.5, 4, 5, 4],
    )
    _p(doc, "Firma del instructor: _________________________   Fecha: _______________")
    doc.add_paragraph()

    _h(doc, "E.5 Registro de asistencia — Perfil: TI Administrador", level=2)
    _p(doc, "Modulo capacitado: Sistema completo con acceso admin — auditoria, soporte, activos, respaldo")
    _p(doc, "Instructor: ___________________________   Fecha: _______________   Duracion: _____ horas")
    _table(doc,
        ["#", "Nombre completo", "Cargo", "Correo corporativo", "Firma"],
        [["1", "", "", "", ""],
         ["2", "", "", "", ""]],
        widths=[1, 4.5, 4, 5, 4],
    )
    _p(doc, "Firma del instructor: _________________________   Fecha: _______________")

    doc.add_paragraph()
    doc.add_paragraph()
    fin = doc.add_paragraph()
    fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fin.add_run("— Fin de los Anexos del Protocolo de Validacion FamSPI v1.0.0 —")
    r.italic = True
    r.font.size = Pt(9)
    r.font.name = "Calibri"
    r.font.color.rgb = GRIS


# ── Construccion ─────────────────────────────────────────────────────────────


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.2)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # Portada general
    doc.add_paragraph()
    doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("ANEXOS — PROTOCOLO DE VALIDACION FAMSPI v1.0.0")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.name = "Calibri"
    tr.font.color.rgb = AZUL
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"Departamento de TI  |  {FECHA}  |  Evidencia extraida del codigo fuente")
    sr.italic = True
    sr.font.size = Pt(10)
    sr.font.name = "Calibri"
    sr.font.color.rgb = GRIS
    doc.add_page_break()

    anexo_a(doc)
    anexo_b(doc)
    anexo_c(doc)
    anexo_d(doc)
    anexo_e(doc)

    doc.save(OUTPUT)
    print(f"Anexos generados: {OUTPUT}")


if __name__ == "__main__":
    build()
